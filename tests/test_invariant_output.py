"""THE fabrication gate: a line identical across disjoint inputs was not derived from them.

THE METHOD

Render the whole application, in all four formats, under N genuinely disjoint
scenarios — no shared CDE, project set, state, sector or dollar amount between
any two. Extract the prose back out of each artifact. Any line that is
byte-identical across all N scenarios cannot have come from those inputs,
because the inputs have nothing in common. It is therefore either

  (a) legitimately invariant — a heading, a column label, placeholder text, or
      a statutory constant that carries its citation inline; every such line is
      on ALLOWLIST below with a one-line justification, or

  (b) an assertion the tool makes about every CDE regardless of who is asking,
      which is the definition of the defect this release exists to remove.

Anything in (b) fails the test.

WHY THIS AND NOT A DENYLIST

tests/test_no_fabricated_output.py is an eighteen-string denylist. Measured
against reintroduced fabrications: 5/5 verbatim caught, 0/8 paraphrases caught.
A denylist can only contain strings somebody already thought of, so it cannot
catch a NEW fabrication by construction. This test inverts the burden — a new
invariant assertion fails by default and has to be argued onto the allowlist by
a human — and it caught all 8 paraphrases the denylist missed.

READING THE ALLOWLIST IS THE HIGHEST-VALUE REVIEW ON THIS PACKAGE. Every entry
is a sentence the tool says to the CDFI Fund about every CDE that ever uses it.

FAILS CLOSED. Fewer than N scenarios, an empty extraction from any artifact, an
empty invariant set, or a missing format all ERROR rather than pass.
"""
from __future__ import annotations

import os
import re

import pytest

from nmtcapp.core.application import Application
from nmtcapp.core.cde import CDEProfile
from nmtcapp.core.pipeline import Pipeline, PipelineProject

FORMATS = ("markdown", "word", "excel", "pdf")

# Four scenarios sharing NO CDE, project, state, sector or dollar amount.
MIN_SCENARIOS = 4


# ---------------------------------------------------------------------------
# Disjoint scenarios
# ---------------------------------------------------------------------------

def _cde(n: int) -> CDEProfile:
    names = {
        1: ("Aurora Borealis Capital CDE, LLC", "CDE-2011-0001", "2011-01-05"),
        2: ("Bayou Delta Growth Fund CDE", "CDE-2015-0222", "2015-07-19"),
        3: ("Cascadia Tribal Lending CDE, Inc.", "CDE-2020-0333", "2020-11-30"),
        4: ("Dakota Prairie Ventures CDE", "CDE-2024-0444", "2024-03-02"),
    }[n]
    markets = {
        1: ["Alaska"],
        2: ["Louisiana", "Mississippi", "Alabama", "Arkansas", "Oklahoma", "Texas"],
        3: ["Oregon", "Washington", "Idaho"],
        4: ["North Dakota", "South Dakota", "Minnesota", "Iowa", "Nebraska",
            "Missouri", "Kansas"],
    }[n]
    return CDEProfile(
        name=names[0], cde_id=names[1], certification_date=names[2],
        mission={
            1: "Finance cold-climate health infrastructure.",
            2: "Rebuild coastal small business after storm loss.",
            3: "Deploy capital with tribal governments on trust land.",
            4: "Back grain-belt manufacturing in non-metro counties.",
        }[n],
        target_markets=markets,
        prior_awards=[] if n in (3, 4) else [
            {"year": 2010 + n, "amount": (7_000_000 * n), "deployment_status": "fully_deployed"}
        ],
        contact={"name": f"Contact {n}", "email": f"c{n}@example.org"},
        governance={"board_members": 4 + n, "community_representatives": n},
        extra={} if n != 2 else {"has_prior_reporting_issues": True},
    )


# Each scenario differs in EVERY dimension a rendered line could key on:
# state set, sector, project type, distress mix, project count, dollar scale,
# jobs, units, square footage and eligibility rate. Where a line is genuinely
# derived from the input it therefore varies and drops out of the invariant
# set — which is what keeps the allowlist honest. If two scenarios agreed on,
# say, "single state", the "expand geographic footprint" recommendation would
# look invariant and demand an allowlist entry it does not deserve.
def _pipeline(n: int) -> Pipeline:
    spec = {
        1: dict(states=["AK"], city="Nome", sector="healthcare",
                project_type="real_estate", distress=["deep", "severe"], count=2,
                cost=3_100_000, qei=2_100_000, jobs=240, tract="02180000100",
                units=0, sqft=1_200, eligible=[True, True]),
        2: dict(states=["LA", "MS", "AL", "AR", "OK", "TX"], city="Houma",
                sector="small_business", project_type="operating_business",
                distress=["severe"], count=6,
                cost=4_700_000, qei=3_300_000, jobs=27, tract="22109040300",
                units=140, sqft=9_500, eligible=[True] * 6),
        3: dict(states=["OR", "WA", "ID"], city="Warm Springs",
                sector="community_facility", project_type="real_estate",
                distress=["lic", "ineligible", "lic"], count=3,
                cost=9_200_000, qei=6_400_000, jobs=43, tract="41013970100",
                units=7, sqft=58_000, eligible=[True, False, True]),
        4: dict(states=["ND", "SD", "MN", "IA", "NE", "MO", "KS"],
                city="Jamestown", sector="clean_energy",
                project_type="operating_business", distress=["deep"], count=7,
                cost=15_800_000, qei=11_900_000, jobs=64, tract="38093960200",
                units=55, sqft=310_000, eligible=[True] * 7),
    }[n]
    projects = []
    for i in range(spec["count"]):
        p = PipelineProject(
            project_id=f"S{n}-{i:02d}",
            project_name=f"Scenario {n} Project {i}",
            qalicb_name=f"S{n}-{i:02d} QALICB LLC",
            address=f"{100 * n + i} Route {n}",
            city=spec["city"], state=spec["states"][i % len(spec["states"])],
            sector=spec["sector"], project_type=spec["project_type"],
            total_project_cost=float(spec["cost"] + i),
            qei_request=float(spec["qei"] + i),
            qlici_amount=float(spec["qei"] + i),
            expected_jobs_created=spec["jobs"] + i,
            expected_jobs_retained=n + i,
            # SUPPLIED, INCLUDING A SUPPLIED ZERO (FIX-2 B-2). This read
            # `spec["units"] + i if spec["units"] else None`, so scenario 1's
            # declared units=0 became "the CDE supplied nothing" — the exact
            # collapse B-2 removed from impact_aggregator, reproduced in the
            # fixture that is supposed to exercise it. Scenario 1 now supplies
            # a genuine 0 and every scenario supplies a real square footage.
            expected_units_built=spec["units"] + i,
            expected_sq_ft=float(spec["sqft"] + i),
        )
        p.census_tract = spec["tract"]
        p.is_nmtc_eligible = spec["eligible"][i]
        p.distress_level = spec["distress"][i % len(spec["distress"])]
        p.is_native_area = n == 3
        p.is_high_migration_rural = n == 4
        p.is_opportunity_zone = n == 2
        p.is_us_territory = False
        p.is_persistent_poverty = n in (1, 2)
        p.is_below_market_rate = n in (1, 4)
        p.is_unrelated_entity = True
        p.geocode_success = True
        projects.append(p)
    pipeline = Pipeline(projects)
    # These fixtures carry verified-shaped data on purpose; the adapter's
    # pre-enriched branch would otherwise route every scenario into degraded
    # mode and the four would then share their degraded banner text.
    pipeline.eligibility_data_status = "ok"
    return pipeline


SCENARIOS = {
    1: (_cde(1), _pipeline(1), 2_000_000.0),
    2: (_cde(2), _pipeline(2), 17_000_000.0),
    3: (_cde(3), _pipeline(3), 19_500_000.0),
    4: (_cde(4), _pipeline(4), 84_000_000.0),
}


# ---------------------------------------------------------------------------
# Allowlist — every line the tool legitimately says to every CDE.
#
# Each entry is (line, justification). A line reaches this list only by being
# one of:
#   HEADING   structural title or table-of-contents entry
#   LABEL     column header, row label or field name
#   PLACEHO   [CDE TO COMPLETE] / [NARRATIVE PLACEHOLDER] text, which is the
#             absence of a claim
#   SOURCED   a statutory or published constant that carries its citation
#             inline in the same line. EXTERNALLY RETRIEVABLE: a reader must be
#             able to go and read the thing the justification names.
#   HOUSE     a value this package sets itself, disclaimed on its face in the
#             rendered line. Split out of SOURCED in 1.2.1 R-2.
#   UNSOURCED research established that NO primary source exists, and the
#             rendered line says so. Also split out of SOURCED in 1.2.1 R-2.
#   TOOL      a statement about this tool, not about the CDE
#   FORMAT    markdown/table scaffolding with no propositional content
# ---------------------------------------------------------------------------

ALLOWLIST_PATH = os.path.join(os.path.dirname(__file__), "invariant_allowlist.txt")


def _load_allowlist() -> dict:
    """Parse the committed allowlist: 'CATEGORY | justification | line'."""
    entries = {}
    with open(ALLOWLIST_PATH, encoding="utf-8") as fh:
        for raw in fh:
            if not raw.strip() or raw.startswith("#"):
                continue
            category, justification, line = raw.rstrip("\n").split(" | ", 2)
            entries[line] = (category, justification)
    return entries


# ---------------------------------------------------------------------------
# Render + extract
# ---------------------------------------------------------------------------

def _extract(fmt: str, path: str) -> str:
    if fmt == "markdown":
        with open(path, encoding="utf-8") as fh:
            return fh.read()
    if fmt == "word":
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        for section in doc.sections:
            parts.extend(p.text for p in section.header.paragraphs)
            parts.extend(p.text for p in section.footer.paragraphs)
        return "\n".join(parts)
    if fmt == "excel":
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                parts.extend(str(v) for v in row if v is not None)
        return "\n".join(parts)
    if fmt == "pdf":
        from pypdf import PdfReader
        return "\n".join((page.extract_text() or "") for page in PdfReader(path).pages)
    raise AssertionError(f"unknown format {fmt}")


def _is_prose(line: str) -> bool:
    """A line worth adjudicating: long enough and wordy enough to assert something.

    Short cells, bare numbers and single words carry no proposition on their
    own; they are covered by the table-label allowlist entries they sit under.

    THE THRESHOLD IS FOUR WORDS, NOT FIVE, AND THAT IS A BUG FIX. At five, this
    rendered line was invisible to BOTH gates:

        **Board Meeting Frequency:** Quarterly

    — thirty-eight characters, four words, and a fabricated answer to a
    governance question the CDE never answered, printed in the section where
    management capacity is scored. It survived six rounds behind a word count.

    Cost of the change, measured: the invariant set grows 172 -> 207, and 33 of
    the 35 additions are bare headings and column labels. That is a cheap price
    for closing the last of the four gaps this gate had.

    Three words was also measured (223 lines) and rejected: the extra 16 are
    fragments like "Native Area Projects" that carry no proposition even in
    principle, and they dilute the list a reviewer has to read.
    """
    if len(line) < 25:
        return False
    if line.startswith(("#", "@@SHEET")):
        return False
    return len(line.split()) >= 4


# ---------------------------------------------------------------------------
# The mask: compare what a line SAYS, not what it says it ABOUT.
#
# Byte-identity was the original test, and it has a hole big enough to have let
# four blockers through at once. This renders for every CDE that ever runs the
# tool:
#
#     f"All {total_projects} projects have completed preliminary underwriting
#      review."
#
# It is invariant in substance and variable in bytes, so intersecting raw lines
# never saw it. Interpolating ANY user value into a fabricated sentence hid the
# sentence. Measured on this fixture set: 43 of the 155 things the tool says
# about every CDE escaped that way.
#
# So mask the interpolated values first, then intersect. A sentence that says
# the same thing with different nouns is the same sentence.
#
# THE VOCABULARY IS DERIVED FROM THE SCENARIOS, NOT GUESSED. Every string below
# is a value these fixtures actually vary — CDE names, project names, cities,
# addresses, sectors — read straight off SCENARIOS. A hand-written regex would
# be simultaneously too broad (eating real words) and too narrow (missing a
# field somebody adds later); this cannot drift from the fixtures because it is
# generated from them.
#
# TWO NARROWINGS, both empirical, both measured rather than assumed:
#
#   1. Matching is WORD-BOUNDED. Substring matching spliced tokens into real
#      words — "QALICB" contains Alabama's "AL", "CONFIDENTIAL" contains "ID"
#      and "AL", and "<SECTOR>" contains "OR", so the mask re-masked its own
#      output. Word boundaries fix all of it.
#   2. Bare two-letter STATE CODES are NOT masked; full state names are. "ID"
#      is a whole word in the column header "Project ID" and "OR"/"IN" appear
#      as English. Masking them moves the count by exactly one line (164 vs
#      163) and that line is a state-summary table row carrying no proposition.
#      Not worth corrupting every header row for.
#
# Effect, measured: invariant prose 116 -> 172 (163 before the seven blocker
# fixes changed what renders). The growth IS the point — each new line is
# something the tool asserts about every CDE that no one had ever had to
# justify. If a future change pushes this past ~200 the mask has become too
# aggressive; fix the mask rather than bulk-approving the allowlist.
# ---------------------------------------------------------------------------

_NUMBER_RE = re.compile(r"\d[\d,]*(?:\.\d+)?")


def _spellings(raw: str) -> set:
    """Every casing a fixture value can reach the page in.

    Renderers title-case sectors ("clean_energy" -> "Clean Energy"), upper-case
    them in headers, and pass them through raw in CSV-ish columns.
    """
    out = {raw}
    if "_" in raw:
        spaced = raw.replace("_", " ")
        out |= {spaced, spaced.title(), spaced.upper(), spaced.capitalize()}
    out |= {raw.title(), raw.upper(), raw.capitalize()}
    return {v for v in out if len(v) >= 3}


def _build_mask():
    """Compile a single-pass masker over every value SCENARIOS varies."""
    buckets = {}

    def add(token, value):
        if value:
            buckets.setdefault(token, set()).add(str(value))

    for _sid, (cde, pipeline, _requested) in SCENARIOS.items():
        for v in _spellings(cde.name):
            add("<CDE>", v)
        add("<CDEID>", cde.cde_id)
        add("<DATE>", cde.certification_date)
        add("<MISSION>", cde.mission)
        add("<MISSION>", cde.mission[:200])
        for market in cde.target_markets:
            add("<STATE>", market)          # full names only — see narrowing 2
        for p in pipeline:
            add("<QALICB>", p.qalicb_name)
            for v in _spellings(p.project_name):
                add("<PROJECT>", v)
            add("<PID>", p.project_id)
            add("<ADDRESS>", p.address)
            for v in _spellings(p.city):
                add("<CITY>", v)
            for v in _spellings(p.sector):
                add("<SECTOR>", v)
            for v in _spellings(p.project_type):
                add("<PTYPE>", v)

    pairs = sorted(
        ((v, tok) for tok, vals in buckets.items() for v in vals),
        key=lambda kv: -len(kv[0]),          # longest first: QALICB name before its ID
    )
    assert pairs, (
        "the mask vocabulary is EMPTY — it is built from SCENARIOS, so an empty "
        "vocabulary means the fixtures stopped exposing the fields it reads and "
        "every interpolated line would silently become variant again"
    )
    lookup = {value: token for value, token in pairs}
    # One alternation, one pass: a token can never be re-masked by a later rule.
    pattern = re.compile(
        r"(?<![A-Za-z0-9])(" + "|".join(re.escape(v) for v, _ in pairs) + r")(?![A-Za-z0-9])"
    )
    return pattern, lookup


_MASK_PATTERN, _MASK_LOOKUP = _build_mask()


def _mask(line: str) -> str:
    """Replace interpolated values with typed placeholders, digits with N."""
    masked = _MASK_PATTERN.sub(lambda m: _MASK_LOOKUP[m.group(1)], line)
    masked = _NUMBER_RE.sub("N", masked)
    return re.sub(r"\s+", " ", masked).strip()


@pytest.fixture(scope="module")
def rendered_lines(tmp_path_factory) -> dict:
    """{scenario_id: set(masked prose lines)} across all four formats."""
    per_scenario = {}
    for sid, (cde, pipeline, requested) in SCENARIOS.items():
        out = tmp_path_factory.mktemp(f"inv{sid}")
        app = Application(cde=cde, requested_allocation=requested)
        app.add_pipeline(pipeline)
        paths = app.generate(str(out), formats=list(FORMATS))

        assert set(paths) == set(FORMATS), (
            f"scenario {sid} rendered {sorted(paths)}, expected all of "
            f"{sorted(FORMATS)} — a format that silently does not render is "
            "a format this gate is not checking"
        )

        lines = set()
        for fmt, path in paths.items():
            assert os.path.exists(path), f"scenario {sid}: {fmt} not written"
            text = _extract(fmt, path)
            assert text.strip(), (
                f"scenario {sid}: {fmt} extracted as empty text — the gate "
                "would then find everything invariant, or nothing"
            )
            lines |= {_mask(ln) for ln in text.splitlines() if ln.strip()}
        per_scenario[sid] = {ln for ln in lines if ln}
    return per_scenario


# ---------------------------------------------------------------------------
# Fail-closed structural checks
# ---------------------------------------------------------------------------

def test_scenarios_are_disjoint():
    """No two scenarios may share a CDE, state, sector or dollar amount."""
    assert len(SCENARIOS) >= MIN_SCENARIOS, (
        f"only {len(SCENARIOS)} scenarios; the method needs at least "
        f"{MIN_SCENARIOS} for 'identical across all of them' to mean anything"
    )
    names, ids, states, sectors, amounts = set(), set(), set(), set(), set()
    for sid, (cde, pipeline, requested) in SCENARIOS.items():
        for bucket, value in ((names, cde.name), (ids, cde.cde_id),
                              (amounts, requested)):
            assert value not in bucket, f"scenario {sid} reuses {value!r}"
            bucket.add(value)
        for p in pipeline:
            assert p.state not in states or p.state in {q.state for q in pipeline}, ""
        scenario_states = {p.state for p in pipeline}
        scenario_sectors = {p.sector for p in pipeline}
        assert not (scenario_states & states), (
            f"scenario {sid} shares a state with an earlier scenario"
        )
        assert not (scenario_sectors & sectors), (
            f"scenario {sid} shares a sector with an earlier scenario"
        )
        states |= scenario_states
        sectors |= scenario_sectors
        for p in pipeline:
            assert p.qei_request not in amounts, (
                f"scenario {sid} reuses QEI {p.qei_request}"
            )


def test_allowlist_is_loadable_and_justified():
    """Every allowlist entry must carry a category and a justification."""
    allow = _load_allowlist()
    assert allow, "the allowlist is empty — every invariant line would fail"
    # DERIVED is new with the mask: a fixed sentence frame whose every figure
    # comes from the CDE's own inputs. Byte-identity never surfaced these,
    # because the interpolated value made every scenario differ.
    valid = {"HEADING", "LABEL", "PLACEHO", "SOURCED", "TOOL", "FORMAT",
             "DERIVED", "HOUSE", "UNSOURCED"}
    for line, (category, justification) in allow.items():
        assert category in valid, f"unknown category {category!r} for {line[:60]!r}"
        assert len(justification.strip()) >= 8, (
            f"allowlist entry has no real justification: {line[:60]!r}"
        )


def test_sourced_means_externally_retrievable():
    """SOURCED is the category a future reviewer trusts. It must stay earned.

    1.2.1 R-2: thirteen entries were filed SOURCED whose "source" was an
    internal defect ID ("1.2.1 B-3"), and five more whose source was an
    assertion that NO source exists ("No primary source defines 'net
    subsidy'"). Not one of those justifications is a false statement — every
    one is accurate about the line it justifies. The problem is the shelf they
    were put on: SOURCED is what a reviewer skims past on the assumption that
    somebody can go and read the cited thing, and an internal ticket number is
    not something anybody outside this repository can read.

    Three categories now carry the distinction, so the shelf means what it
    says:
        SOURCED    go and read it: a document, a year, a section
        HOUSE      this package set it, and the rendered line admits that
        UNSOURCED  nobody published it, and here is the research establishing
                   that nobody did

    UNSOURCED is the fourth category this file did not have. It is not a
    softer SOURCED: it is a stronger claim, because it asserts a negative that
    somebody had to go and check. "The CY 2024-2025 Application does not use
    the phrase 'net subsidy' at all (zero occurrences across 142 pages)" is a
    finding, and filing it as a citation hid the work that produced it.
    """
    allow = _load_allowlist()
    internal = []
    for line, (category, justification) in allow.items():
        if category != "SOURCED":
            continue
        j = justification.lower()
        if re.search(r"\b1\.[0-9]+\.[0-9]+ [a-z]-[0-9]", j):
            internal.append(f"internal defect ID: {justification[:70]}")
        elif "no primary source" in j or "publishes no" in j:
            internal.append(f"asserts no source exists: {justification[:70]}")
        elif "this model's" in j or "this tool's own" in j or "market assumption" in j:
            internal.append(f"a house value: {justification[:70]}")
    assert not internal, (
        f"{len(internal)} entr(ies) are filed SOURCED whose source is not "
        "externally retrievable. Use HOUSE (this package set it and the line "
        "says so) or UNSOURCED (no primary source exists and here is the "
        "research):\n  " + "\n  ".join(internal)
    )


def test_house_and_unsourced_are_both_populated():
    """Fail closed: an empty category is a category nobody is using.

    If HOUSE or UNSOURCED ever empties out, either the entries moved back to
    SOURCED — which is the defect R-2 fixed — or the disclosures they describe
    stopped rendering.
    """
    allow = _load_allowlist()
    counts = {}
    for _line, (category, _j) in allow.items():
        counts[category] = counts.get(category, 0) + 1
    for category in ("HOUSE", "UNSOURCED"):
        assert counts.get(category, 0) > 0, (
            f"no entry is filed {category}. This package renders both a credit "
            "price it invented and a row whose term of art no primary source "
            "defines; if neither is on this list, either the categories were "
            "collapsed back into SOURCED or those lines stopped rendering."
        )


def test_mask_actually_masks():
    """Fail closed: a mask that silently stops masking restores the old hole.

    This is the failure mode this portfolio keeps producing — a guard that
    passes because it checks nothing. If the vocabulary stopped covering CDE
    names, or the number regex stopped matching, every interpolated line would
    become variant again and four blockers' worth of surface would go dark
    without a single test turning red.
    """
    cde, pipeline, _ = SCENARIOS[1]
    project = next(iter(pipeline))

    # The exact shape that hid for six rounds.
    sentence = f"All {len(list(pipeline))} projects have completed review."
    assert _mask(sentence) == "All N projects have completed review.", _mask(sentence)

    for value, token in ((cde.name, "<CDE>"), (cde.cde_id, "<CDEID>"),
                         (project.city, "<CITY>"), (project.project_name, "<PROJECT>")):
        masked = _mask(f"prefix {value} suffix")
        assert token in masked, f"{value!r} was not masked to {token}: {masked!r}"
        assert value not in masked, f"{value!r} survived masking: {masked!r}"

    # And it must NOT eat ordinary words. "QALICB" contains Alabama's "AL";
    # "CONFIDENTIAL" contains "ID" and "AL". Substring matching spliced tokens
    # into both before word boundaries were added.
    for intact in ("QALICB Name", "CONFIDENTIAL", "TOTALS", "SUMMARY", "Project ID"):
        assert _mask(intact) == intact, f"mask corrupted {intact!r} -> {_mask(intact)!r}"


def test_mask_does_not_collapse_everything():
    """The mask must discriminate: two genuinely different claims stay different."""
    a = _mask("This pipeline is concentrated in deep distress.")
    b = _mask("This pipeline is concentrated in severe distress.")
    assert a != b, "the mask erased a real difference between two claims"


def test_extraction_is_not_empty(rendered_lines):
    """Fail closed: an empty extraction would make the gate vacuous."""
    assert len(rendered_lines) >= MIN_SCENARIOS
    for sid, lines in rendered_lines.items():
        assert len(lines) > 100, (
            f"scenario {sid} yielded only {len(lines)} lines across four "
            "formats — extraction is broken, not the output"
        )




# ---------------------------------------------------------------------------
# N-WRAP: a piece of a ruled sentence is ruled.
#
# ADDED IN 1.3.0 FIX-2 B1, AND THE MEASUREMENT THAT FORCED IT. Fixing the PDF's
# key/value tables so their cells WRAP turned four long unwrapped cells into 91
# short visual lines. All 91 are contiguous fragments of lines already on the
# allowlist — verified, zero orphans — and the whole-cell parents still render
# (Word and Markdown carry them unwrapped), so nothing became dead.
#
# The alternative was 91 new allowlist entries. It was rejected because it
# measures the renderer instead of the claim: the same Q25 basis note is ONE
# line in Word, ONE in Markdown and nine in the PDF, and requiring nine
# separate rulings for one sentence would put lines like
#
#     Areas; Federal Medically Underserved Areas;
#
# into the file whose own header says "READING THIS FILE IS THE HIGHEST-VALUE
# REVIEW ON THIS PACKAGE". A fragment that carries no proposition on its own
# cannot be reviewed on its own, and 91 of them would bury the 231 that can.
#
# The file was already conceding the point by hand: four entries in it are
# justified as "Wrapped continuation of ..." and one says outright that N2
# "misses it only because the PDF extractor wraps the line away from its
# opening bracket". This is that concession, derived instead of typed.
#
# WHAT IT CANNOT DO. The fragment must be a CONTIGUOUS substring of a line a
# human already ruled, and strictly shorter than it. A sentence the tool starts
# saying that is not part of one it already says is not a fragment of anything
# and is still reported. test_the_wrap_narrowing_cannot_absorb_a_new_claim is
# the proof, and the gate prints how many lines the narrowing absorbed so the
# number cannot grow unnoticed.
# ---------------------------------------------------------------------------

#: Below this, a "fragment" is short enough to be an accidental substring of an
#: unrelated ruled line rather than a wrap of it. _is_prose already floors
#: candidates at 25 characters; this is belt and braces at the narrowing.
MIN_FRAGMENT_CHARS = 20


def wrapped_fragments(lines, ruled) -> dict:
    """Return ``{line: parent}`` for each line that is a piece of a ruled line.

    Example::

        wrapped_fragments({"a piece of the whole"}, {"a piece of the whole thing"})
    """
    absorbed = {}
    for line in lines:
        if len(line) < MIN_FRAGMENT_CHARS:
            continue
        for parent in ruled:
            if len(parent) > len(line) and line in parent:
                absorbed[line] = parent
                break
    return absorbed


# ---------------------------------------------------------------------------
# The gate
# ---------------------------------------------------------------------------

def test_no_unallowed_invariant_assertion(rendered_lines):
    """Any prose line identical across all disjoint scenarios must be allowlisted."""
    invariant = set.intersection(*rendered_lines.values())
    assert invariant, (
        "NO line was invariant across the scenarios. Headings alone should be. "
        "Extraction or rendering is broken — this gate must not pass vacuously."
    )

    candidates = {ln for ln in invariant if _is_prose(ln)}
    allow = _load_allowlist()
    unallowed = set(candidates) - set(allow)
    absorbed = wrapped_fragments(unallowed, set(allow))
    print(f"  N-WRAP absorbed {len(absorbed)} wrapped fragment(s) of "
          f"{len(set(absorbed.values()))} ruled line(s)")
    unallowed = sorted(unallowed - set(absorbed))

    assert not unallowed, (
        f"{len(unallowed)} line(s) are byte-identical across "
        f"{len(rendered_lines)} disjoint CDEs and pipelines, and are not on "
        f"the reviewed allowlist ({len(allow)} entries).\n\n"
        "A line that does not vary with the input was not derived from the "
        "input. Either it is a heading/label/placeholder/sourced constant — "
        f"add it to {os.path.basename(ALLOWLIST_PATH)} with a category and a "
        "justification — or the tool is asserting it about every CDE that "
        "ever runs it, which is the defect this release exists to remove.\n\n"
        + "\n".join(f"  {ln}" for ln in unallowed)
    )


def test_allowlist_has_no_dead_entries(rendered_lines):
    """An allowlist entry that no longer renders is stale and must be removed.

    Without this the allowlist only ever grows, and a reviewer cannot tell
    which entries still describe the shipped output.
    """
    invariant = {ln for ln in set.intersection(*rendered_lines.values()) if _is_prose(ln)}
    dead = sorted(set(_load_allowlist()) - invariant)
    assert not dead, (
        f"{len(dead)} allowlist entr(ies) no longer appear as invariant lines. "
        "Remove them so the list stays a description of what actually "
        "renders:\n" + "\n".join(f"  {ln}" for ln in dead)
    )


def test_the_wrap_narrowing_cannot_absorb_a_new_claim(rendered_lines):
    """N-WRAP must swallow wraps and nothing else.

    A gate that has never been seen to refuse is a gate nobody has evidence
    about. Three cases, all against the real allowlist: a genuine fragment is
    absorbed, a sentence that is not a piece of any ruled line is not, and a
    fragment shorter than :data:`MIN_FRAGMENT_CHARS` is not — because a short
    string can be an accidental substring of an unrelated ruling.
    """
    allow = set(_load_allowlist())
    parent = max(allow, key=len)

    genuine = parent[: len(parent) - 10]
    assert wrapped_fragments({genuine}, allow) == {genuine: parent}, (
        "N-WRAP did not absorb a contiguous piece of a ruled line — the wraps "
        "it exists for would each need a hand-written entry"
    )

    invented = "This tool ranks your pipeline in the top decile of applicants."
    assert wrapped_fragments({invented}, allow) == {}, (
        "N-WRAP absorbed a sentence that is not a fragment of any ruled line. "
        "It would then hide exactly the class this gate exists to catch."
    )

    assert wrapped_fragments({parent[:MIN_FRAGMENT_CHARS - 1]}, allow) == {}, (
        "N-WRAP absorbed a fragment below the length floor, where a substring "
        "match is as likely to be a coincidence as a wrap"
    )
