"""THE CONSTANT GATE: every published constant, asserted as the string it prints.

WHAT THIS CLOSES

tests/test_invariant_output.py is the fabrication gate and it works, but it
collapses every digit run to "N" before intersecting::

    _NUMBER_RE = re.compile(r"\\d[\\d,]*(?:\\.\\d+)?")
    masked = _NUMBER_RE.sub("N", masked)

THE MASK IS CORRECT AND MUST NOT BE WEAKENED. Without it, every dollar figure
varies per scenario, nothing is invariant, and the gate cannot run at all. But
digit masking necessarily erases every constant the document prints, so that
gate is structurally blind to all of them. Measured on the 1.2.0 tree by
mutating the rendered 85%/20% commitment to 55%/5%::

    [baseline 85%/20%]  invariant=207  UNALLOWED=0  DEAD=0
    [mutated  55%/ 5%]  invariant=207  UNALLOWED=0  DEAD=0

Identical. The round-nine swapped distress thresholds mask to the same string
as the correct legend. So do IRC §45D -> §42D, the 7-year compliance period,
$0.83, the ACS vintage, the impact bands and the readiness weights. The
attribution gate normalises digits the same way and its own header concedes
that citations are "recorded, not verified".

So: a separate gate, over values instead of clauses.

THE METHOD

tests/pinned_constants.txt lists, for each published constant that reaches a
rendered surface, the exact literal string that constant must print, the
surfaces it must print on, and its source. This module renders the surfaces and
asserts each string is present.

FOUR RULES, each of which this package has previously violated somewhere:

  1. ASSERT AGAINST THE RENDERED ARTIFACT. Never against the constant. A test
     that reads SEVERE_DISTRESS_MIN_PCT and compares it to itself cannot fail —
     that is the shape of the version guard that read installed metadata, the
     ast.parse().body scan, the ``|| true`` grep swallow and the test doubles
     that defined a dropped dependency field. Every expected string in the pin
     file is a literal, typed by hand, and nothing in this module imports the
     constants it pins.

  2. FAIL CLOSED. An empty pin set errors. A pin whose string no longer renders
     errors (a stale pin is a pin that stopped guarding). A pin with no source
     errors. A HOUSE pin whose disclosure does not co-render errors.

  3. EVERY ENTRY CARRIES A SOURCE, and where one cannot be given the entry says
     HOUSE and the RENDERED TEXT must say so too. That is the attribution
     allowlist's policy, applied to values.

  4. THE LIST IS DERIVED, NOT INHERITED. test_every_consumed_constant_is_pinned
     sweeps nmtcapp/data/{schema,benchmark_thresholds}.py for every module-level
     constant, finds the ones any module outside nmtcapp/data/ references, and
     requires each to be either pinned or explicitly waived with a reason. A
     constant added tomorrow fails this test until somebody adjudicates it.

WHAT A PIN CANNOT DO. A pin proves the string renders. It does not prove the
constant produced it. Where a renderer duplicated a constant as a display
literal, the pin would have passed over a mutated constant — so 1.2.1 removed
the duplications it found rather than pinning around them: win_probability's
"40-point minimum"/"85-point minimum" and excel_builder's weight_map, both of
which printed one value while the package gated on another. If you add a pin,
check the call site interpolates.

A PIN MUST ALSO BE ABLE TO FAIL ON A REARRANGEMENT, NOT ONLY ON A DELETION.
Pinning the bare string "Deep Distress" proves the phrase is somewhere in the
document. It does not prove it is on the right ROW: swap the two values in
renderers/styles.DISTRESS_DISPLAY and both phrases still appear, on each
other's projects, and a "Deep Distress" pin passes. Every label pin below is
therefore anchored to the fixture project whose distress_level produced it, so
the swap moves the label away from the anchor and the pin fails. Verified by
executing the swap: the majority of the distress-label pins go red. The number
of such pins is DERIVED by _sweep_census()['distress_label_pins'] and is not
stated here — see test_the_sweep_states_no_count_it_has_not_derived.

WHAT WAS NEVER A CANDIDATE FOR PINNING, AND IS NOW (1.2.1 S-2)

The gate's scope through 1.2.1-rc was nmtcapp/data/{schema,benchmark_thresholds}
only. That is where the NUMBERS live, and the gate was built for numbers. But
the thing a CDE reads is not only a number — it is also the WORD a number is
printed under, and the mapping from an internal key to that word lives in a
renderer, not in nmtcapp/data/. Three mutations were run against the 1.2.1
tree and all three passed 937 green:

    distress_table._ELIGIBILITY_SOURCE  2016-2020 ACS -> 2011-2015 ACS
    styles.DISTRESS_DISPLAY             swap the deep and severe LABELS
    recommendations.py                  hardcoded federal 85% -> 55%

The middle one is the reason the scope changed. It leaves Section B's narrative
share correct (it is computed from the distress_level KEY) while Appendices A,
B and D relabel every project, so a reviewer summing the attachment's Deep rows
gets a different number from the one the narrative claims. Reproduced on a
five-project fixture: narrative 19.4%, attachment 47.2%, one document.

It got through because a LABEL DICT had never been a candidate for pinning —
not because anybody skipped it. So the sweep no longer asks "is this constant
in nmtcapp/data/". It asks the artifact: DOES ANY STRING THIS CONSTANT HOLDS
APPEAR IN THE RENDERED DOCUMENT? A constant whose strings appear must be
adjudicated. A constant whose strings do not is adjudicated BY THE FIXTURE, on
every run, and needs no hand-written waiver — which matters, because a
hand-written "this does not render" is exactly the claim that goes stale
silently. See test_every_rendering_constant_is_pinned_or_waived.
"""
from __future__ import annotations

import ast
import os
import re
import subprocess
import tempfile

import pytest

from nmtcapp.core.application import Application
from nmtcapp.core.cde import CDEProfile
from nmtcapp.core.pipeline import Pipeline, PipelineProject

PIN_PATH = os.path.join(os.path.dirname(__file__), "pinned_constants.txt")

DOCUMENT_SURFACES = ("markdown", "word", "excel", "pdf")
# recommendations: RecommendationSet.summary(), reached through the public
# app.recommendations() and printed by the Streamlit scorer. It is a SEVENTH
# rendered surface and no gate looked at it, which is how a hardcoded federal
# 85% survived being changed to 55% with nothing red (1.2.1 L-3).
#
# top_tier_*: the same two text surfaces rendered from a SECOND fixture, one
# that actually reaches Top Tier. HOUSE_TOP_TIER_AGGREGATE_MIN and
# HOUSE_TOP_TIER_SECTION_MIN were waived on the grounds that they "render only when a
# fixture reaches Top Tier", which was an admission that the fixture was too
# weak rather than a statement that the constants reach no surface. They gate
# the tier LABEL, the label renders, and a fixture that sits just above the bar
# makes both constants bite: raise either and the label changes.
TEXT_SURFACES = (
    "cli_summary", "win_score", "recommendations",
    "top_tier_win_score", "top_tier_recommendations",
)
# NOT text. The number FORMAT applied to a cell is something the reader sees
# and no text extractor can: openpyxl's data_only read returns 6000000 whether
# the cell prints "$6,000,000", "600000000.0%" or "6000000". FMT_PCT landing on
# a dollar column is a units defect that changes every figure on the page and
# leaves the extracted text byte-identical, so every excel_builder.FMT_*
# constant was invisible to a gate built on extracted text. This surface is
# "<sheet> · <column header> · <number format>", one line per column, which
# makes a format pinnable TO A COLUMN rather than merely present in the file.
SHAPE_SURFACES = ("excel_cell_formats",)
ALL_SURFACES = DOCUMENT_SURFACES + TEXT_SURFACES + SHAPE_SURFACES

DATA_MODULES = {
    "schema": os.path.join("nmtcapp", "data", "schema.py"),
    "benchmark_thresholds": os.path.join("nmtcapp", "data", "benchmark_thresholds.py"),
}

# The whole shipped package, for the rendering sweep. DATA_MODULES above stays
# as the scope of the older CONSUMED sweep, which asks a different question
# (is a published value read by another module) and catches constants that
# reach a score without reaching a string.
PACKAGE_ROOT = "nmtcapp"

# Trees that CONSUME a constant. Through 1.2.1-rc this was nmtcapp/ alone, so a
# constant read only by a Streamlit page — a surface a CDE looks at — counted
# as consumed by nobody and needed no adjudication.
CONSUMER_ROOTS = ("nmtcapp", "streamlit_app")

# A string shorter than this matches by accident. "name", "state", "deep" and
# "none" all appear in any rendered application for reasons having nothing to
# do with the constant that happens to contain them; at a 4-character floor the
# sweep demanded rows for a colour map and a field-name list on exactly those
# collisions. The floor is 12 rather than a larger round number because the
# SHORTEST published label the sweep must not lose is "Deep Distress", which is
# 13 characters — a floor above it would drop the very constant this sweep was
# widened to catch. Measured across the whole package: floor 4 -> 30 rows
# demanded, floor 8 -> 23, floor 12 -> 19, floor 16 -> 17 and "Deep Distress"
# gone. 12 is the largest floor that still holds the case.
MIN_RENDERED_STRING = 12

# A HOUSE pin must render its own disclaimer on the same surface. These are the
# phrases that count; a HOUSE pin whose surface carries none of them fails.
_DISCLOSURE_PHRASES = (
    "not a cdfi fund parameter",
    "this tool's own",
    "unsourced house heuristic",
    "the cdfi fund publishes no",
    "not a cdfi fund threshold",
    "not a federal figure",
)


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

class Pin:
    __slots__ = ("constant", "surfaces", "source", "expected", "lineno")

    def __init__(self, constant, surfaces, source, expected, lineno):
        self.constant = constant
        self.surfaces = surfaces
        self.source = source
        self.expected = expected
        self.lineno = lineno

    @property
    def is_house(self) -> bool:
        return self.source.startswith("HOUSE:")

    def __repr__(self) -> str:
        return f"<Pin {self.constant} @{self.lineno}>"


def _load_registry():
    """Parse tests/pinned_constants.txt into (pins, waivers, known).

    THREE ROW TYPES, NOT TWO (FIX-2 G-4).

        CONSTANT | ...   a PIN: this string must render, here
        WAIVE | ...      nothing to guard: the constant reaches no surface
        KNOWN | ...      it DOES reach a surface, it is not guarded, and the
                         row says why not

    The third exists because twelve rows were filed as waivers under a section
    header reading "consumed constants that reach no rendered surface" while
    their own text opened "Renders as the '/ 10' denominator of a sub-score
    line". Both statements were on the same screen and they contradicted each
    other; the header is the half a reviewer skims. Splitting the row type
    makes the distinction structural instead of a convention, and
    test_no_waiver_describes_something_that_renders keeps it that way.

    KNOWN adjudicates for the sweep exactly as WAIVE does — the point is not to
    demand different behaviour from the gate, it is to stop a deferred defect
    being filed as an absence.
    """
    pins, waivers = [], {}
    known = {}
    with open(PIN_PATH, encoding="utf-8") as fh:
        for lineno, raw in enumerate(fh, start=1):
            line = raw.rstrip("\n")
            if not line.strip() or line.lstrip().startswith("#"):
                continue
            # maxsplit, so the LAST field may itself contain " | ". A markdown
            # table row is the only string that anchors a cell label to the row
            # it sits on ("| 39035103200 | Y | Deep Distress |"), and without
            # this the distress-label pins could not be expressed on the
            # markdown surface at all. Missing fields still error: len < 3 or
            # len < 4 below.
            for prefix, bucket in (("WAIVE | ", waivers), ("KNOWN | ", known)):
                if line.lstrip().startswith(prefix):
                    parts = line.split(" | ", 2)
                    assert len(parts) == 3, (
                        f"{PIN_PATH}:{lineno}: a {prefix.strip(' |')} row is "
                        f"'{prefix.strip(' |')} | CONSTANT | reason', got "
                        f"{len(parts)} fields"
                    )
                    bucket[parts[1].strip()] = parts[2].strip()
                    break
            if line.lstrip().startswith(("WAIVE | ", "KNOWN | ")):
                continue
            parts = line.split(" | ", 3)
            assert len(parts) == 4, (
                f"{PIN_PATH}:{lineno}: a pin row is "
                f"'CONSTANT | SURFACES | SOURCE | expected', got {len(parts)} fields"
            )
            constant, surfaces, source, expected = (p.strip() for p in parts)
            pins.append(Pin(
                constant=constant,
                surfaces=tuple(s.strip() for s in surfaces.split(",")),
                source=source,
                expected=expected,
                lineno=lineno,
            ))
    return pins, waivers, known


PINS, WAIVERS, KNOWN = _load_registry()


# ---------------------------------------------------------------------------
# Fixture: one fully-verified application, rendered to every surface.
#
# Deliberately NOT the four disjoint scenarios from test_invariant_output. That
# gate needs disjointness; this one needs the non-degraded path, because half
# the constants below are suppressed when eligibility data is missing or a
# project is unverified — and a gate that silently ran on the degraded path
# would pass while pinning nothing.
# ---------------------------------------------------------------------------

def _fixture_application() -> Application:
    cde = CDEProfile(
        name="Pinned Constants Test CDE, LLC",
        cde_id="CDE-2019-0777",
        certification_date="2019-06-14",
        mission="Fixture profile for the constant-pinning gate.",
        target_markets=["Ohio", "Kentucky", "West Virginia"],
        prior_awards=[{"year": 2019, "amount": 40_000_000,
                       "deployment_status": "fully_deployed"}],
        contact={"name": "Pin Fixture", "email": "pins@example.org"},
        governance={"board_members": 9, "community_representatives": 4},
    )
    projects = []
    for i, (state, tract, level) in enumerate((
        ("OH", "39035103200", "deep"),
        ("KY", "21111010100", "severe"),
        ("WV", "54039000200", "lic"),
        ("OH", "39049003400", "deep"),
    )):
        p = PipelineProject(
            project_id=f"PIN-{i:02d}",
            project_name=f"Pin Fixture Project {i}",
            qalicb_name=f"Pin Fixture {i} QALICB LLC",
            address=f"{100 + i} Pin Street",
            city="Fixtureville",
            state=state,
            sector="healthcare",
            project_type="real_estate",
            total_project_cost=float(9_000_000 + i),
            qei_request=float(6_000_000 + i),
            qlici_amount=float(6_000_000 + i),
            expected_jobs_created=40 + i,
            expected_jobs_retained=10 + i,
            expected_units_built=None,
            expected_sq_ft=float(15_000 + i),
        )
        p.census_tract = tract
        p.is_nmtc_eligible = True
        p.distress_level = level
        p.is_native_area = i == 0
        p.is_high_migration_rural = False
        p.is_opportunity_zone = False
        p.is_us_territory = False
        p.is_persistent_poverty = True
        p.is_below_market_rate = True
        p.is_unrelated_entity = True
        p.geocode_success = True
        projects.append(p)
    pipeline = Pipeline(projects)
    pipeline.eligibility_data_status = "ok"

    app = Application(cde=cde, requested_allocation=24_000_000.0)
    app.add_pipeline(pipeline)
    return app


def _top_tier_application() -> Application:
    """A pipeline strong enough to be classified Top Tier.

    R-1: HOUSE_TOP_TIER_AGGREGATE_MIN and HOUSE_TOP_TIER_SECTION_MIN were waived because
    "they render only when a fixture actually reaches Top Tier". That is a
    statement about the FIXTURE, not about the constants — the constants decide
    which of three tier labels a CDE is told it is in, and the label is printed
    on the score block and on the recommendation set's overall assessment. A
    waiver saying "our fixture is too weak to reach the branch" waives the test,
    not the constant.

    So: a fixture that reaches the branch. Every project is Deep Distress, the
    CDE has three fully-deployed prior awards, and the flexible-product and
    unrelated-entity attributes are set, which is what the Business Strategy
    and Community Outcomes sub-scorers read. The pins that hang off this
    fixture assert the LABEL "Top Tier", so raising either constant above the
    score drops the label to "Highly Qualified" and the pin fails.
    """
    cde = CDEProfile(
        name="Top Tier Fixture CDE, LLC",
        cde_id="CDE-2014-0555",
        certification_date="2014-03-03",
        mission="Fixture profile that reaches the Top Tier gating branch.",
        target_markets=["Ohio", "Kentucky", "West Virginia", "Tennessee",
                        "Indiana", "Michigan", "Pennsylvania"],
        prior_awards=[
            {"year": 2019, "amount": 60_000_000, "deployment_status": "fully_deployed"},
            {"year": 2021, "amount": 65_000_000, "deployment_status": "fully_deployed"},
            {"year": 2023, "amount": 70_000_000, "deployment_status": "fully_deployed"},
        ],
        contact={"name": "Top Tier Fixture", "email": "toptier@example.org"},
        governance={"board_members": 11, "community_representatives": 7},
        extra={
            "years_in_operation": 12,
            "has_own_capital_at_risk": True,
            "products_below_market_pct": 1.0,
            "products_flexible_indicia_count": 7,
            "pipeline_pct_identified": 1.0,
            "track_record_pipeline_alignment_pct": 1.0,
            "track_record_deployment_pct": 1.0,
            "dbc_focus_years": 9,
            "dbc_dollar_volume_pct": 0.95,
            "unrelated_entities_pct": 1.0,
            "lic_board_representation_pct": 0.64,
            "has_community_engagement_track_record": True,
            "has_quantified_outcomes": True,
            "has_third_party_validation": True,
        },
    )
    tracts = [
        ("OH", "39035103200"), ("KY", "21111010100"), ("WV", "54039000200"),
        ("TN", "47157003100"), ("IN", "18097353100"), ("MI", "26163516900"),
        ("PA", "42101014000"), ("OH", "39049003400"),
    ]
    projects = []
    for i, (state, tract) in enumerate(tracts):
        p = PipelineProject(
            project_id=f"TOP-{i:02d}",
            project_name=f"Top Tier Fixture Project {i}",
            qalicb_name=f"Top Tier {i} QALICB LLC",
            address=f"{500 + i} Summit Street",
            city="Peakville",
            state=state,
            sector=("healthcare", "affordable_housing", "education",
                    "community_facility", "clean_energy")[i % 5],
            project_type="real_estate",
            total_project_cost=float(12_000_000 + i),
            qei_request=float(8_000_000 + i),
            qlici_amount=float(8_000_000 + i),
            expected_jobs_created=90 + i,
            expected_jobs_retained=30 + i,
            expected_units_built=40 + i,
            expected_sq_ft=float(30_000 + i),
        )
        p.census_tract = tract
        p.is_nmtc_eligible = True
        p.distress_level = "deep"
        p.is_native_area = i % 2 == 0
        p.is_high_migration_rural = i % 3 == 0
        p.is_opportunity_zone = True
        p.is_us_territory = False
        p.is_persistent_poverty = True
        p.is_below_market_rate = True
        p.is_unrelated_entity = True
        p.geocode_success = True
        projects.append(p)
    pipeline = Pipeline(projects)
    pipeline.eligibility_data_status = "ok"

    app = Application(cde=cde, requested_allocation=64_000_000.0)
    app.add_pipeline(pipeline)
    return app


def _extract(fmt: str, path: str) -> str:
    if fmt == "markdown":
        with open(path, encoding="utf-8") as fh:
            return fh.read()
    if fmt == "word":
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        for section in doc.sections:
            parts.extend(p.text for p in section.header.paragraphs)
            parts.extend(p.text for p in section.footer.paragraphs)
        return "\n".join(parts)
    if fmt == "excel":
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                parts.extend(str(v) for v in row if v is not None)
        return "\n".join(parts)
    if fmt == "pdf":
        from pypdf import PdfReader
        return "\n".join(_strip_pdf_chrome(page.extract_text() or "")
                         for page in PdfReader(path).pages)
    raise AssertionError(f"unknown format {fmt}")


#: The running footer ReportLab draws onto every page but the cover:
#: "<CDE name>  —  NMTC <round> Application  |  CONFIDENTIAL" and "Page <n>",
#: matched by shape rather than by the fixture's own strings so the pattern
#: cannot go stale against a renamed fixture.
_PDF_CHROME = re.compile(
    r"^(?:.*\|\s*CONFIDENTIAL\s*|Page \d+\s*)$", re.MULTILINE
)


def _strip_pdf_chrome(page_text: str) -> str:
    """Remove the running header/footer from one page's extracted text.

    A PIN IS A CLAIM ABOUT THE DOCUMENT, NOT ABOUT THE PAGE IT LANDED ON.
    1.3.0 FIX-2 B1 made the Section B key/value table wrap, and the Q25 basis
    note is long enough that its row now breaks across a page. pypdf then
    extracts the footer of page N BETWEEN the two halves of the sentence:

        ... at least TWO of items 6-12 (25% poverty /
        Great Lakes Regional Capital CDE, LLC — NMTC CY2025 Application |
        CONFIDENTIAL Page 6
        70% median family income / 1.25x unemployment; ...

    ``_normalise`` collapses whitespace so a pin survives a LINE wrap; it
    cannot survive a PAGE wrap, because what interrupts the string is not
    whitespace but chrome. Dropping the chrome — which is drawn onto the canvas
    and is not part of any pinned constant — restores the contiguity the pin
    asserts, and does so for every future pin that spans a page break rather
    than only for this one.

    Example::

        _strip_pdf_chrome("body text\nPage 6")   # -> 'body text'
    """
    return _PDF_CHROME.sub("", page_text).strip()


def _excel_cell_formats(path: str) -> str:
    """"<sheet> · <column header> · <number format>", one line per data column.

    The number format is the only thing about the Excel attachment that the
    reader sees and _extract cannot: openpyxl with data_only=True hands back
    6000000 regardless of whether the cell prints $6,000,000 or 600000000.0%.
    Building this surface makes the excel_builder.FMT_* constants
    pinnable to the COLUMN they format, so swapping FMT_CURRENCY for FMT_PCT on
    the QEI column fails a pin instead of shipping a page of percentages.

    The header is taken to be the last row above the first cell that carries a
    non-General format, which is how these sheets are laid out (a title block,
    then one header row, then data). A column whose data cells are all text
    contributes nothing, which is correct: there is no format to pin.
    """
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    lines = []
    for ws in wb.worksheets:
        rows = list(ws.iter_rows())
        for col_idx in range(ws.max_column):
            header, seen = None, []
            for row in rows:
                if col_idx >= len(row):
                    continue
                cell = row[col_idx]
                if cell.value is None:
                    continue
                fmt = cell.number_format
                if fmt in ("General", "@"):
                    header = str(cell.value)
                    continue
                if header and fmt not in seen:
                    seen.append(fmt)
                    lines.append(f"{ws.title} · {header} · {fmt}")
    return "\n".join(lines)


def _normalise(text: str) -> str:
    """Collapse whitespace so a pin survives PDF line wrapping and cell padding.

    Word puts the whole methodology block in one paragraph, ReportLab wraps it
    at the column width, and pypdf reintroduces the wrap as newlines. Comparing
    raw text would make every pin a test of the page width.
    """
    return re.sub(r"\s+", " ", text)


@pytest.fixture(scope="module")
def rendered() -> dict:
    """{surface: normalised text} for every surface in ALL_SURFACES."""
    import tempfile
    out = tempfile.mkdtemp(prefix="nmtcapp-pins-")
    app = _fixture_application()
    paths = app.generate(out, formats=list(DOCUMENT_SURFACES))

    assert set(paths) == set(DOCUMENT_SURFACES), (
        f"rendered {sorted(paths)}, expected all of {sorted(DOCUMENT_SURFACES)} — "
        "a format that silently does not render is a format this gate is not "
        "checking, which is exactly how the docs sample-output hook shipped two "
        "of four formats under a --strict build"
    )

    surfaces = {}
    for fmt, path in paths.items():
        text = _extract(fmt, path)
        assert text.strip(), f"{fmt} extracted as empty — the gate would pin nothing"
        surfaces[fmt] = _normalise(text)

    # NOT normalised: this surface is line-structured ("sheet · header · fmt")
    # and _shape_tokens reads it a line at a time. Collapsing the newlines would
    # turn it into one line and the token set into one token.
    surfaces["excel_cell_formats"] = _excel_cell_formats(paths["excel"])

    analysis = app.analyze()
    surfaces["cli_summary"] = _normalise(
        analysis.pipeline_result.summary() + "\n" + analysis.readiness_score.summary()
    )
    surfaces["win_score"] = _normalise(app.score_win_probability().summary())
    surfaces["recommendations"] = _normalise(app.recommendations().summary())

    top = _top_tier_application()
    top_score = top.score_win_probability()
    # FAIL CLOSED. Every pin on the two top_tier_* surfaces exists to make
    # HOUSE_TOP_TIER_AGGREGATE_MIN and HOUSE_TOP_TIER_SECTION_MIN bite. If the fixture stops
    # reaching the branch — because a sub-scorer changed, or because somebody
    # raised the constants and the fixture no longer clears them — those pins
    # would go quietly unreachable and the constants would be back to being
    # ungated. This assert says so out loud instead.
    assert top_score.tier == "Top Tier", (
        f"the Top Tier fixture scored {top_score.aggregate_base_score}/100 and was "
        f"classified {top_score.tier!r}, not 'Top Tier'. Either a sub-scorer "
        "changed, or HOUSE_TOP_TIER_AGGREGATE_MIN / HOUSE_TOP_TIER_SECTION_MIN moved. Both "
        "are release-blocking: the pins that gate those two constants hang off "
        "this branch, and a fixture that no longer reaches it un-gates them "
        "silently — which is the reason those constants were waived in the "
        "first place. Strengthen the fixture; do not weaken the assert."
    )
    surfaces["top_tier_win_score"] = _normalise(top_score.summary())
    surfaces["top_tier_recommendations"] = _normalise(top.recommendations().summary())

    for name in TEXT_SURFACES + SHAPE_SURFACES:
        assert surfaces[name].strip(), f"{name} rendered empty"
    assert set(surfaces) == set(ALL_SURFACES), (
        f"rendered {sorted(surfaces)}, ALL_SURFACES is {sorted(ALL_SURFACES)}. A "
        "surface named in ALL_SURFACES but never rendered would let every pin "
        "that names it pass by KeyError-free accident."
    )
    return surfaces


# ---------------------------------------------------------------------------
# Fail-closed structural checks
# ---------------------------------------------------------------------------

def test_registry_is_not_empty():
    """An empty pin set must ERROR, not pass.

    This is the property every gate in this package has failed at least once:
    the check runs, finds nothing to check, and reports success.
    """
    assert PINS, (
        f"{os.path.basename(PIN_PATH)} contains no pins. An empty registry "
        "would make this gate vacuous — every constant in the package could "
        "change and nothing would turn red."
    )
    assert len(PINS) >= 10, (
        f"only {len(PINS)} pins. The sweep in "
        "test_every_consumed_constant_is_pinned found more constants than that "
        "reaching a rendered surface; a registry this small means rows were "
        "deleted rather than adjudicated."
    )


def test_every_pin_names_a_surface_and_a_source():
    for pin in PINS:
        assert pin.surfaces, f"{pin} names no surface"
        for surface in pin.surfaces:
            assert surface in ALL_SURFACES, (
                f"{pin}: unknown surface {surface!r}; known: {ALL_SURFACES}"
            )
        assert pin.expected, f"{pin} pins an empty string"
        assert len(pin.source) >= 20, (
            f"{pin}: source is {pin.source!r}. Every pin carries a real source "
            "or says HOUSE and explains what the rendered text must admit."
        )
        if pin.is_house:
            assert len(pin.source) > len("HOUSE:") + 20, (
                f"{pin}: a HOUSE pin must state what the rendered text says"
            )
        else:
            assert re.search(r"\d", pin.source), (
                f"{pin}: a cited source must carry a digit — a year, section, "
                f"question or column. 'CDFI Fund guidance' is not a citation. "
                f"Got: {pin.source!r}"
            )


def test_waivers_carry_reasons():
    assert WAIVERS, "no waivers recorded; the sweep found consumed constants"
    for constant, reason in list(WAIVERS.items()) + list(KNOWN.items()):
        assert len(reason) >= 30, (
            f"row for {constant} has no real reason: {reason!r}"
        )


def test_known_rows_exist_and_are_a_separate_shelf():
    """FIX-2 G-4. The deferred-defect shelf must not empty back into WAIVE.

    If KNOWN empties, either the twelve sub-score denominators were fixed —
    in which case they get pins, not waivers — or somebody moved them back
    under a header that says they reach no rendered surface, which is the
    misfiling this row type exists to prevent.
    """
    assert KNOWN, (
        "no KNOWN rows. This package defers twelve rendering constants on "
        "purpose (eleven sub-score denominators paired with hardcoded caps, "
        "plus the QEI coverage band). If they were genuinely fixed they "
        "belong in the pin section; they do not belong under WAIVE."
    )
    overlap = sorted(set(KNOWN) & set(WAIVERS))
    assert not overlap, (
        f"filed as both waived and known-deferred: {overlap}. One of the two "
        "rows is wrong — a constant either reaches a surface or it does not."
    )
    pinned = {p.constant for p in PINS}
    assert not (set(KNOWN) & pinned), sorted(set(KNOWN) & pinned)


def test_every_row_sits_under_a_section_header_that_agrees_with_it():
    """A row's TYPE and the header above it must say the same thing (FIX-2 G-4).

    Twelve rows sat under

        "# --- Waivers: consumed constants that reach no rendered surface ---"

    while their own text opened

        "Renders as the '/ 10' denominator of a sub-score line in
         WinProbabilityScore.summary()."

    Both on one screen, contradicting each other, and the header is the half a
    reviewer skims. The REASONING in those rows is right — thirteen sub-scorers
    cap at hardcoded literals, so pinning a denominator freezes half of a pair
    whose other half is a typed number, and unifying them changes scoring,
    which a patch release is not the place for. What was wrong was the filing:
    a waiver claims there is nothing to guard, and these say there is something
    unguarded and here is why. That is a deferred defect.

    THE GUARD IS STRUCTURAL, NOT TEXTUAL. A regex over the reason prose cannot
    tell "the constant renders" from "something adjacent to the constant
    renders", and both sentences are legitimate — schema.GRADE_THRESHOLDS is
    correctly waived with the words "the grade LETTER renders". So this checks
    the one thing that is exactly checkable: a WAIVE row may not sit under the
    known-and-deferred header, and a KNOWN row may not sit under the waiver
    header. Moving the twelve back is then a test failure rather than a
    judgement call.
    """
    waiver_header = "reach no rendered surface"
    known_header = "Known and left alone"

    section = None
    misfiled = []
    with open(PIN_PATH, encoding="utf-8") as fh:
        for lineno, raw in enumerate(fh, start=1):
            line = raw.rstrip("\n")
            if line.startswith("# --- "):
                section = line
                continue
            if not line.strip() or line.lstrip().startswith("#"):
                continue
            constant = line.split(" | ")[1].strip() if " | " in line else line[:40]
            if line.startswith("KNOWN | ") and (
                section is None or known_header not in section
            ):
                misfiled.append(
                    f"  {PIN_PATH}:{lineno} KNOWN row {constant} under "
                    f"{section!r}"
                )
            if line.startswith("WAIVE | ") and section and known_header in section:
                misfiled.append(
                    f"  {PIN_PATH}:{lineno} WAIVE row {constant} under the "
                    "known-and-deferred header"
                )
            if line.startswith("KNOWN | ") and section and waiver_header in section:
                misfiled.append(
                    f"  {PIN_PATH}:{lineno} KNOWN row {constant} under the "
                    "waiver header, which says these reach no rendered surface"
                )

    assert section is not None, (
        "the registry has no '# --- ' section headers at all, so this check "
        "has nothing to compare a row against and would pass vacuously."
    )
    assert not misfiled, (
        f"{len(misfiled)} row(s) sit under a section header that contradicts "
        "their row type:\n" + "\n".join(misfiled)
        + "\n\nA WAIVE row says there is nothing to guard. A KNOWN row says "
        "there is something unguarded and why. Filing the second under a "
        "header announcing the first is how twelve deferred defects came to "
        "read as absences."
    )


# ---------------------------------------------------------------------------
# The gate
# ---------------------------------------------------------------------------

def test_pinned_constants_render_verbatim(rendered):
    """Every pinned constant must appear, as its literal string, on its surface.

    A failure here means one of two things and both are release-blocking:
    the constant changed and the pin was not re-adjudicated, or the constant
    changed and the DOCUMENT is now wrong.
    """
    failures = []
    for pin in PINS:
        for surface in pin.surfaces:
            if pin.expected not in rendered[surface]:
                failures.append(
                    f"  {PIN_PATH}:{pin.lineno}  {pin.constant}\n"
                    f"    surface : {surface}\n"
                    f"    expected: {pin.expected!r}\n"
                    f"    source  : {pin.source[:110]}"
                )
    assert not failures, (
        f"{len(failures)} pinned constant(s) do not render as pinned.\n\n"
        "Either a published constant changed value and the document is now "
        "wrong, or the rendered wording changed and the pin is stale. Both "
        "require a human to look at the source column before touching this "
        "file.\n\n" + "\n".join(failures)
    )


def test_house_pins_carry_their_disclaimer(rendered):
    """A HOUSE value must be disclosed as one, on the surface that prints it.

    Same policy as the attribution allowlist: an entry claiming HOUSE while the
    document reads as a federal figure is the defect, not the fix. 1.1.5
    printed this package's own 50%/75% bands under the labels "CDFI Fund
    Competitive Minimum" and "CDFI Fund Target".
    """
    failures = []
    for pin in PINS:
        if not pin.is_house:
            continue
        for surface in pin.surfaces:
            haystack = rendered[surface].lower()
            if not any(phrase in haystack for phrase in _DISCLOSURE_PHRASES):
                failures.append(
                    f"  {pin.constant} on {surface}: HOUSE value with no "
                    f"disclosure anywhere on that surface"
                )
    assert not failures, (
        "HOUSE-sourced constants render without saying so:\n" + "\n".join(failures)
    )


# ---------------------------------------------------------------------------
# Proximity: a disclosure 25 pages from its claim is not a disclosure
#
# test_house_pins_carry_their_disclaimer above flattens the whole document into
# one haystack and asks whether a disclosure phrase appears ANYWHERE in it. That
# is a presence check standing in for a proximity property, and it could not
# fail the thing it was written to catch: on 1.5.0 the readiness grade was
# printed on page 1 of the PDF and its disclosure was the last thing before the
# end of page 26 — 24,739 characters away, 98% of the document — and this gate
# was green. Measured on the same build: Word 16,657 characters (60% of the
# document), Markdown 25,179 (88%), Excel 48. Excel was the only surface that
# put the disclosure where the claim was, and it is the shape the other three
# now match.
#
# A GATE DEMONSTRATED ONLY GREEN IS NOT KNOWN TO WORK, so
# test_proximity_gate_fails_when_the_disclosure_is_far_away drives this same
# helper with a document whose disclosure sits at the far end and asserts it
# reports a failure. Without that, "the gate passes" and "the gate cannot fail"
# look identical from the outside — which is exactly how the flattened check
# above survived twenty-two releases.
# ---------------------------------------------------------------------------

# THE GATE MATCHED ON DISTANCE ALONE, AND COULD NOT FAIL ON MARKDOWN
# (1.5.1 audit, F3). The version shipped at 4bd26ab searched for ANY entry of
# _DISCLOSURE_PHRASES near a readiness claim. Those phrases are generic --
# "this tool's own" is the broadest of them -- so the gate was satisfied by a
# disclosure about a DIFFERENT claim that happened to sit nearby.
#
# PROVED BY REVERSION, not argued. Deleting readiness_weights_note() from
# markdown_builder's executive summary restores a 25,218-character gap between
# the readiness score and its weighting disclosure, and THE GATE STAYS GREEN:
# the Key Strengths heading two lines below reads "(this tool's own assessment
# against its own thresholds ...)", contains "this tool's own", and is measured
# at 69 characters. A disclosure about the STRENGTHS LIST was accepted as the
# disclosure for the READINESS SCORE.
#
# So the wrong-disclosure attack needed no synthetic document -- it was live on
# the shipped tree. And a gate that matches on distance alone is a presence
# check with extra steps: exactly the defect
# test_house_pins_carry_their_disclaimer had, one abstraction up. Replacing a
# flattened haystack with a windowed haystack does not make a check specific.
#
# THE PROPERTY IS NOW: a readiness claim must be near the disclosure that
# EXPLAINS IT -- text rendered by one of the three functions whose subject is
# the readiness score. The anchors are derived from those functions at import
# time, not typed here, so rewording a disclosure moves the gate with it and
# cannot silently un-gate it.

#: The functions whose output IS the readiness disclosure. Anything else on the
#: page may be a true disclosure of something, but it does not disclose this.
def _readiness_disclosure_anchors() -> tuple:
    from nmtcapp.renderers._methodology import (
        readiness_inline_qualifier,
        readiness_weights_note,
        readiness_weights_sheet_note,
    )
    return tuple(
        _normalise(fn()).lower()[:_ANCHOR_CHARS]
        for fn in (
            readiness_weights_note,
            readiness_inline_qualifier,
            readiness_weights_sheet_note,
        )
    )


#: How much of each disclosure must match. Long enough that no generic phrase
#: can satisfy it (asserted below), short enough to survive a renderer wrapping
#: the sentence across a table cell or a PDF line break.
_ANCHOR_CHARS = 40

#: Maximum characters between a readiness claim and the disclosure that
#: explains it.
#:
#: DERIVED, NOT CHOSEN (1.5.1 audit, F3). The value here was 1,200 -- "roughly
#: a rendered paragraph or two", which is a round number governing a gate, i.e.
#: a house constant that had not been declared one. It is replaced by the
#: length of the disclosure itself: a claim and its disclosure are in the same
#: passage when less text separates them than the disclosure contains. Past
#: that, the reader has crossed more unrelated material than explanation, which
#: is the thing being measured.
#:
#: It moves automatically if the disclosure is reworded, and it needs no
#: headroom argument because it is not a budget. Worst distance actually
#: rendered across all four surfaces at 1.5.1: 214 characters (PDF), against a
#: derived limit of 363 -- reported by the gate's own failure message, so a
#: regression toward the limit is visible rather than silent.
def _disclosure_proximity_limit() -> int:
    from nmtcapp.renderers._methodology import readiness_weights_note
    return len(_normalise(readiness_weights_note()))


#: Where the readiness grade is claimed, as it renders after normalisation.
_READINESS_CLAIM = re.compile(
    r"(readiness (?:grade|assessment|score)|application readiness)", re.I
)


def _disclosure_proximity_failures(surfaces: dict) -> list:
    """Report every readiness claim with no READINESS disclosure near it.

    Shared by the live gate and by its own red-proofs below, so the proofs
    exercise the real matcher rather than a paraphrase of it.
    """
    failures = []
    anchors = _readiness_disclosure_anchors()
    limit = _disclosure_proximity_limit()
    for surface, text in surfaces.items():
        low = text.lower()
        disclosure_at = [
            m.start()
            for phrase in anchors
            for m in re.finditer(re.escape(phrase), low)
        ]
        for claim in _READINESS_CLAIM.finditer(low):
            if not disclosure_at:
                failures.append(
                    f"  {surface} @{claim.start()}: readiness claim "
                    f"{claim.group(0)!r} — NO READINESS disclosure anywhere on "
                    "this surface (a house disclaimer about something else "
                    "does not count)"
                )
                break
            nearest = min(abs(d - claim.start()) for d in disclosure_at)
            if nearest > limit:
                failures.append(
                    f"  {surface} @{claim.start()}: readiness claim "
                    f"{claim.group(0)!r} — nearest READINESS disclosure is "
                    f"{nearest:,} characters away (derived limit {limit:,} = "
                    "the rendered length of readiness_weights_note())"
                )
    return failures


def test_the_readiness_anchors_cannot_be_satisfied_by_a_generic_disclaimer():
    """THE PROPERTY THAT MAKES THIS A MATCHED CHECK RATHER THAN A NEARBY ONE.

    If a readiness anchor were a substring of one of the generic
    _DISCLOSURE_PHRASES -- or short enough to appear inside an unrelated house
    disclaimer -- the gate would silently degrade back into the proximity-only
    check that could not fail on markdown. Asserted, not assumed.
    """
    anchors = _readiness_disclosure_anchors()
    assert len(anchors) == 3, anchors
    for anchor in anchors:
        assert len(anchor) == _ANCHOR_CHARS, (anchor, len(anchor))
        for generic in _DISCLOSURE_PHRASES:
            assert anchor not in generic, (
                f"readiness anchor {anchor!r} is contained in the generic "
                f"house phrase {generic!r}, so any house disclaimer anywhere "
                "would satisfy the readiness gate."
            )
        assert "readiness" in anchor or "this tool's own unsourced" in anchor, (
            f"anchor {anchor!r} does not identify the readiness score as its "
            "subject, so it cannot establish that the disclosure explains THIS "
            "claim."
        )

    # And the live Key Strengths disclaimer -- the exact string that satisfied
    # the old gate at 69 characters -- must NOT satisfy the new one.
    key_strengths = (
        "key strengths (this tool's own assessment against its own "
        "thresholds — not a cdfi fund evaluation):"
    )
    assert not any(a in key_strengths for a in anchors), (
        "the Key Strengths disclaimer still satisfies the readiness anchors. "
        "That string is a disclosure ABOUT THE STRENGTHS LIST, and accepting "
        "it for the readiness score is the defect F3 recorded."
    )


def test_proximity_gate_fails_on_markdown_when_only_the_readiness_note_is_removed():
    """THE MARKDOWN RED PROOF, on real rendered output, not a synthetic.

    F3's finding was that reverting ONLY the markdown disclosure fix left the
    gate green, because the Key Strengths disclaimer sat 69 characters from the
    claim and contained "this tool's own". This drives the real matcher with
    the real markdown surface, with only the readiness note removed, and
    asserts it now reports the true gap.

    A synthetic document cannot prove this: the whole point of F3 is that the
    live document already had a decoy in it. This uses the live one.
    """
    import re as _re

    from nmtcapp.renderers import markdown_builder as _mb

    src = _extract_markdown_exec_summary_source(_mb)
    assert "readiness_weights_note()" in src, (
        "markdown_builder's executive summary no longer calls "
        "readiness_weights_note(); this proof has lost its subject."
    )

    app = _fixture_application()
    out = tempfile.mkdtemp(prefix="nmtcapp-f3-")
    paths = app.generate(out, formats=["markdown"])
    with open(paths["markdown"], encoding="utf-8") as handle:
        text = handle.read()

    # Remove ONLY the executive-summary readiness note, exactly as reverting
    # the T3 markdown fix would. The methodology copy far below is untouched,
    # and so is the Key Strengths disclaimer that decoyed the old gate.
    from nmtcapp.renderers._methodology import readiness_weights_note

    note = readiness_weights_note()
    assert text.count(note) >= 2, (
        f"expected the readiness note at both the exec summary and the "
        f"methodology block, found {text.count(note)}"
    )
    reverted = text.replace(f"*{note}*\n\n", "", 1)
    assert reverted != text, "the exec-summary note was not removed"

    failures = _disclosure_proximity_failures({"markdown": _normalise(reverted)})
    assert failures, (
        "THE GATE DID NOT FAIL with the markdown readiness disclosure removed. "
        "It is matching on distance alone again: the Key Strengths disclaimer "
        "~69 characters below the claim is being accepted as the readiness "
        "score's disclosure."
    )
    assert any("characters away" in f for f in failures), failures
    # And the untouched document must pass, so this is not failing everything.
    assert not _disclosure_proximity_failures({"markdown": _normalise(text)})


def _extract_markdown_exec_summary_source(module) -> str:
    import inspect

    return inspect.getsource(module)


#: The Streamlit pages, as a proximity surface. Added by the 1.5.2 audit's F2.
#:
#: THE GATE THAT BLOCKS EXACTLY THIS COULD NOT SEE THE PAGE. Through 1.5.2
#: ``test_readiness_disclosure_is_adjacent_to_the_claim`` iterated the four
#: DOCUMENT_SURFACES only, and no other gate covered Streamlit. Meanwhile
#: ``1_Pipeline_Analyzer.py`` rendered the composite three ways -- the metric,
#: the "Overall readiness grade" fallback, and the six-component bar chart --
#: and carried no readiness disclosure at all. CLI and markdown went from a
#: one-component notice to a six-component deduction table in the same release;
#: Streamlit went from nothing to nothing, under a shipped refusal claim that
#: reads "a tool may decline to advise, it may not deduct silently".
#:
#: WHY AST AND NOT A RENDER. Driving Streamlit needs a script run and a running
#: server, and every other surface here is text. What a reader sees on this
#: page is the sequence of literals handed to the render calls, so that is what
#: is extracted, IN SOURCE ORDER, which is what makes a character distance
#: between a claim and its disclosure meaningful.
#:
#: WHAT IT CANNOT SEE, DECLARED. Text that arrives through a variable or a
#: function call -- ``readiness_inline_qualifier()`` among them -- is not a
#: literal in this file. Calls to the three _methodology disclosure functions
#: are therefore substituted with their VALUES, by name, so the anchors they
#: carry are measured where they actually render. Any other interpolated text
#: is invisible to this surface, and a disclosure smuggled in through one would
#: NOT satisfy this gate. That is the conservative direction.
_ST_RENDER_CALLS = frozenset({
    "markdown", "write", "caption", "info", "warning", "error", "success",
    "subheader", "header", "title", "text", "code", "latex", "metric",
    "expander", "badge", "toast", "popover", "metric_classification",
})


def _streamlit_disclosure_substitutions() -> dict:
    from nmtcapp.renderers._methodology import (
        readiness_inline_qualifier,
        readiness_weights_note,
        readiness_weights_sheet_note,
    )
    return {
        fn.__name__: fn()
        for fn in (readiness_inline_qualifier, readiness_weights_note,
                   readiness_weights_sheet_note)
    }


def _call_name(node) -> str:
    fn = node.func
    if isinstance(fn, ast.Attribute):
        return fn.attr
    if isinstance(fn, ast.Name):
        return fn.id
    return ""


def _streamlit_page_surface(path: str) -> str:
    """Every string a Streamlit page hands to a render call, in source order.

    ORDER COMES FROM AN ORDERED TRAVERSAL, NOT FROM ast POSITIONS, and that is
    not a style preference -- it is a portability bug this gate shipped once and
    CI caught on three interpreters.

    Before PEP 701 (Python 3.12) the expressions inside an f-string do not carry
    their own true line and column; they are parsed from a synthesised
    sub-source and the positions can land far from where the text renders.
    Sorting chunks by ``(lineno, col_offset)`` therefore produced a sane surface
    on 3.12 and a scrambled one on 3.9, 3.10 and 3.11 -- where an interpolated
    disclosure sorted thousands of characters away from the claim it sits
    beside, and the gate failed on a page that is correct.

    So: the render CALLS are ordered by position, which is a statement-level
    fact every version gets right, and within each call the arguments are walked
    depth-first through ``iter_child_nodes``, which yields ``JoinedStr.values``
    in textual order by construction.

    WHAT IT CANNOT SEE, DECLARED. Text arriving through a variable is not a
    literal here. Calls to the three _methodology disclosure functions are
    substituted with their VALUES so the anchors they carry are measured where
    they render; any OTHER interpolation is invisible, and a disclosure
    smuggled in through one would NOT satisfy this gate. That is the
    conservative direction.
    """
    subs = _streamlit_disclosure_substitutions()
    with open(path, encoding="utf-8") as fh:
        tree = ast.parse(fh.read())

    chunks: list = []

    def _emit_in_order(node):
        """Depth-first, field order — the order a reader meets the text."""
        if isinstance(node, ast.Constant) and isinstance(node.value, str):
            chunks.append(node.value)
            return
        if isinstance(node, ast.Call) and _call_name(node) in subs:
            chunks.append(subs[_call_name(node)])
            return
        for child in ast.iter_child_nodes(node):
            _emit_in_order(child)

    calls = [
        node for node in ast.walk(tree)
        if isinstance(node, ast.Call) and _call_name(node) in _ST_RENDER_CALLS
    ]
    # Statement-level positions, which every supported interpreter reports
    # correctly. Nested render calls (a column's .caption inside a with-block)
    # sort by their own position, which is where they render.
    calls.sort(key=lambda n: (n.lineno, n.col_offset))
    for node in calls:
        for arg in list(node.args) + [kw.value for kw in node.keywords]:
            _emit_in_order(arg)

    return _normalise(" ".join(chunks))


#: Computed here rather than through _repo_root(), which this module defines
#: two hundred lines further down.
_STREAMLIT_PAGES_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "streamlit_app", "pages",
)


def _streamlit_surfaces() -> dict:
    """{surface_name: extracted text} for every Streamlit page."""
    out = {}
    for name in sorted(os.listdir(_STREAMLIT_PAGES_DIR)):
        if not name.endswith(".py"):
            continue
        out[f"streamlit:{name}"] = _streamlit_page_surface(
            os.path.join(_STREAMLIT_PAGES_DIR, name)
        )
    return out


@pytest.mark.skipif(
    not os.path.isdir(_STREAMLIT_PAGES_DIR),
    reason="streamlit_app/ absent (installed tree or unpacked sdist, not a checkout)",
)
def test_the_streamlit_extractor_keeps_interpolated_text_in_place():
    """THE PORTABILITY PROPERTY, NAMED RATHER THAN LEFT INCIDENTAL.

    The first version of this extractor sorted its chunks by
    ``(lineno, col_offset)``. Before PEP 701 (Python 3.12) the expressions
    inside an f-string do not carry their true position, so an interpolated
    disclosure sorted THOUSANDS of characters from the claim beside which it
    actually renders: green on 3.12, red on 3.9, 3.10 and 3.11, on a page that
    was correct. CI caught it; nothing in this file asked the question.

    So it is asked directly, on a synthetic whose answer does not depend on how
    any real page happens to be written. If this fails, the surface's character
    offsets are fiction and every distance the proximity gate reports is too.
    """
    import textwrap

    page = os.path.join(_STREAMLIT_PAGES_DIR, "1_Pipeline_Analyzer.py")
    surface = _streamlit_page_surface(page)
    qualifier = _normalise(_streamlit_disclosure_substitutions()[
        "readiness_inline_qualifier"
    ]).lower()
    low = surface.lower()
    assert qualifier in low, (
        "the analyzer page's interpolated readiness qualifier does not appear "
        "in the extracted surface at all — the substitution is not firing"
    )

    claim = _READINESS_CLAIM.search(low)
    assert claim, "no readiness claim in the extracted analyzer surface"
    nearest = min(
        abs(m.start() - claim.start())
        for m in re.finditer(re.escape(qualifier), low)
    )
    limit = _disclosure_proximity_limit()
    assert nearest <= limit, (
        f"the interpolated qualifier lands {nearest:,} characters from the "
        f"page's first readiness claim (limit {limit:,}). On the page itself "
        "it renders in the same column as the metric. The extractor is "
        "ordering interpolations by ast position again, which is unreliable "
        "before Python 3.12 — order must come from the traversal.\n\n"
        + textwrap.shorten(surface[:400], 380)
    )


@pytest.mark.skipif(
    not os.path.isdir(_STREAMLIT_PAGES_DIR),
    reason="streamlit_app/ absent (installed tree or unpacked sdist, not a checkout)",
)
def test_readiness_disclosure_is_adjacent_to_the_claim(rendered):
    """Every readiness claim carries a disclosure within reach of it.

    SEVEN SURFACES NOW, NOT FOUR (1.5.2 audit F2). The four generated documents
    plus every Streamlit page, because the page carrying the LARGEST rendered
    instance of this claim was the one surface the gate could not see.
    """
    surfaces = {
        name: rendered[name]
        for name in ("word", "pdf", "excel", "markdown")
        if name in rendered
    }
    assert surfaces, "no document surfaces rendered — this gate would pin nothing"

    streamlit_surfaces = _streamlit_surfaces()
    assert streamlit_surfaces, (
        "no Streamlit page was extracted — this half of the gate would pin "
        "nothing, which is the state F2 found it in"
    )
    # FAIL CLOSED ON THE EXTRACTOR ITSELF. If the AST walk stopped finding
    # rendered strings — a renamed render call, a page rewritten around a
    # helper — every Streamlit surface would go empty and pass vacuously.
    analyzer = streamlit_surfaces.get("streamlit:1_Pipeline_Analyzer.py", "")
    assert _READINESS_CLAIM.search(analyzer), (
        "the Pipeline Analyzer surface contains no readiness claim at all. "
        "That page renders the grade three ways, so the extractor is broken, "
        "not the page."
    )
    surfaces.update(streamlit_surfaces)

    failures = _disclosure_proximity_failures(surfaces)
    assert not failures, (
        "Readiness claims render too far from their disclosure. A CDE who reads "
        "the grade and stops has been told a house heuristic is an assessment:\n"
        + "\n".join(failures)
    )


def test_proximity_gate_fails_when_the_disclosure_is_far_away():
    """THE RED PROOF. Drive the gate with a document shaped like 1.5.0's PDF.

    A synthetic surface: the readiness claim at the top, the disclosure at the
    far end, nothing in between. This is 1.5.0's actual geometry — claim on
    page 1, disclosure on page 26 — and the flattened presence check called it
    green. If this test ever passes-by-not-failing, the proximity gate has
    stopped measuring proximity.
    """
    far = (
        "application readiness assessment grade b 83.0/100 "
        + ("filler " * 4_000)
        + "this tool's own unsourced house heuristic"
    )
    failures = _disclosure_proximity_failures({"synthetic_far": far})
    assert failures, (
        "the proximity gate did NOT fail on a disclosure ~28,000 characters "
        "from its claim — it is not measuring proximity"
    )
    assert "characters away" in failures[0], failures

    # And the same document with the disclosure moved adjacent must pass,
    # so the gate is not simply failing everything it is shown.
    near = (
        "application readiness assessment grade b 83.0/100 "
        "this tool's own unsourced house heuristic "
        + ("filler " * 4_000)
    )
    assert not _disclosure_proximity_failures({"synthetic_near": near})


# ---------------------------------------------------------------------------
# Derivation: the list must not be a list somebody once wrote down
#
# BOTH SWEEPS BELOW READ THE SOURCE TREE, so both are checkout-only.
#
# release.yml's sdist job deliberately runs the suite from a directory holding
# ONLY what the tarball shipped, with nmtcapp/ ABSENT from the working
# directory (a bare nmtcapp/ there would shadow the installed package as a
# namespace package and turn the job into a test of the tarball's own source
# tree). So _repo_root()/nmtcapp does not exist there, and it never can.
#
# This is not a cosmetic skip. Before 1.2.1 the consumed sweep did not skip and
# did not fail either: os.walk over a directory that does not exist yields
# nothing, silently, so _consumed() returned [] for every constant, nothing was
# ever "consumed", and the sweep reported success having adjudicated zero
# constants. It failed only because _module_constants() opened schema.py first
# and raised FileNotFoundError — an accident of ordering standing in for a
# guard. Both are fixed: the sweeps skip where the tree is absent, and
# _consumer_roots() raises rather than walking nothing where it is present.
# ---------------------------------------------------------------------------

def _repo_root() -> str:
    return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


_SOURCE_TREE_PRESENT = os.path.isdir(os.path.join(_repo_root(), PACKAGE_ROOT))

_needs_source_tree = pytest.mark.skipif(
    not _SOURCE_TREE_PRESENT,
    reason=(
        f"{PACKAGE_ROOT}/ is not present next to tests/ (this is an unpacked "
        "sdist or an installed tree, not a checkout). These two sweeps read the "
        "package's SOURCE to enumerate its constants; release.yml's sdist job "
        "deliberately runs with no source tree in the working directory, so the "
        "question is not the sdist's to answer. Every pin still runs there — "
        "only the derivation of the list is skipped."
    ),
)


def _module_constants() -> dict:
    """{qualified_name: module} for every module-level UPPER_CASE assignment."""
    found = {}
    root = _repo_root()
    for mod, rel in DATA_MODULES.items():
        path = os.path.join(root, rel)
        tree = ast.parse(open(path, encoding="utf-8").read())
        for node in tree.body:
            targets = (
                node.targets if isinstance(node, ast.Assign)
                else [node.target] if isinstance(node, ast.AnnAssign)
                else []
            )
            for t in targets:
                if isinstance(t, ast.Name) and t.id.isupper():
                    found[f"{mod}.{t.id}"] = mod
    return found


def _package_constants() -> dict:
    """{qualified_name: repo-relative path} for ALL of nmtcapp/, not just data/.

    Same shape as _module_constants but over the whole shipped package, because
    the label dict that let the deep/severe swap through lives in a renderer.

    NO CALLER MAY HAND-TYPE THIS COUNT. See :func:`_sweep_census`.
    """
    found = {}
    root = _repo_root()
    for dirpath, _dirs, files in os.walk(os.path.join(root, PACKAGE_ROOT)):
        for f in sorted(files):
            if not f.endswith(".py"):
                continue
            path = os.path.join(dirpath, f)
            mod = os.path.splitext(f)[0]
            tree = ast.parse(open(path, encoding="utf-8").read())
            for node in tree.body:
                targets = (
                    node.targets if isinstance(node, ast.Assign)
                    else [node.target] if isinstance(node, ast.AnnAssign)
                    else []
                )
                for t in targets:
                    if isinstance(t, ast.Name) and t.id.isupper() and t.id != "__all__":
                        found[f"{mod}.{t.id}"] = os.path.relpath(path, root)
    return found


def _sweep_census() -> dict:
    """Every figure this module would otherwise state about its own scope.

    A HAND-TYPED COUNT INSIDE A GATE IS A CLAIM LIKE ANY OTHER (FIX-2 G-1).

    Three sites in this file, and one paragraph of CHANGELOG.md, said the
    package sweep covers **133** constants. It covered 149 — and every one of
    those figures had been measured on 324e9cd, the artifact the hostile audit
    rejected, then carried forward unremeasured. The same round's other
    published figures were wrong the same way: "93 outside data/" read 77,
    "124 with streamlit_app/" read 108, "7 FMT_* constants" read 8,
    "DISTRESS_DISPLAY has 9 rows" read 15.

    Not one of them could fail. They were prose inside a file whose entire
    premise is that its numbers are evidence — which is FLOOR=440's shape one
    round later, and the reason this function exists rather than a corrected
    literal.

    So nothing states a count any more. Every figure is derived here, on the
    tree under test, and the two places that need one — the fail-closed floor
    below and the docstring shown to a reviewer — read it from this dict.
    """
    consts = _package_constants()
    data_prefix = "data" + os.sep
    fmt_names = [n for n in consts if n.startswith("excel_builder.FMT_")]

    from nmtcapp.renderers.styles import DISTRESS_DISPLAY

    return {
        "package_constants": len(consts),
        "outside_data": sum(1 for p in consts.values()
                            if not p[len(PACKAGE_ROOT) + 1:].startswith(data_prefix)),
        "consumer_roots": len(CONSUMER_ROOTS),
        "fmt_constants": len(fmt_names),
        "distress_display_rows": len(DISTRESS_DISPLAY),
        "distress_label_pins": sum(
            1 for p in PINS if p.constant.startswith("styles.DISTRESS_DISPLAY")
        ),
        "pinned_rows": len(PINS),
        "waivers": len(WAIVERS),
        "known_deferred": len(KNOWN),
    }


def _resolves(dotted: str) -> bool:
    """Can this dotted name be imported and walked to a real object?

    ``module.NAME`` for a module-level constant, ``module.Class.attr`` for a
    class attribute. Used by test_every_pinned_constant_name_resolves, which
    must not reduce "exists" to "was collected by the constant sweep" —
    the sweep only walks UPPERCASE module-level assignments, and a rendered
    heading is allowed to come from somewhere else.
    """
    import importlib

    parts = dotted.split(".")
    root = _repo_root()
    for dirpath, _dirs, files in os.walk(os.path.join(root, PACKAGE_ROOT)):
        if f"{parts[0]}.py" not in files:
            continue
        rel = os.path.relpath(os.path.join(dirpath, parts[0]), root)
        modname = rel.replace(os.sep, ".")
        try:
            obj = importlib.import_module(modname)
        except Exception:                                  # pragma: no cover
            continue
        for attr in parts[1:]:
            obj = getattr(obj, attr, None)
            if obj is None:
                break
        else:
            return True
    return False


def _consumer_roots() -> list:
    """The trees a constant can be consumed from. Raises if one is missing.

    FAILS CLOSED, and the reason is specific rather than defensive: os.walk over
    an absent directory yields no files and no error, so a root that quietly
    stopped existing would make every constant look unconsumed and the sweep
    would pass having checked nothing.
    """
    root = _repo_root()
    paths = []
    for name in CONSUMER_ROOTS:
        p = os.path.join(root, name)
        if not os.path.isdir(p):
            raise AssertionError(
                f"consumer root {name}/ is missing from {root}. os.walk would "
                "return nothing from it WITHOUT error, so the consumed sweep "
                "would report every constant unconsumed and pass having "
                "adjudicated none. Fix the root or remove it from "
                "CONSUMER_ROOTS deliberately."
            )
        paths.append(p)
    return paths


def _consumed(constant_name: str) -> list:
    """Files outside nmtcapp/data/ that reference this constant by name.

    Scans streamlit_app/ as well as nmtcapp/ (1.2.1 S-2): a constant read only
    by a Streamlit page reaches a surface a CDE looks at, and under the old
    nmtcapp/-only scan it counted as consumed by nobody.
    """
    root = _repo_root()
    bare = constant_name.split(".", 1)[1]
    pattern = re.compile(rf"\b{re.escape(bare)}\b")
    hits = []
    for tree_root in _consumer_roots():
        for dirpath, _dirs, files in os.walk(tree_root):
            if os.path.join("nmtcapp", "data") in dirpath:
                continue
            for f in files:
                if not f.endswith(".py"):
                    continue
                p = os.path.join(dirpath, f)
                with open(p, encoding="utf-8") as fh:
                    if pattern.search(fh.read()):
                        hits.append(os.path.relpath(p, root))
    return sorted(hits)


def _adjudicated(name: str) -> bool:
    """Is this exact constant name pinned or waived?

    SUBSCRIPTS DO NOT PROPAGATE UPWARD (1.2.1 S-2). Through 1.2.1-rc this
    compared against ``{p.constant.split("[")[0] for p in PINS}``, so pinning
    ONE key of a dict marked the WHOLE dict adjudicated — every other key,
    including keys added later, inherited a pin that says nothing about them.
    NMTC_PROGRAM_CONSTRAINTS has six keys and four pins; the two unpinned keys
    were adjudicated only because somebody had also written waivers for them by
    hand, not because the sweep asked.

    Now ``X[k]`` adjudicates ``X[k]`` and nothing else, and the bare name ``X``
    adjudicates the whole constant — which is correct for a scalar and is a
    deliberate, visible choice for a dict.
    """
    return (name in {p.constant for p in PINS}
            or name in WAIVERS or name in KNOWN)


@_needs_source_tree
def test_every_consumed_constant_is_pinned_or_waived():
    """Sweep the data modules; every consumed constant must be adjudicated.

    THIS IS THE PART THAT KEEPS THE LIST HONEST. The starting list for 1.2.1
    came from the 1.2.0 audit, and an inherited list goes stale the first time
    somebody adds a constant. This test walks nmtcapp/data/ and requires that
    every constant any other module reads is either pinned to a rendered string
    or waived with a stated reason.
    """
    constants = _module_constants()
    assert len(constants) >= 20, (
        f"the data-module sweep found only {len(constants)} constants. "
        "nmtcapp/data/{schema,benchmark_thresholds}.py carry far more than "
        "that; a number this small means the parse found nothing and the "
        "sweep is about to pass vacuously."
    )

    unadjudicated = []
    for name in sorted(constants):
        if _adjudicated(name):
            continue
        consumers = _consumed(name)
        if consumers:
            unadjudicated.append(f"  {name}  read by: {', '.join(consumers)}")

    assert not unadjudicated, (
        f"{len(unadjudicated)} constant(s) in nmtcapp/data/ are read by other "
        "modules and are neither pinned to a rendered string nor waived with a "
        f"reason in {os.path.basename(PIN_PATH)}.\n\n"
        "Add a pin (with the literal string it prints and its source) or a "
        "WAIVE row saying why it reaches no rendered surface. Do not delete "
        "this test to make it pass.\n\n" + "\n".join(unadjudicated)
    )


def _strings_in(value, depth: int = 0):
    """Every string reachable inside a constant, keys and values alike."""
    if depth > 3:
        return
    if isinstance(value, str):
        yield value
    elif isinstance(value, dict):
        for k, v in value.items():
            yield from _strings_in(k, depth + 1)
            yield from _strings_in(v, depth + 1)
    elif isinstance(value, (list, tuple, set, frozenset)):
        for v in value:
            yield from _strings_in(v, depth + 1)


def _rendered_blob(rendered: dict) -> str:
    return "\n".join(rendered[s] for s in DOCUMENT_SURFACES + TEXT_SURFACES)


def _shape_tokens(rendered: dict) -> set:
    """The number formats actually applied, as exact tokens.

    Matched EXACTLY rather than as substrings with a length floor, because a
    number format is a token and not prose: '"$"#,##0' is eight characters,
    below MIN_RENDERED_STRING, and would never trip the substring sweep — while
    a document containing that exact token by accident is not a thing that
    happens. Without this the excel_cell_formats surface would exist and demand
    nothing, which is a gate that cannot fail dressed as a new surface.
    """
    tokens = set()
    for surface in SHAPE_SURFACES:
        for line in rendered[surface].splitlines():
            if " · " in line:
                tokens.add(line.rsplit(" · ", 1)[-1].strip())
    return tokens


@_needs_source_tree
def test_every_rendering_constant_is_pinned_or_waived(rendered):
    """ASK THE ARTIFACT, NOT THE DIRECTORY: does this constant's text appear?

    This is the sweep that would have caught the deep/severe label swap, and
    the reason it is asked this way rather than by widening DATA_MODULES.

    Widening DATA_MODULES to every module that renders was measured first:
    97 constants outside nmtcapp/data/ would each have needed a hand-written
    waiver, most of them saying "this is a colour" or "this is a point size".
    A file of 139 rows where two-thirds are ceremony is a file nobody reads,
    and a waiver nobody reads is the rubber stamp this gate exists to avoid.

    So the non-rendering side is adjudicated BY THE FIXTURE instead. For every
    module-level constant in nmtcapp/, the sweep asks whether any string it
    holds appears in the rendered document. If none does, the constant needs no
    row: the artifact says so on every run, and it cannot go stale the way a
    written waiver can. If one does, the constant publishes text and must be
    pinned or waived by hand.

    NO COUNT IS STATED HERE ON PURPOSE (FIX-2 G-1). This docstring used to
    open a sentence with "Measured on this tree" and follow it with 133 — a
    figure taken on 324e9cd, the rejected artifact, and 149 by the time
    anybody read it. See :func:`_sweep_census` and
    ``test_the_sweep_states_no_count_it_has_not_derived``, which prints the
    live figures and fails if a literal comes back into this file.

    WHAT THIS SWEEP STILL CANNOT SEE, stated rather than implied:

      * A constant that shapes the artifact WITHOUT contributing text. Colours,
        fonts and page margins are the harmless case; the
        excel_builder.FMT_* number formats are not, because a percent format on
        a dollar column changes every figure a reviewer reads and leaves the
        extracted text identical. That class is why the excel_cell_formats
        surface exists — the formats ARE text there, keyed to their column, so
        the FMT_* constants land in this sweep like any other.
      * A constant that renders only on a path this fixture does not take.
        benchmarks._METHODOLOGY reaches app.benchmark(), which app.generate()
        never calls, so it is silent here and waived by hand instead.
      * A string shorter than MIN_RENDERED_STRING. See that constant.
    """
    constants = _package_constants()
    # THE FLOOR IS A FLOOR, NOT A CENSUS (FIX-2 G-1). It used to read "this
    # tree has ~133" in its own failure message, a figure measured on a
    # rejected artifact and never remeasured — a hand-typed count inside a
    # fail-closed assert, in a file whose premise is that its numbers are
    # evidence. The floor's job is to catch a walk that found nothing; it has
    # no business asserting what the tree contains. The census is derived.
    census = _sweep_census()
    assert census["package_constants"] == len(constants)
    assert len(constants) >= 100, (
        f"the package sweep found only {len(constants)} constants under "
        f"{PACKAGE_ROOT}/. A number this small means the walk found nothing "
        "and the sweep is about to pass vacuously."
    )

    blob = _rendered_blob(rendered)
    assert len(blob) > 20_000, (
        f"the rendered surfaces total {len(blob)} characters. A truncated or "
        "degraded render would make every constant look silent and this sweep "
        "would demand nothing of anybody."
    )
    shape_tokens = _shape_tokens(rendered)
    assert len(shape_tokens) >= 3, (
        f"the excel_cell_formats surface yielded {len(shape_tokens)} distinct "
        "number formats. This workbook applies several; a set this small means "
        "the surface parsed nothing and every FMT_* constant is about to look "
        "silent."
    )

    import importlib
    unadjudicated = []
    for name in sorted(constants):
        if _adjudicated(name):
            continue
        rel = constants[name]
        modname = rel[:-3].replace(os.sep, ".")
        try:
            module = importlib.import_module(modname)
        except Exception:                              # pragma: no cover
            continue
        value = getattr(module, name.split(".", 1)[1], None)
        hits = [
            s for s in set(_strings_in(value))
            if (len(s) >= MIN_RENDERED_STRING and _normalise(s) in blob)
            or s in shape_tokens
        ]
        if hits:
            shown = ", ".join(repr(h[:60]) for h in sorted(hits)[:3])
            unadjudicated.append(f"  {name}  ({rel})\n      renders: {shown}")

    assert not unadjudicated, (
        f"{len(unadjudicated)} constant(s) under {PACKAGE_ROOT}/ publish text "
        "into the rendered application and are neither pinned nor waived in "
        f"{os.path.basename(PIN_PATH)}.\n\n"
        "This sweep does not care which directory the constant lives in. It "
        "asks whether a string it holds reaches the page a CDE files. The "
        "deep/severe label swap reached four appendices and passed 937 tests "
        "because a LABEL DICT in a renderer had never been a candidate for "
        "pinning.\n\n"
        "Pin it to the string it prints — anchored to the row it prints on, if "
        "it is a label — or waive it with a reason a reviewer can check. Do "
        "not delete this test to make it pass.\n\n" + "\n".join(unadjudicated)
    )


@_needs_source_tree
def test_every_dict_key_the_package_reads_is_adjudicated():
    """A subscripted row covers ONE key. Nothing else covers a key at all.

    ``schema.NMTC_PROGRAM_CONSTRAINTS[credit_rate]`` says nothing about
    ``[leverage_ratio_typical]``, and before 1.2.1 the sweep treated it as
    though it did — pinning any one key marked the whole dict adjudicated,
    including keys added afterwards.

    THE FIX FOR THAT DISABLED THIS GATE ENTIRELY (FIX-2 G-2). The rows S-2
    added to enable key checking are what turned it off: every one of the ten
    dicts acquired a BARE-NAME row, and this loop opened with
    ``if _adjudicated(name): continue``. Measured on the branch head:

        dicts in DATA_MODULES        10   (56 string keys)
        short-circuited by bare name 10
        dicts actually examined       0
        keys actually examined        0

    A gate that cannot fail, inside the gate built to close the last gate that
    could not fail — and with no floor to say so.

    A BARE NAME ADJUDICATES THE DICT AS A VALUE, NOT ITS KEYS. That distinction
    is the whole point. Seven of the ten bare-name rows are WAIVERS reading
    "the dict itself renders nowhere", which is a claim ABOUT THE DICT OBJECT;
    whether ``IMPACT_BENCHMARKS["jobs_per_million_qei_low"]`` reaches the page
    is a different question, and it is the question this test asks. So there is
    no short-circuit any more: every key the package subscripts carries its own
    row, whatever its parent says.
    """
    import importlib
    adjudicated_keys = {
        n for n in ({p.constant for p in PINS} | set(WAIVERS) | set(KNOWN))
        if "[" in n
    }
    # Read each consumer file once. The old loop re-read the whole tree per
    # key — 56 keys x every .py file — which is why it was never noticed that
    # the outer `continue` meant it read nothing at all.
    sources = []
    for tree_root in _consumer_roots():
        for dirpath, _dirs, files in os.walk(tree_root):
            for f in sorted(files):
                if f.endswith(".py"):
                    sources.append(
                        open(os.path.join(dirpath, f), encoding="utf-8").read()
                    )

    missing = []
    dicts_examined = 0
    keys_examined = 0
    for name, mod in sorted(_module_constants().items()):
        bare = name.split(".", 1)[1]
        value = getattr(importlib.import_module(f"nmtcapp.data.{mod}"), bare, None)
        if not isinstance(value, dict):
            continue
        dicts_examined += 1
        for key in value:
            if not isinstance(key, str):
                continue
            keys_examined += 1
            # Only keys the package actually reads. A declared-and-unread key
            # publishes nothing and is the dict's business, not this gate's.
            pattern = re.compile(
                rf"\b{re.escape(bare)}\b\s*\[\s*[\"']{re.escape(key)}[\"']"
            )
            read = any(pattern.search(text) for text in sources)
            if read and f"{name}[{key}]" not in adjudicated_keys:
                missing.append(f"  {name}[{key}]  is read but not adjudicated")

    # THE FLOOR THIS GATE DID NOT HAVE. Zero examined is not zero defects; it
    # is a gate that has stopped running, and it reported success for a whole
    # release.
    assert dicts_examined > 0, (
        "this gate examined ZERO dicts. DATA_MODULES holds ten. Either the "
        "module walk broke or a short-circuit is back — and either way the "
        "test was about to pass having adjudicated nothing, which is the "
        "exact failure it was built to close."
    )
    assert keys_examined > 0, (
        f"this gate examined {dicts_examined} dict(s) and ZERO keys. A dict "
        "sweep that reaches no key adjudicates nothing."
    )
    assert sources, (
        "no consumer source files were read, so every key would look unread "
        "and this gate would demand nothing of anybody."
    )

    assert not missing, (
        f"{len(missing)} dict key(s) are subscripted by the package and carry "
        "no row of their own. A pin on a SIBLING key does not adjudicate "
        "them, and neither does a row on the BARE DICT NAME — that row is a "
        f"claim about the dict object.\n\n(examined {dicts_examined} dicts, "
        f"{keys_examined} keys)\n\n" + "\n".join(missing)
    )


def test_no_waiver_shadows_a_live_pin():
    """A constant cannot be both pinned and waived — one of them is a lie."""
    both = sorted({p.constant for p in PINS} & set(WAIVERS))
    assert not both, (
        "these constants are both pinned and waived; the waiver claims they "
        f"reach no rendered surface while the pin asserts they do: {both}"
    )


@_needs_source_tree
def test_every_pinned_constant_name_resolves():
    """A pin's CONSTANT column must name something that exists.

    The registry's own format note says CONSTANT is "the dotted name of the
    constant this row pins, for the reviewer and for the coverage meta-test".
    Five 1.2.1 rows named modules the package does not have — ``_statute.``
    and ``_workbook.`` — which are prose labels for quotations, not constants.
    Neither sweep could ever match them, so they sat outside the derivation
    while looking like they were inside it.

    A row that pins a QUOTATION rather than a constant is legitimate; it just
    has to say so. The QUOTE: prefix is that statement.
    """
    known = set(_package_constants())
    unknown = []
    for pin in PINS:
        name = pin.constant
        if name.startswith("QUOTE:"):
            continue
        base = name.split("[", 1)[0]
        if base in known:
            continue
        # A PRODUCER NEED NOT BE A MODULE-LEVEL CONSTANT (FIX-2 G-5 sweep).
        # _package_constants() walks module-level UPPERCASE assignments, which
        # is the right scope for the SWEEP — but it is the wrong test for
        # "does this name exist". The five section headings a reviewer
        # navigates by come from a class attribute
        # (SectionEPriorAwards.title), not from a constant; they were pinned
        # against styles.SECTION_META, a dict READ BY NOTHING whose "E" entry
        # had already drifted from the class, and the pins passed anyway. So
        # this resolves the dotted name for real instead of asking whether the
        # sweep happens to have collected it.
        if _resolves(base):
            continue
        unknown.append(name)
    unknown = sorted(set(unknown))
    assert not unknown, (
        "these pin rows name a constant that does not exist anywhere under "
        f"{PACKAGE_ROOT}/, so no sweep can ever reach them:\n  "
        + "\n  ".join(unknown)
        + "\n\nEither correct the name, or prefix it 'QUOTE:' if the row pins a "
        "quotation this package holds as inline text rather than as a constant."
    )


# ---------------------------------------------------------------------------
# FIX-2 G-1: a hand-typed count inside a gate is a claim like any other
# ---------------------------------------------------------------------------

@_needs_source_tree
def test_the_sweep_states_no_count_it_has_not_derived():
    """Scope figures must be DERIVED here, never typed into prose.

    Three sites in this file said the package sweep covers 133 constants while
    it covered 149, and CHANGELOG.md said the same. Every one of those figures
    was measured on 324e9cd — the artifact the hostile audit rejected — and
    carried forward unremeasured, which is FLOOR=440's shape one round later
    inside the file whose whole premise is that its numbers are evidence.

    A literal count cannot fail. So this test does two things:

      1. prints the live census, so a reviewer or a CHANGELOG author can copy
         a figure that is true of the tree in front of them rather than of a
         tree somebody measured once, and
      2. fails if a bare scope integer reappears in this module's prose.

    It deliberately does NOT assert an expected value for any census figure.
    An expected value is the thing being removed.
    """
    census = _sweep_census()
    for key, value in sorted(census.items()):
        print(f"  {key:26s} {value}")

    # Every figure must be positive, or the derivation itself has broken and
    # the census would read as authoritative while measuring nothing.
    for key, value in census.items():
        assert value > 0, f"{key} derived as {value}; the sweep measured nothing"

    # WHAT IS FORBIDDEN IS THE CLAIM SHAPE, NOT THE DIGIT. This file records
    # the stale figures on purpose — a reader needs to know 133 was checked
    # and found wrong — so a bare-integer denylist would forbid the history
    # along with the claim. These patterns match a CURRENT-TENSE assertion
    # about the tree under test, which is the thing that must be derived.
    source = open(__file__, encoding="utf-8").read()
    source = source.split("def test_the_sweep_states_no_count_it_has_not_derived")[0]
    claim_shapes = {
        r"this tree has ~?\d": "asserts a package-sweep count",
        r"[Mm]easured on this tree: ~?\d": "asserts a measured count",
        r"\d+ constants are swept": "asserts a swept-constant count",
        r"the (?:eight|seven|nine|ten|eleven|twelve|\d+) (?:excel_builder\.)?FMT_\*":
            "asserts how many number-format constants exist",
        r"(?:eight|nine|ten|eleven|twelve|\d+) of the "
        r"(?:eight|nine|ten|eleven|twelve|\d+)": "asserts a pin-count ratio",
    }
    reappeared = []
    for pattern, why in claim_shapes.items():
        for m in re.finditer(pattern, source):
            line = source[:m.start()].count("\n") + 1
            reappeared.append(f"line {line}: {m.group(0)!r} — {why}")
    assert not reappeared, (
        "a current-tense scope count is stated in this module's prose:\n  "
        + "\n  ".join(reappeared) + "\n\n"
        "Derive it in _sweep_census() and interpolate, or do not state it. "
        "A figure a reader cannot check is a claim, and this file exists "
        "because claims in gates are what this package keeps shipping. "
        "Recording a PAST figure as history is fine; asserting a present one "
        "is not."
    )


@_needs_source_tree
def test_the_changelogs_sweep_figures_match_the_tree():
    """CHANGELOG.md quoted the same rejected-artifact counts to the reader.

    A release note is the one surface a CDE or a downstream maintainer reads
    without running anything, so a figure there has to be true of the tree it
    describes. This asserts only the figures the CHANGELOG actually states.
    """
    census = _sweep_census()
    path = os.path.join(_repo_root(), "CHANGELOG.md")
    text = open(path, encoding="utf-8").read()

    claimed = re.findall(r"and (\d+) constants are swept", text)
    assert claimed, (
        "CHANGELOG.md no longer states a swept-constant count in the form "
        "this test recognises. If the sentence was reworded, reword the "
        "regex; do not delete the check."
    )
    for value in claimed:
        assert int(value) == census["package_constants"], (
            f"CHANGELOG.md says {value} constants are swept; the tree has "
            f"{census['package_constants']}. The 1.2.1 figure (133) was "
            "measured on 324e9cd, the rejected artifact."
        )


def test_the_changelogs_review_process_sweep_matches_the_tree():
    """The sweep's own headline count, re-derived rather than trusted.

    1.3.0 B1. The Review Process sweep reported "72 mentions across 68 lines".
    Measured across the cycle's commits, NO TREE YIELDS 72/68: 03261c1 — the
    commit that made the claim — gives 75/71, and 63443cc, the pre-round base,
    gives 67/64. The figure was neither the before nor the after, and it sat
    in the release note looking like a measurement.

    Nothing renders off it. That is the point: it is the sweep's statement of
    how much it looked at, and a sweep that miscounts its own corpus is one
    whose coverage claim cannot be checked. Every other hand-typed count in
    this package has gone stale — FLOOR four times, 133 in three places, the
    swept-constant census — and each got a derivation gate instead. This is
    that gate for this count.
    """
    roots = ["nmtcapp", "streamlit_app", "docs", "README.md"]
    present = [r for r in roots
               if os.path.exists(os.path.join(_repo_root(), r))]
    if "docs" not in present:
        pytest.skip("no docs/ tree — an unpacked sdist prunes it, and the "
                    "count in the CHANGELOG is of the repository")

    mentions = 0
    lines = 0
    for root in present:
        full = os.path.join(_repo_root(), root)
        paths = []
        if os.path.isfile(full):
            paths = [full]
        else:
            for dirpath, dirnames, filenames in os.walk(full):
                dirnames[:] = [d for d in dirnames if d != "__pycache__"]
                paths.extend(os.path.join(dirpath, f) for f in filenames)
        for path in sorted(paths):
            try:
                body = open(path, encoding="utf-8").read()
            except (UnicodeDecodeError, OSError):
                continue  # binary; grep --binary-files=without-match skips it
            for line in body.splitlines():
                hits = line.count("Review Process")
                if hits:
                    mentions += hits
                    lines += 1

    text = open(os.path.join(_repo_root(), "CHANGELOG.md"), encoding="utf-8").read()
    claimed = re.findall(r"\*\*(\d+) mentions across (\d+) lines\*\*", text)
    assert claimed, (
        "CHANGELOG.md no longer states the Review Process sweep's corpus size "
        "in the form this test recognises. If the sentence was reworded, "
        "reword the regex; do not delete the check."
    )
    for claimed_mentions, claimed_lines in claimed:
        assert (int(claimed_mentions), int(claimed_lines)) == (mentions, lines), (
            f"CHANGELOG.md says {claimed_mentions} mentions across "
            f"{claimed_lines} lines; the tree has {mentions} across {lines}. "
            "Re-derive it with the two greps recorded beside the claim, or "
            "correct the claim. Do not widen this test."
        )


#: The CHANGELOG's statement of how far the rendered baseline moved, in the one
#: form this test recognises — a blockquote, unwrapped to a single line first::
#:
#:     **13 insertions, 4 deletions** in `tests/rendered_baseline/`, measured
#:     `63443cc`..`ff49064`, in `excel.txt`, `markdown.txt` and `word.txt`.
#:
#: BOTH ENDPOINTS ARE NAMED so the claim can still be re-derived years later,
#: after the branch it was written on has been merged and moved past. The
#: second may be ``HEAD``, which means the working tree and is what an open
#: branch states.
_BASELINE_DELTA_RE = re.compile(
    r"\*\*(\d+) insertions?, (\d+) deletions?\*\* in `tests/rendered_baseline/`, "
    r"measured `([0-9a-fA-F]{7,40})`\.\.`(HEAD|[0-9a-fA-F]{7,40})`, in (.+?)\.?$"
)


def _git(*args) -> str:
    return subprocess.run(
        ["git", *args], cwd=_repo_root(), capture_output=True, text=True,
    ).stdout


def _blockquotes(text: str):
    """Yield each markdown blockquote as one unwrapped line.

    The claim is written as a quote and wraps over three source lines; matching
    it raw would make the regex a test of where the author pressed return.

    Example::

        list(_blockquotes("> a\n> b\n\nc"))   # -> ['a b']
    """
    current = []
    for line in text.splitlines():
        if line.startswith(">"):
            current.append(line.lstrip(">").strip())
        elif current:
            yield " ".join(current)
            current = []
    if current:
        yield " ".join(current)


@_needs_source_tree
def test_the_changelogs_rendered_baseline_delta_matches_the_tree():
    """The baseline-movement figures must be the tree's, not a person's memory.

    1.3.0 FIX-2 B3. The 1.3.0 entry claimed "seven insertions, five deletions"
    of a tree that yields 13 and 4, under a table headed "every changed line
    classified, zero unexplained" that listed three rows which have never
    existed in this branch and omitted the eight that do. SIXTH stale
    hand-typed count of the cycle, produced by the commit whose own narrative
    is about the fifth.

    A baseline delta is the one figure in a release note a reviewer uses to
    decide whether to read the rendered diff. Being wrong about it in the safe
    direction — claiming a smaller change than happened — is the direction that
    stops the diff being read.

    So it is derived. Every claim in the CHANGELOG names both endpoints and the
    files it says moved; this re-runs ``git diff --numstat`` between them and
    fails on either the counts or the set.
    """
    if not os.path.isdir(os.path.join(_repo_root(), ".git")):
        pytest.skip("not a git checkout — this claim is about the repository")

    text = open(os.path.join(_repo_root(), "CHANGELOG.md"), encoding="utf-8").read()
    claims = [m.groups() for m in
              (_BASELINE_DELTA_RE.search(q) for q in _blockquotes(text)) if m]
    assert claims, (
        "CHANGELOG.md no longer states a rendered-baseline delta in the form "
        "this test recognises. If the sentence was reworded, reword the regex; "
        "do not delete the check — this figure has been wrong once already."
    )

    for ins, dels, base, head, surfaces in claims:
        # A SHALLOW CLONE CANNOT ANSWER THIS, and saying so is not the same as
        # passing. `actions/checkout` defaults to fetch-depth: 1, under which
        # the commits the claim names are simply absent and `git diff` returns
        # nothing at all — indistinguishable, to an assertion, from a claim
        # about a movement that never happened. CI is configured with
        # fetch-depth: 0 for exactly this gate and
        # test_ci_fetches_enough_history_to_answer_this_gate holds that line,
        # so this skip is a courtesy to a contributor's `--depth 1` and can
        # never become the way the suite goes green.
        unreachable = [
            sha for sha in (base, head) if sha != "HEAD"
            and subprocess.run(["git", "cat-file", "-e", f"{sha}^{{commit}}"],
                               cwd=_repo_root(), capture_output=True).returncode
        ]
        if unreachable:
            pytest.skip(
                f"{', '.join(unreachable)} not in this clone — a shallow "
                "checkout cannot re-derive a diff between two commits. CI "
                "sets fetch-depth: 0 so this gate runs there."
            )

        args = ["diff", "--numstat", base] + ([head] if head != "HEAD" else [])
        raw = _git(*args, "--", "tests/rendered_baseline/")
        assert raw.strip(), (
            f"git produced no diff between {base} and {head} for "
            "tests/rendered_baseline/ — both commits are present, so the claim "
            "describes a movement that did not happen."
        )
        rows = [ln.split("\t") for ln in raw.strip().splitlines()]
        got_ins = sum(int(r[0]) for r in rows if r[0] != "-")
        got_dels = sum(int(r[1]) for r in rows if r[1] != "-")
        got_files = {os.path.basename(r[2]) for r in rows}
        named = set(re.findall(r"`([a-z]+\.txt)`", surfaces))

        assert (got_ins, got_dels) == (int(ins), int(dels)), (
            f"CHANGELOG.md says {ins} insertions and {dels} deletions in "
            f"tests/rendered_baseline/ between {base} and {head}; the tree has "
            f"{got_ins} and {got_dels}. Re-derive it:\n\n"
            f"    git diff --numstat {base} {'' if head == 'HEAD' else head} "
            "-- tests/rendered_baseline/\n\n"
            "Do not widen this test. Five hand-typed counts in this package "
            "have gone stale and every one was wrong in the direction that "
            "never fails."
        )
        assert got_files == named, (
            f"CHANGELOG.md says the movement between {base} and {head} is in "
            f"{sorted(named)}; the tree moved {sorted(got_files)}. Which "
            "surfaces moved is the claim a reviewer checks a renderer fix "
            "against — B1 was asserted to touch the PDF only."
        )


#: One row of a baseline-movement class table: | Class | Lines | +/- | ... |
_CLASS_ROW_RE = re.compile(
    r"^\|\s*(?!Class\b)(?!-)(.+?)\s*\|\s*(\d+)\s*\|\s*\+(\d+)\s*/\s*−(\d+)\s*\|",
    re.M,
)


def test_the_changelogs_baseline_class_table_adds_up():
    """The BREAKDOWN, not just the total — the half the gate above never read.

    WHY THIS EXISTS (1.3.1 G7). The 1.3.0 entry's FIX-2 movement table
    classified 194 changed lines into five rows and stated 10 / 13 / 96 / 9 /
    66. **Every row was wrong, and they summed to 194** — the correct total,
    which is why nothing caught it:
    ``test_the_changelogs_rendered_baseline_delta_matches_the_tree`` parses the
    BLOCKQUOTE and asserts the total and the file set. It was never given the
    table. A hand-typed breakdown under a machine-checked total can be
    arbitrarily wrong so long as it adds up, and this one was — five figures,
    each off, cancelling.

    WHAT THIS CAN AND CANNOT DERIVE, STATED RATHER THAN IMPLIED.

    Four of the five classes are SEMANTIC — "Section B's four distress row
    labels wrap" is a reading of the diff, not a rule over it, and no assertion
    can re-derive it without re-doing the reading. One is not: `Item`/`Value`
    gaining a leading space in extraction is exactly "the changed line, stripped,
    is `Item` or `Value`". So:

      1. every row's +/− split sums to its Lines figure;
      2. the rows' Lines figures sum to insertions + deletions from the tree;
      3. the `Item`/`Value` row is RE-DERIVED from the diff and must match;
      4. therefore the other four rows are pinned to the residual, jointly.

    (3) is what would have caught what shipped. (4) is weaker than (3) and is
    written down as weaker: three of those four rows could still be wrong in
    ways that cancel. Making them derivable means giving each row a matching
    rule in the CHANGELOG itself, which is a change to how release notes are
    written and is not a patch-release change.
    """
    if not os.path.isdir(os.path.join(_repo_root(), ".git")):
        pytest.skip("not a git checkout — this claim is about the repository")

    text = open(os.path.join(_repo_root(), "CHANGELOG.md"), encoding="utf-8").read()
    claims = [(m, q) for q, m in
              ((q, _BASELINE_DELTA_RE.search(q)) for q in _blockquotes(text)) if m]
    assert claims, "no rendered-baseline delta claim found to check a table against"

    checked = 0
    for match, _quote in claims:
        ins, dels, base, head, _surfaces = match.groups()
        # The table follows its blockquote in the source. The blockquote is
        # matched UNWRAPPED, so it is located here by the one fragment that
        # survives the unwrap intact rather than by the whole match.
        needle = f"**{ins} insertions, {dels} deletions**"
        if needle not in text:
            continue
        after = text[text.index(needle) + len(needle):]
        # THE WINDOW IS BOUNDED BY THE TABLE, NOT BY A MAGIC 4,000 (1.5.1
        # audit). This read ``after[:4000]``, and a 4,000-character window is a
        # silent row filter: prose added between the blockquote and the table
        # pushes the LAST rows out of the window, the sum drops by exactly
        # those rows, and the failure message blames the arithmetic. That
        # happened in this round -- the gate reported "13 rows sum to 152"
        # against a 15-row table, and the two missing rows were the two most
        # recently added.
        #
        # A gate that adjudicates a SUBSET and reports on the whole is the
        # defect this file exists to chase; it does not get to have one. The
        # window now runs to the end of the table -- the first blank line after
        # the last row -- so every row is either counted or absent, and a table
        # that grows cannot silently fall out of range.
        # Bound the region to THIS claim: stop at the next release heading or
        # the next delta claim, so a table belonging to a different entry can
        # never be counted against this one.
        stop = min(
            [m.start() for m in re.finditer(r"\n##\s", after)]
            + [m.start() for m in re.finditer(r"insertions,\s*\d+\s*deletions", after)]
            + [len(after)]
        )
        region = after[:stop]
        if "| Class | Lines |" in region:
            table = region[region.index("| Class | Lines |"):]
            end = re.search(r"\n\s*\n", table)
            table = table[: end.start()] if end else table
        else:
            table = region
        rows = _CLASS_ROW_RE.findall(table)
        if not rows:
            continue
        checked += 1

        unreachable = [
            sha for sha in (base, head) if sha != "HEAD"
            and subprocess.run(["git", "cat-file", "-e", f"{sha}^{{commit}}"],
                               cwd=_repo_root(), capture_output=True).returncode
        ]
        if unreachable:
            pytest.skip(f"{', '.join(unreachable)} not in this clone")

        total = int(ins) + int(dels)
        for label, lines, plus, minus in rows:
            assert int(plus) + int(minus) == int(lines), (
                f"class row {label!r} says {lines} lines but +{plus}/−{minus} "
                f"is {int(plus) + int(minus)}. The row does not agree with itself."
            )
        summed = sum(int(r[1]) for r in rows)
        assert summed == total, (
            f"the class table's {len(rows)} rows sum to {summed}; the "
            f"blockquote says {ins} insertions and {dels} deletions, which is "
            f"{total}. Re-derive the rows against the diff:\n\n"
            f"    git diff -U0 {base} {'' if head == 'HEAD' else head} "
            "-- tests/rendered_baseline/pdf.txt"
        )

        # (3) — the one row that is a rule and not a reading.
        args = ["diff", "-U0", base] + ([head] if head != "HEAD" else [])
        diff = _git(*args, "--", "tests/rendered_baseline/")
        changed = [ln for ln in diff.splitlines()
                   if ln[:1] in "+-" and not ln[:3] in ("+++", "---")]
        assert changed, "the diff produced no changed lines to classify"
        derived = sum(1 for ln in changed if ln[1:].strip() in ("Item", "Value"))
        stated = [int(lines) for label, lines, _p, _m in rows
                  if "Item" in label and "Value" in label]
        assert len(stated) == 1, (
            "the class table no longer has exactly one `Item`/`Value` row. "
            "That row is the only one this gate can re-derive; if it was "
            "merged away, say so here rather than letting the check evaporate."
        )
        assert stated[0] == derived, (
            f"the class table says {stated[0]} `Item`/`Value` line(s); the "
            f"tree has {derived}. This is the exact figure that shipped wrong "
            "in 1.3.0 (10, against 16), under a total that was right.\n\n"
            f"    git diff -U0 {base} {'' if head == 'HEAD' else head} "
            "-- tests/rendered_baseline/ | grep -E '^[+-] *(Item|Value)$' | wc -l"
        )

    assert checked, (
        "no baseline-delta blockquote is followed by a class table in the form "
        "this gate reads. The table is where 1.3.0's five wrong figures lived; "
        "if it was reformatted, reformat the regex — do not drop the check."
    )


def test_ci_fetches_enough_history_to_answer_this_gate():
    """CI must clone deeply enough for the baseline-delta gate to run.

    The gate above skips on a shallow clone rather than failing, because a
    contributor running `--depth 1` should get "cannot answer", not a red that
    says the CHANGELOG is wrong. That courtesy is also the way the gate could
    quietly stop asking: `actions/checkout` defaults to fetch-depth: 1, so
    deleting one line from ci.yml turns the check into a permanent skip and
    nothing anywhere says so.

    THIS GUARD WAS ITSELF THE SHAPE IT WAS WRITTEN TO PREVENT (1.3.1 G1).
    Through 1.3.0 it read::

        text = open(path).read()
        assert "fetch-depth: 0" in text

    — a substring test against the WHOLE FILE. ci.yml has two jobs. The line
    commented out, or moved onto the `docs` job while the `test` job's checkout
    lost it, satisfied that assertion exactly as well as the correct file did,
    and the gate against "the skip becomes the pass" could not tell the two
    apart. It parses the workflow now and asks the one question that matters:
    does the job that RUNS THE SUITE check out with full history.

    The pattern is on record in this package's CHANGELOG under "a gate that
    cannot fail is also a green tick". A tally used to stand here and in four
    other files; it is gone rather than restated, because a hand-maintained
    count restated five times is the defect the entry it points at is about.
    """
    yaml = pytest.importorskip("yaml")

    path = os.path.join(_repo_root(), ".github", "workflows", "ci.yml")
    if not os.path.exists(path):
        pytest.skip("no .github/workflows/ci.yml — not a checkout of the repo")

    with open(path, encoding="utf-8") as fh:
        workflow = yaml.safe_load(fh)

    jobs = workflow.get("jobs") or {}
    assert TEST_JOB_NAME in jobs, (
        f"ci.yml has no `{TEST_JOB_NAME}` job. This gate names the job that "
        "runs the suite; if it was renamed, rename it here too rather than "
        f"letting the check evaporate. Jobs present: {sorted(jobs)}"
    )

    checkouts = [
        step for step in (jobs[TEST_JOB_NAME].get("steps") or [])
        if str(step.get("uses", "")).startswith("actions/checkout")
    ]
    assert checkouts, (
        f"the `{TEST_JOB_NAME}` job has no actions/checkout step at all, so "
        "there is no clone for the baseline-delta gate to read."
    )

    for step in checkouts:
        depth = (step.get("with") or {}).get("fetch-depth")
        assert depth == 0, (
            f"the `{TEST_JOB_NAME}` job's checkout requests fetch-depth="
            f"{depth!r}; it must be 0. Without full history, "
            "test_the_changelogs_rendered_baseline_delta_matches_the_tree "
            "skips on every CI run and the CHANGELOG's baseline figures go "
            "back to being unchecked prose.\n\n"
            "A `fetch-depth: 0` ANYWHERE ELSE IN THE FILE DOES NOT COUNT. That "
            "is what this gate used to accept, and the docs job carries its "
            "own checkout."
        )


#: The ci.yml job that runs the test suite — the one whose clone depth decides
#: whether the baseline-delta gate can answer. Named once so the two assertions
#: above cannot drift onto different jobs.
TEST_JOB_NAME = "test"


# ---------------------------------------------------------------------------
# FIX-2 G-3: a pin must assert the VALUE, not merely its presence
# ---------------------------------------------------------------------------

# (label, the pinned literal, a regex matching every place the surface prints
#  a value of this kind). The pattern must match the VALUE, in group 1, not
#  the whole phrase — the test compares what it captured against the pin.
_EXCLUSIVE_VALUES = (
    (
        "credit price",
        "0.83",
        r"credit price(?: of)? \$([\d.]+)\s*(?:/|per\s+)?(?:credit|NMTC)",
    ),
    (
        "CDE fee rate",
        "2.5",
        r"CDE fee(?: rate)?(?: of)?[: ]+\s*([\d.]+)%\s*of QEI",
    ),
    (
        "compliance period",
        "7",
        r"([\d]+)[- ]year compliance period",
    ),
)


@pytest.mark.parametrize("label,expected,pattern", _EXCLUSIVE_VALUES,
                         ids=[v[0] for v in _EXCLUSIVE_VALUES])
def test_a_pinned_value_has_no_rival_on_the_same_surface(rendered, label, expected, pattern):
    """"$0.83 appears somewhere" is not "the document says $0.83".

    THE HOLE, MEASURED (probe N1, on the branch head). Setting
    renderers/_methodology's credit price to a literal 0.95 produced ONE FILING
    saying $0.83 in Section D and $0.95 in the methodology appendix two pages
    later — and 955 tests stayed green. The registry pin for
    NMTC_PROGRAM_CONSTRAINTS[standard_credit_price] passed because "$0.83" was
    still in the document; the cross-surface check passed because the credit
    price was excused in consistency_check._UNPAIRED as "a per-credit rate, not
    a pipeline total".

    The registry already states the correct rule and applied it to all nine
    distress rows: a pin must be ANCHORED, not merely present. This is that
    rule applied to the financial rates — every value of the pinned kind that
    the surface prints must BE the pinned value.

    (The shipped validator now compares the two surfaces directly as well —
    consistency_check._SECTION_D_TO_METHODOLOGY — so a CDE with no test suite
    gets the same protection. This test guards the registry's own claim.)
    """
    compiled = re.compile(pattern, re.IGNORECASE)
    seen_anywhere = False
    disagreements = []
    for surface in DOCUMENT_SURFACES:
        found = compiled.findall(rendered[surface])
        if not found:
            continue
        seen_anywhere = True
        rivals = sorted({v for v in found if v.rstrip("0").rstrip(".")
                         != expected.rstrip("0").rstrip(".")})
        if rivals:
            disagreements.append(
                f"  {surface}: prints {sorted(set(found))} for the {label}; "
                f"the pinned value is {expected!r}"
            )

    assert seen_anywhere, (
        f"no rendered surface prints a {label} in the form this test "
        f"recognises ({pattern!r}). Either the wording changed — update the "
        "pattern — or the value stopped rendering, in which case the registry "
        "row that pins it is stale. A pattern that matches nothing is a test "
        "that checks nothing."
    )
    assert not disagreements, (
        f"the same document states more than one {label}:\n"
        + "\n".join(disagreements)
        + "\n\nA pin that only asks whether its literal appears SOMEWHERE "
        "passes over this. The document has to say one thing."
    )
