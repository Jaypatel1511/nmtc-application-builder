"""THE REGRESSION GATE: a rendered line may not change without being reviewed.

WHY THIS EXISTS

Every other gate in this package asks whether a rendered line is ENTITLED to be
there. The invariance gate asks whether a line was derived from the input. The
constant gate asks whether a published value prints as pinned. The attribution
gate asks whether a claim carries a citation. The cross-surface check asks
whether two surfaces agree.

None of them asks WHAT CHANGED.

That is not a theoretical gap. Two consecutive rounds each shipped a defect
created by the round's own fix, in the renderers, and every gate stayed green:

    1.2.1 B-5's fix (splitting the 85%/20% distress bars onto their own rows)
        created a filing that stated 0.0% severe distress in Section B and
        flagged the same projects "Yes" in Appendix B — because the row it
        added read the EXCLUSIVE severe bucket under an INCLUSIVE label.

    1.2.1 B-1's fix (unifying cell formatting across the four renderers)
        created prior-award years rendering as 2,019 on markdown, Word and PDF
        — because the plain-number arm of the new formatter ends in
        ``f"{value:,}"`` and a year is not a quantity.

Both were found by a human reading the output. A byte-diff of all four formats
between v1.2.0 and the branch head, from one fixed fixture, showed both in
ninety seconds. Nobody had ever diffed this package's output against its own
last release.

READING DOES NOT SCALE. This is that diff, kept.

THE METHOD

One fixed, fully-populated fixture — NOT Pipeline.sample(), whose rows changed
between releases, which would report input edits as output changes and bury
everything else. Render all four formats. Project each artifact to text,
INCLUDING the parts that are not text: Excel number formats are extracted
alongside cell values, because the year defect is a number-format defect and
the cell value 2019 is correct either way.

Normalise only what is genuinely non-deterministic — timestamps, temp paths,
today's date. Diff every remaining line against a committed baseline.

A changed line fails. The fix is to read the diff and, if the change is
intended, update the baseline in the same commit — which puts every rendered
change in front of a reviewer in the pull request, as a diff of the DOCUMENT
rather than of the code that produces it.

This is the invariant allowlist's shape applied to CHANGE rather than to
INVARIANCE, and for the same reason: a new line fails by default and has to be
argued past a human.

FAILS CLOSED. A missing format, an empty extraction, a baseline that is absent
or suspiciously short, or a fixture that stops populating a field all ERROR
rather than pass.
"""
from __future__ import annotations

import datetime
import difflib
import os
import re

import pytest

from nmtcapp.core.application import Application
from nmtcapp.core.cde import CDEProfile
from nmtcapp.core.pipeline import Pipeline, PipelineProject

FORMATS = ("markdown", "word", "excel", "pdf")

BASELINE_DIR = os.path.join(os.path.dirname(__file__), "rendered_baseline")

# Below this, the baseline is not describing a whole application and the diff
# would pass on a document that had mostly stopped rendering.
_MIN_BASELINE_LINES = {
    "markdown": 300,
    "word": 180,
    "excel": 500,
    "pdf": 700,
}


# ---------------------------------------------------------------------------
# THE FIXTURE. Written once, here, and deliberately not Pipeline.sample().
#
# Every optional field is populated on at least one project, and the rows vary
# in the dimensions the two known regressions lived in:
#
#   distress  deep AND severe both present, so the subset relation is visible
#   units     PRJ-D03 supplies a genuine 0; PRJ-D01/02/05/07/08 supply nothing;
#             the rest supply a positive count — all three states in one table
#   tract     11-digit GEOIDs, which must never acquire a thousands separator
#   awards    three prior awards, whose YEARS are the identifier column that
#             regressed
#
# NO DATE IN THIS FIXTURE MAY EVER EQUAL THE DAY THE SUITE RUNS. THAT IS A
# CONSTRAINT, NOT A COINCIDENCE — DO NOT ADD A ROW DATED NEAR TODAY.
#
# The project dates are in 2099/2100 for that reason and no other. They are
# deliberately absurd: a closing target seventy-odd years out is unmistakably
# synthetic, and it cannot collide with `date.today()` in any working lifetime.
# The CDE's certification_date and its three award years are historical, which
# is the same protection running the other way.
#
# The first version of this fixture dated PRJ-D08's closing target 2026-08-15,
# the day the baseline happened to be generated. The normaliser then replaced
# today's value blindly wherever it appeared, so that CELL normalised to a
# placeholder on the 15th and rendered literally on the 16th, and two of the
# four baselines went red for a reason with nothing to do with the document.
#
# The normaliser no longer looks at the value, so a collision is survivable —
# but a fixture date equal to the run date also makes the committed baseline
# ambiguous to read, and belt and braces is the right number of belts here.
# test_no_fixture_date_could_be_mistaken_for_the_run_date enforces it.
#
# Ordering constraint that must be preserved: validation/consistency_check
# requires construction_start <= operations_start per project.
# ---------------------------------------------------------------------------

def _cde() -> CDEProfile:
    return CDEProfile(
        name="Great Lakes Regional Capital CDE, LLC",
        cde_id="CDE-2016-0777",
        certification_date="2016-04-11",
        mission=(
            "Deploy New Markets Tax Credit capital into severely distressed "
            "census tracts across the eastern Great Lakes, with a standing "
            "preference for operating businesses that hire locally."
        ),
        target_markets=["Michigan", "Ohio", "Pennsylvania", "New York"],
        prior_awards=[
            {"year": 2019, "amount": 45_000_000, "deployment_status": "fully_deployed"},
            {"year": 2021, "amount": 60_000_000, "deployment_status": "fully_deployed"},
            {"year": 2023, "amount": 55_000_000, "deployment_status": "partially_deployed"},
        ],
        contact={
            "name": "Dana Okonkwo",
            "email": "dokonkwo@greatlakesregional.example.org",
            "phone": "216-555-0142",
            "title": "Chief Lending Officer",
        },
        governance={
            "board_members": 11,
            "community_representatives": 5,
            "advisory_board_members": 9,
            "board_meeting_frequency": "Bi-monthly",
        },
        website="https://greatlakesregional.example.org",
        extra={
            "years_operating": 9,
            "prior_award_count": 3,
            "has_prior_reporting_issues": False,
            "cde_type": "nonprofit_controlled",
        },
    )


# id, name, city, state, sector, ptype, cost, qei, jobs, retained,
# units, sqft, tract, distress, eligible
_SPEC = [
    ("PRJ-D01", "Erie Shoreline Health Pavilion", "Erie", "PA",
     "healthcare", "real_estate", 18_400_000, 13_000_000, 96, 34,
     None, 74_500.0, "42049000300", "deep", True),
    ("PRJ-D02", "Toledo Advanced Components Plant", "Toledo", "OH",
     "small_business", "operating_business", 24_750_000, 17_500_000, 143, 61,
     None, 210_000.0, "39095002400", "severe", True),
    ("PRJ-D03", "Buffalo East Side Grocery Anchor", "Buffalo", "NY",
     "other", "operating_business", 9_300_000, 6_600_000, 52, 18,
     0, 31_200.0, "36029006702", "severe", True),
    ("PRJ-D04", "Flint Riverfront Mixed-Income Lofts", "Flint", "MI",
     "affordable_housing", "real_estate", 31_200_000, 22_000_000, 27, 6,
     118, 165_000.0, "26049000900", "deep", True),
    ("PRJ-D05", "Youngstown Workforce Training Institute", "Youngstown", "OH",
     "education", "real_estate", 12_650_000, 8_900_000, 41, 15,
     None, 44_800.0, "39099810900", "lic", True),
    ("PRJ-D06", "Saginaw Community Wellness Center", "Saginaw", "MI",
     "community_facility", "real_estate", 7_450_000, 5_200_000, 33, 11,
     24, 28_300.0, "26145000200", "severe", True),
    ("PRJ-D07", "Scranton Solar Manufacturing Line", "Scranton", "PA",
     "clean_energy", "operating_business", 15_900_000, 11_200_000, 78, 29,
     None, 96_400.0, "42069100900", "deep", True),
    ("PRJ-D08", "Rochester Main Street Small Business Fund", "Rochester", "NY",
     "small_business", "operating_business", 5_800_000, 4_100_000, 62, 24,
     None, 12_700.0, "36055004400", "lic", True),
]


def _pipeline() -> Pipeline:
    projects = []
    for i, row in enumerate(_SPEC):
        (pid, name, city, state, sector, ptype, cost, qei, jobs, retained,
         units, sqft, tract, distress, eligible) = row
        p = PipelineProject(
            project_id=pid,
            project_name=name,
            qalicb_name=f"{name} QALICB, LLC",
            address=f"{1100 + 25 * i} Commerce Street",
            city=city,
            state=state,
            sector=sector,
            project_type=ptype,
            total_project_cost=float(cost),
            qei_request=float(qei),
            qlici_amount=float(qei),
            expected_jobs_created=jobs,
            expected_jobs_retained=retained,
            expected_units_built=units,
            expected_sq_ft=sqft,
            # Fixed, not derived from today — a date that moves would make
            # every run a diff. 2099/2100 so that none of them can EVER equal
            # the run date; see the fixture header. Ordering preserved:
            # construction_start <= operations_start, which consistency_check
            # requires, and "2099-.." sorts before "2100-.." as a string.
            closing_target_date=f"2099-{(i % 12) + 1:02d}-15",
            construction_start=f"2099-{(i % 12) + 1:02d}-28",
            operations_start=f"2100-{(i % 12) + 1:02d}-01",
        )
        p.census_tract = tract
        p.is_nmtc_eligible = eligible
        p.distress_level = distress
        p.is_native_area = (i == 5)
        p.is_high_migration_rural = (i == 6)
        p.is_opportunity_zone = (i % 3 == 0)
        p.is_us_territory = False
        p.is_persistent_poverty = (i % 4 == 0)
        p.is_below_market_rate = (i % 2 == 0)
        p.is_unrelated_entity = True
        p.geocode_success = True
        projects.append(p)
    pipeline = Pipeline(projects)
    # Same reasoning as tests/test_invariant_output.py: these rows carry
    # verified-shaped data on purpose, so the adapter's pre-enriched branch
    # must not route the whole document into the degraded banner and make the
    # baseline a picture of the failure mode instead of the document.
    pipeline.eligibility_data_status = "ok"
    return pipeline


REQUESTED_ALLOCATION = 70_000_000.0
APPLICATION_ROUND = "CY 2026"


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------

def _extract(fmt: str, path: str) -> str:
    if fmt == "markdown":
        with open(path, encoding="utf-8") as fh:
            return fh.read()

    if fmt == "word":
        from docx import Document
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(f"P|{p.style.name}|{p.text}")
        for t_i, table in enumerate(doc.tables):
            for r_i, row in enumerate(table.rows):
                cells = [c.text.replace("\n", "\\n") for c in row.cells]
                parts.append(f"T{t_i}|R{r_i}|" + "|".join(cells))
        for section in doc.sections:
            for p in section.header.paragraphs:
                if p.text.strip():
                    parts.append(f"HDR|{p.text}")
            for p in section.footer.paragraphs:
                if p.text.strip():
                    parts.append(f"FTR|{p.text}")
        return "\n".join(parts)

    if fmt == "excel":
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        parts = []
        for ws in wb.worksheets:
            parts.append(f"@@SHEET {ws.title}")
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    # THE NUMBER FORMAT IS PART OF THE OUTPUT. openpyxl's value
                    # read returns 2019 whether the cell displays "2019" or
                    # "2,019", so a value-only projection could not see B-3 at
                    # all — the defect that motivated this file.
                    parts.append(
                        f"{ws.title}!{cell.coordinate}|"
                        f"{type(cell.value).__name__}|"
                        f"fmt={cell.number_format}|{cell.value}"
                    )
        return "\n".join(parts)

    if fmt == "pdf":
        from pypdf import PdfReader
        pages = []
        for i, page in enumerate(PdfReader(path).pages):
            pages.append(f"@@PAGE {i + 1}")
            pages.append(page.extract_text() or "")
        return "\n".join(pages)

    raise AssertionError(f"unknown format {fmt}")


_ISO_TS = re.compile(r"\d{4}-\d{2}-\d{2}[T ]\d{2}:\d{2}:\d{2}(?:\.\d+)?")

# Either form a renderer stamps the run date in.
_ANY_DATE = (
    r"(?:\d{4}-\d{2}-\d{2}"
    r"|(?:January|February|March|April|May|June|July|August|September"
    r"|October|November|December) \d{1,2}, \d{4})"
)

# ---------------------------------------------------------------------------
# THE RUN-DATE STAMP, MATCHED BY WHERE IT SITS — NOT BY WHAT IT SAYS.
#
# This was `text.replace(date.today().isoformat(), "<TODAY>")`: a blind
# value replace across the whole document. It made the SUITE'S RESULT DEPEND ON
# THE CALENDAR.
#
# The fixture's PRJ-D08 carried closing_target_date "2026-08-15". The baseline
# was generated on 2026-08-15, so that literal matched today and normalised to
# "<TODAY>"; on 2026-08-16 the same cell rendered literally and the markdown
# and excel baselines mismatched. Two formats red, for a reason with nothing to
# do with what the gate exists to watch — which is the recorded reason gates
# get bypassed, and the exact inverse of every defect this release closed.
#
# So the value is never consulted. Each pattern below is anchored to the words
# a renderer puts BESIDE its stamp, and names the emission site it covers. A
# fixture date can now equal the run date without being touched, because
# nothing looks at what the date says.
#
# Coverage is asserted, not assumed: test_every_generation_stamp_site_is_
# anchored walks the renderers for `date.today()` call sites and fails if one
# is not matched by a pattern here.
# ---------------------------------------------------------------------------
_STAMP_ON_LINE = re.compile(
    r"(?P<lead>"
    r"Prepared:\s*"                                   # markdown:112, excel:171
    r"|Generated\s+"                                  # excel:369, excel:437
    r"|nmtc-application-builder v(?:[\d.]+|<VERSION>)\*{0,2}\s+on\s+"  # md:246,
    r"|Preparation Date:\|"                           # word:178  word:518,
    r")"                                              #            pdf:715
    r"(?P<date>" + _ANY_DATE + r")"
)

# pdf:468 is the one stamp with no lead word ON ITS OWN LINE. The PDF cover
# table extracts as two lines — the label, then the date alone — so it is
# anchored to the PRECEDING line instead. A bare date is normalised only
# directly under this label, never on its own.
_PDF_STAMP_LABEL = "Preparation Date:"
_BARE_DATE = re.compile(r"^\s*" + _ANY_DATE + r"\s*$")

# ---------------------------------------------------------------------------
# THE PACKAGE VERSION IS ENVIRONMENT STATE, NOT RENDERED CONTENT.
#
# The renderers stamp `nmtcapp.__version__`, which resolves through
# importlib.metadata — i.e. from the INSTALLED DISTRIBUTION'S metadata, not
# from the code under test. Measured on one interpreter, changing only the
# working directory:
#
#     run from the repo (carries nmtc_application_builder.egg-info)  v1.2.1
#     run from a clean clone (egg-info is gitignored)                v1.2.0
#
# So a clean clone of this branch failed three baselines — markdown, word and
# pdf — for a reason that has nothing to do with the document. Same class as
# the run-date defect this file was just fixed for: the gate's result
# depending on something other than its subject.
#
# The RENDERER is right; a CDE's filing should say which build produced it.
# The GATE must not depend on it, so the version is normalised the same way
# the run date is: anchored to the package name, so it cannot eat anything
# else. Version correctness is gated by tests/test_version.py, which is where
# it belongs.
# ---------------------------------------------------------------------------
_VERSION_STAMP = re.compile(
    r"(?P<lead>nmtc-application-builder v)\d+\.\d+\.\d+(?:\.\w+)?"
)

VERSION_STAMP = "<VERSION>"

RUN_DATE = "<RUNDATE>"


def _normalise(text: str, outdir: str) -> str:
    """Erase ONLY what is genuinely non-deterministic.

    Deliberately narrow. Every normalisation is a line this gate stops
    watching, so the list is temp paths, ISO timestamps, the run-date stamp
    and the package version — and nothing else. The last two are properties
    of the RUN and the INSTALL, not of the document. The fixture's own dates (certification_date, closing
    target, construction start) are OUTPUT and must survive verbatim; that is
    what the anchoring above protects.
    """
    text = text.replace(outdir, "<OUTDIR>")
    text = re.sub(r"/(?:private/)?(?:tmp|var)/[^\s\"']*", "<TMPPATH>", text)
    text = _ISO_TS.sub("<TIMESTAMP>", text)
    text = _VERSION_STAMP.sub(lambda m: m.group("lead") + VERSION_STAMP, text)

    out = []
    previous = ""
    for line in text.split("\n"):
        if previous.strip() == _PDF_STAMP_LABEL and _BARE_DATE.match(line):
            out.append(RUN_DATE)
        else:
            out.append(_STAMP_ON_LINE.sub(
                lambda m: m.group("lead") + RUN_DATE, line
            ))
        previous = line
    return "\n".join(out)


def _render_projections(outdir: str) -> dict:
    """Render the fixture into outdir and return {format: normalised text}."""
    app = Application(
        cde=_cde(),
        requested_allocation=REQUESTED_ALLOCATION,
        application_round=APPLICATION_ROUND,
    )
    app.add_pipeline(_pipeline())
    paths = app.generate(outdir, formats=list(FORMATS))

    assert set(paths) == set(FORMATS), (
        f"rendered {sorted(paths)}, expected all of {sorted(FORMATS)}. A "
        "format that silently does not render is a format this gate is not "
        "watching, and the missing one is where the next regression lands."
    )

    projected = {}
    for fmt in FORMATS:
        assert os.path.exists(paths[fmt]), f"{fmt} was not written"
        text = _normalise(_extract(fmt, paths[fmt]), outdir)
        assert text.strip(), (
            f"{fmt} extracted as empty text. The diff would then be empty and "
            "this gate would report no change on a document that renders "
            "nothing."
        )
        projected[fmt] = text.rstrip("\n") + "\n"
    return projected


# The four modules that stamp the run date. Frozen together, because a stamp
# left un-frozen in one of them is exactly what the date-independence test
# below is looking for.
_STAMPING_MODULES = (
    "nmtcapp.renderers.markdown_builder",
    "nmtcapp.renderers.word_builder",
    "nmtcapp.renderers.excel_builder",
    "nmtcapp.renderers.pdf_builder",
)


def _freeze_today(monkeypatch, iso: str) -> None:
    """Make every renderer believe today is `iso`.

    Each module does `from datetime import date`, so each holds its own
    binding and each must be patched. No dependency on freezegun or on a
    faked system clock — the point is that this runs anywhere, including CI.
    """
    import importlib

    year, month, day = (int(part) for part in iso.split("-"))

    class _Frozen(datetime.date):
        @classmethod
        def today(cls):
            return datetime.date(year, month, day)

    for name in _STAMPING_MODULES:
        module = importlib.import_module(name)
        assert hasattr(module, "date"), (
            f"{name} no longer imports `date`; the freeze would silently miss "
            "whatever stamps the run date there now"
        )
        monkeypatch.setattr(module, "date", _Frozen)


@pytest.fixture(scope="module")
def rendered(tmp_path_factory) -> dict:
    """{format: normalised text projection} for the one fixed fixture."""
    return _render_projections(str(tmp_path_factory.mktemp("baseline")))


# ---------------------------------------------------------------------------
# Fail-closed structural checks
# ---------------------------------------------------------------------------

def test_the_fixture_populates_every_field_the_defects_lived_in():
    """A fixture that stops exercising a field stops guarding it.

    Each assertion here names a defect the baseline exists to catch. If the
    fixture is edited and one of these stops holding, the corresponding class
    goes unwatched while the diff still passes.
    """
    pipeline = _pipeline()
    projects = list(pipeline)

    levels = {p.distress_level for p in projects}
    assert {"deep", "severe"} <= levels, (
        f"the fixture must contain BOTH deep and severe projects (B-1): {levels}"
    )

    units = [p.expected_units_built for p in projects]
    assert None in units, "no project leaves units unsupplied (B-2)"
    assert 0 in units, "no project supplies a genuine zero units (B-2)"
    assert any(u for u in units if u), "no project supplies positive units (B-2)"

    assert _cde().prior_awards, "no prior awards, so no Award Year column (B-3)"
    assert all(len(str(p.census_tract)) == 11 for p in projects), (
        "a census tract is not an 11-digit GEOID; the identifier-column class "
        "would go unexercised (B-3)"
    )


def test_a_baseline_exists_for_every_format():
    """A missing baseline must ERROR, not be created silently on the next run.

    A gate that writes its own expected output the first time it fails is a
    gate that can never fail. Regenerating is a deliberate act — see
    ``test_the_rendered_output_matches_the_reviewed_baseline``'s message.
    """
    assert os.path.isdir(BASELINE_DIR), (
        f"{BASELINE_DIR} does not exist. It is committed on purpose: the "
        "baseline IS the review artifact."
    )
    for fmt in FORMATS:
        path = os.path.join(BASELINE_DIR, f"{fmt}.txt")
        assert os.path.exists(path), f"no committed baseline for {fmt}: {path}"
        with open(path, encoding="utf-8") as fh:
            lines = fh.read().splitlines()
        assert len(lines) >= _MIN_BASELINE_LINES[fmt], (
            f"the {fmt} baseline holds {len(lines)} lines, below the "
            f"{_MIN_BASELINE_LINES[fmt]} floor. A truncated baseline is a "
            "diff that compares almost nothing."
        )


# ---------------------------------------------------------------------------
# THE RUN DATE MUST NOT REACH THE COMPARISON (FIX-2 follow-up)
#
# The first version of this gate normalised the stamp with
#
#     text.replace(datetime.date.today().isoformat(), "<TODAY>")
#
# and the fixture dated PRJ-D08's closing target 2026-08-15. The baseline was
# generated on 2026-08-15, so that CELL matched today and normalised away; on
# 2026-08-16 it rendered literally and markdown and excel went red. The suite's
# result depended on the calendar — a gate failing for a reason unrelated to
# its subject, which is the recorded reason gates get bypassed.
#
# These three tests exist so that cannot recur, and so that nobody has to take
# "it passed on the day I wrote it" as evidence.
# ---------------------------------------------------------------------------

# Three distinct dates, and the THIRD IS ONE THE FIXTURE ITSELF CONTAINS
# (PRJ-D08's closing_target_date). That is the case the old normaliser got
# wrong, so it is the case this test must cover explicitly rather than by luck.
_PROOF_DATES = ("2026-08-16", "2027-03-01", "2099-08-15")


def test_the_projection_is_independent_of_the_run_date(tmp_path, monkeypatch):
    """Render under three different "todays" and require identical output.

    THE INVARIANT, STATED DIRECTLY: nothing about the day the suite runs may
    survive into the text this gate compares. Freezing the four stamping
    modules and rendering three times is the only way to establish that; a
    gate that has only ever been run on one day has been tested on one day.
    """
    projections = {}
    for iso in _PROOF_DATES:
        with monkeypatch.context() as frozen:
            _freeze_today(frozen, iso)
            out = tmp_path / iso
            out.mkdir()
            projections[iso] = _render_projections(str(out))

    reference_iso = _PROOF_DATES[0]
    reference = projections[reference_iso]
    for iso in _PROOF_DATES[1:]:
        for fmt in FORMATS:
            diff = list(difflib.unified_diff(
                reference[fmt].splitlines(), projections[iso][fmt].splitlines(),
                fromfile=f"{fmt} rendered on {reference_iso}",
                tofile=f"{fmt} rendered on {iso}",
                lineterm="", n=1,
            ))
            assert not diff, (
                f"the {fmt} projection changed when only the RUN DATE changed "
                f"({reference_iso} -> {iso}). Either a generation stamp is not "
                "anchored in _STAMP_ON_LINE, or a fixture value collided with "
                "one of these dates and was normalised away. Both make this "
                "gate's result depend on the calendar.\n\n" + "\n".join(diff)
            )

    # And it must not pass by normalising everything into oblivion.
    for fmt in FORMATS:
        assert RUN_DATE in reference[fmt], (
            f"no run-date stamp was found in the {fmt} projection at all. "
            "Three identical renders prove nothing if the normaliser has "
            "erased the stamps along with everything else."
        )


# The fixture's dates must sit OUTSIDE any window a run could plausibly fall
# in. Stated as a window rather than as "!= today" on purpose: see the test.
_RUN_WINDOW = (2026, 2098)


def test_no_fixture_date_could_be_mistaken_for_the_run_date():
    """The fixture's own constraint, enforced rather than commented.

    NOTE WHAT THIS TEST DOES NOT DO: it never calls ``date.today()``.

    The obvious form of this check is ``assert no fixture date == today``, and
    that form is itself the defect being fixed — it passes on 364 days and
    fails on one, so the suite's result would again depend on the calendar.
    Measured, before this was rewritten: with the fixture dated 2099-08-15 and
    the clock faked to 2099-08-15, the baseline comparison PASSED (the
    normaliser no longer reads the value) and this check was the only thing
    red. A guard that manufactures the failure mode it is guarding against is
    worse than no guard.

    So the property asserted is structural and calendar-free: every date in the
    fixture is either clearly historical or clearly synthetic — outside
    _RUN_WINDOW, the span a run could plausibly fall in. A row added with a
    near-today date fails here on every day, including the day it is added,
    which is when it is cheap to fix.
    """
    dates = {"cde.certification_date": _cde().certification_date}
    for award in _cde().prior_awards:
        dates[f"award.{award['year']}"] = f"{award['year']}-01-01"
    for project in _pipeline():
        for field in ("closing_target_date", "construction_start",
                      "operations_start"):
            dates[f"{project.project_id}.{field}"] = getattr(project, field)

    assert len(dates) > 20, (
        f"only {len(dates)} fixture dates were collected; this check is not "
        "reaching the fixture and would pass having examined almost nothing"
    )

    low, high = _RUN_WINDOW
    inside = sorted(
        f"{label}={value}" for label, value in dates.items()
        if low <= int(str(value)[:4]) <= high
    )
    assert not inside, (
        f"{len(inside)} fixture date(s) fall inside the plausible run window "
        f"{low}-{high}: {inside}\n\n"
        "The fixture header states this constraint. Dates in it are 2099/2100 "
        "(projects) or historical (the CDE's certification and award years) "
        "SO THAT none of them can be confused with the day the suite runs — "
        "in the committed baseline a reader has to be able to tell a fixture "
        "value from a normalised stamp at a glance. Move the value."
    )

    # The ordering consistency_check enforces, asserted here so a date edit
    # cannot quietly produce a document that fails validation.
    for project in _pipeline():
        assert project.construction_start <= project.operations_start, (
            f"{project.project_id}: construction_start "
            f"{project.construction_start} is after operations_start "
            f"{project.operations_start}; consistency_check rejects that"
        )


def test_the_projection_is_independent_of_the_installed_version(tmp_path, monkeypatch):
    """The clean-clone failure, reproduced and pinned.

    `nmtcapp.__version__` resolves through importlib.metadata, so it is a fact
    about the INSTALLED DISTRIBUTION rather than about the code under test. On
    one interpreter, changing only the working directory:

        run from the repo (carries an egg-info)   v1.2.1
        run from a clean clone (egg-info ignored)  v1.2.0

    Three baselines went red on a clean clone for that reason alone. Same class
    as the run-date defect — the gate's result depending on something other
    than its subject — and it is fixed the same way, by normalising the stamp
    rather than by pinning an environment fact into a committed file.
    """
    import nmtcapp

    projections = {}
    for version in ("1.2.0", "1.2.1", "99.0.0.dev"):
        monkeypatch.setattr(nmtcapp, "__version__", version)
        out = tmp_path / version
        out.mkdir()
        projections[version] = _render_projections(str(out))

    reference_version = "1.2.0"
    reference = projections[reference_version]
    for version, projection in projections.items():
        if version == reference_version:
            continue
        for fmt in FORMATS:
            diff = list(difflib.unified_diff(
                reference[fmt].splitlines(), projection[fmt].splitlines(),
                fromfile=f"{fmt} built as v{reference_version}",
                tofile=f"{fmt} built as v{version}",
                lineterm="", n=1,
            ))
            assert not diff, (
                f"the {fmt} projection changed when only the INSTALLED "
                f"VERSION changed (v{reference_version} -> v{version}). A "
                "version stamp is not anchored in _VERSION_STAMP, so a clean "
                "clone — or any install whose metadata differs from this "
                "tree — fails this gate for a reason that has nothing to do "
                "with the document.\n\n" + "\n".join(diff)
            )

    # And not by erasing the stamps along with everything else.
    for fmt in ("markdown", "word", "pdf"):
        assert VERSION_STAMP in reference[fmt], (
            f"no version stamp was found in the {fmt} projection. Three "
            "identical renders prove nothing if the normaliser removed them."
        )


def test_every_generation_stamp_site_is_anchored():
    """A new `date.today()` in a renderer must not silently leak into the diff.

    test_the_projection_is_independent_of_the_run_date catches any stamp that
    RENDERS under this fixture. It cannot catch one on a branch the fixture
    does not take. So the renderers are also read directly: every
    ``date.today()`` call site must carry one of the anchor phrases
    _STAMP_ON_LINE keys on, within the call's own line or the one above it
    (pdf_builder puts the lead text on the preceding line).
    """
    import importlib

    anchors = ("Prepared:", "Generated ", "Generated by",
               "nmtc-application-builder", "Preparation Date:")

    unanchored = []
    sites = 0
    for name in _STAMPING_MODULES:
        module = importlib.import_module(name)
        path = module.__file__
        with open(path, encoding="utf-8") as fh:
            lines = fh.read().splitlines()
        for i, line in enumerate(lines):
            if "date.today()" not in line:
                continue
            sites += 1
            window = (lines[i - 1] if i else "") + "\n" + line
            if not any(a in window for a in anchors):
                unanchored.append(
                    f"  {os.path.basename(path)}:{i + 1}: {line.strip()[:90]}"
                )

    assert sites, (
        "no `date.today()` call sites were found in any renderer. Either the "
        "stamps moved somewhere this scan cannot see — in which case the "
        "normaliser's anchors are guarding nothing — or the scan is broken."
    )
    assert not unanchored, (
        f"{len(unanchored)} of {sites} run-date stamp(s) carry none of the "
        f"anchor phrases {anchors}:\n" + "\n".join(unanchored)
        + "\n\nAn un-anchored stamp is a moving value in the committed "
        "baseline: it normalises on the day the baseline is written and "
        "renders literally every day after. Give it a lead word and add the "
        "pattern to _STAMP_ON_LINE."
    )


# ---------------------------------------------------------------------------
# The gate
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("fmt", FORMATS)
def test_the_rendered_output_matches_the_reviewed_baseline(rendered, fmt):
    """Every changed rendered line, in front of a human, in the same commit."""
    path = os.path.join(BASELINE_DIR, f"{fmt}.txt")
    with open(path, encoding="utf-8") as fh:
        expected = fh.read().splitlines()
    actual = rendered[fmt].splitlines()

    diff = list(difflib.unified_diff(
        expected, actual,
        fromfile=f"baseline/{fmt}.txt", tofile=f"rendered/{fmt}.txt",
        lineterm="", n=2,
    ))
    changed = [ln for ln in diff
               if ln.startswith(("+", "-")) and not ln.startswith(("+++", "---"))]

    assert not diff, (
        f"{len(changed)} rendered line(s) changed in the {fmt} output and the "
        "reviewed baseline does not record the change.\n\n"
        "EVERY LINE BELOW IS EITHER AN INTENDED FIX OR A REGRESSION NOBODY "
        "CHOSE. Two consecutive releases shipped a defect created by that "
        "release's own fix — a document stating 0.0% and 40% severe distress, "
        "and prior-award years rendering as 2,019 — and every other gate in "
        "this package stayed green through both, because no gate asks what "
        "CHANGED.\n\n"
        "Read the diff. If the change is intended, regenerate the baseline in "
        "the SAME commit so the rendered change appears in the pull request "
        "as a diff of the document:\n\n"
        "    python -m tests.regen_rendered_baseline\n\n"
        + "\n".join(diff)
    )
