"""THE ROUND GATE: the hash pin cannot fail on staleness, so something else must.

THE GATE THAT CANNOT FAIL

``renderers/_question_25`` pins the CY 2024-2025 Application's SHA-256. That pin
is CORRECT -- re-downloaded and re-verified this round, byte count and all 64
hex characters identical. It is also, on its own, the nineteenth recorded
instance in this project of a gate that cannot fail, and the FIRST one caught
before it fired rather than after.

The reason is that it answers the wrong question:

    the hash asks   "is this the document we read?"          -> yes, forever
    nobody asks     "is this the round the CDE FILES?"       -> no, since
                                                                 Dec 2025

When CY 2026 publishes, the hash still matches, the suite is still green, and
every citation in the package is stale. A hash pins an ARTIFACT against
corruption. It says nothing about that artifact's RELEVANCE, and relevance is
the property that decays.

ATTACKING THE OBVIOUS FIX

The candidate design was "a round label and a published-status assertion pinned
beside the hash". Half right. ``_round_provenance`` now carries
``CITED_ROUND``, ``CITED_ROUND_STATUS`` and ``UPCOMING_MATERIALS_PUBLISHED``,
which makes the staleness a FACT IN THE CODE rather than an omission -- a real
improvement, because a reader of the module now sees it.

But as a GATE it is worthless on its own, and saying so is the point:

    ``UPCOMING_MATERIALS_PUBLISHED = False`` is a sentence somebody typed.
    Nothing flips it. A test asserting the code agrees with itself is the
    tautology ``test_version_sync`` already is (see 1.5.0 S6), and a dated
    assertion nobody re-reads is just another gate that cannot fail.

SO: CAN A TEST DISTINGUISH "CY 2026 HAS NOT PUBLISHED" FROM "NOBODY HAS LOOKED
SINCE AUGUST"?

**Offline, no. Categorically.** Publication is a fact about the world. An
offline test can only read what somebody wrote down, so it can only ever
re-report the writing. No amount of cleverness gets around that, and a gate
that implied otherwise would be worse than none.

What an offline test CAN do is fail on the SECOND condition. That is what this
module does, and it is the honest half:

  ``test_the_round_claim_has_not_expired`` fails when today is past
  ``RECHECK_AFTER``. It detects STALE LOOKING, not stale facts. It cannot tell
  you CY 2026 published yesterday. It CAN tell you nobody has checked in three
  months, which is the failure mode that actually produced this defect -- the
  CY 2024-2025 citation did not go wrong because anyone decided wrongly, it
  went wrong because the world moved and no one was scheduled to look.

  IT DEGRADES TO RITUAL IF BUMPED ROTE, and pretending otherwise would be the
  same self-flattery this suite keeps auditing out. Two things make a rote bump
  less likely and neither makes it impossible: the date lives in the same file
  as the round claim, so the bumper is looking at the claim; and the failure
  message names the two URLs and the specific things to look for, so checking
  properly is cheaper than inventing a reason not to.

  ``test_live_cdfi_fund_check`` is the half that CAN answer the real question:
  it fetches the CDFI Fund program page and asserts CY 2026 materials are still
  absent. It is NOT A GATE -- it is opt-in (``-m network``) and skipped by
  default, because a suite whose greenness depends on a federal website being
  up is a suite that goes red for reasons having nothing to do with this
  repository. A skipped-by-default test cannot fail in CI, and calling it a
  gate would be the vacuity this file exists to name. It is a TOOL, run when
  the expiry fires.

Together: the expiry says WHEN to look, the live check answers WHAT IS TRUE,
and neither pretends to be the other.
"""
from __future__ import annotations

import datetime as _dt
import os
import re

import pytest

from nmtcapp.renderers import _round_provenance as rp

_REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def _iso(value: str) -> _dt.date:
    return _dt.date(*(int(part) for part in value.split("-")))


#: Every format the application renders to. The gate below asserts the round
#: provenance reaches ALL of them, and the tuple is written out rather than
#: imported so that a format being dropped from the renderer cannot silently
#: shrink what this gate checks.
_ALL_FORMATS = ("markdown", "word", "excel", "pdf")

#: Phrases that must survive into every rendered artifact. Each one is a
#: DISTINCT fact a reader loses if it is missing, not a restatement:
#:   - which round, and that it is over
#:   - that the round a CDE will actually file has no materials yet
#:   - the certification deadline, which is external and hard
#:   - the third obligation, which binds prior Allocatees (F7)
_PROVENANCE_FACTS = (
    ("closed and awarded", "which round this encodes, and that it is closed"),
    ("NOT YET PUBLISHED", "that the round a CDE will file has no materials yet"),
    ("August 31, 2026", "the AMIS CDE certification deadline"),
    ("Subsidiary CDE", "the prior-Allocatee Subsidiary CDE obligation"),
)


def _render_all_formats(tmp_path) -> dict:
    """Render the sample application to every format; return {fmt: text}."""
    from nmtcapp.core.application import Application
    from nmtcapp.core.cde import CDEProfile
    from nmtcapp.core.pipeline import Pipeline

    app = Application(cde=CDEProfile.sample(), requested_allocation=65_000_000)
    app.add_pipeline(Pipeline.sample(n=20))
    paths = app.generate(str(tmp_path), formats=list(_ALL_FORMATS))

    # FAILS CLOSED. A format that does not render is a format this gate would
    # otherwise pass by not looking at.
    assert set(paths) == set(_ALL_FORMATS), (
        f"rendered {sorted(paths)}, expected {sorted(_ALL_FORMATS)} — a format "
        "that silently does not render is a format this gate is not checking"
    )

    out = {}
    for fmt, path in paths.items():
        if fmt == "markdown":
            with open(path, encoding="utf-8") as fh:
                text = fh.read()
        elif fmt == "word":
            from docx import Document
            doc = Document(path)
            parts = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            text = "\n".join(parts)
        elif fmt == "excel":
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True)
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    parts.extend(str(v) for v in row if v is not None)
            text = "\n".join(parts)
        else:
            from pypdf import PdfReader
            text = "\n".join((page.extract_text() or "") for page in PdfReader(path).pages)

        assert text.strip(), f"{fmt} extracted as empty text"
        # PDF and Word wrap; compare on collapsed whitespace so a line break
        # inside a phrase is not read as the phrase being absent.
        out[fmt] = re.sub(r"\s+", " ", text)
    return out


def test_the_round_provenance_reaches_all_four_formats(tmp_path):
    """The disclosure must render in EVERY format, not three of four.

    THIS GATE IS THE DELIVERABLE, NOT THE EXCEL BLOCK IT CHECKS (1.5.0 B1).

    Through 1.5.0 the note reached markdown, Word and PDF. The workbook
    carried none of it -- and carried the citation anyway, in the present
    tense and in the Fund's voice, at row 4 of the 'Q25 Basis Note' sheet:
    "Question 25 of the CY 2024-2025 NMTC Allocation Application (printed
    pp. 38-41) sets both commitments." A CDE who opened only the workbook got
    a federal citation with no notice that the round had closed on
    29 Jan 2025, that CY 2026 was unpublished, or that a hard external
    certification deadline fell on 31 Aug 2026.

    Excel is the format most likely to be circulated internally and pasted
    from, so it was the worst of the four to leave silent. It was silent
    because the gates counted formats that HAD the note rather than formats
    that MUST. Adding the block without adding this test would fix the site
    and leave the class -- which is the shape this project has recorded
    repeatedly, and the reason the assertion below is parameterised over
    _ALL_FORMATS rather than naming the workbook.
    """
    rendered = _render_all_formats(tmp_path)

    missing = [
        (fmt, phrase, why)
        for fmt, text in rendered.items()
        for phrase, why in _PROVENANCE_FACTS
        if phrase not in text
    ]
    assert not missing, (
        "round provenance is missing from rendered output:\n\n"
        + "\n".join(
            f"  {fmt:<9} lacks {phrase!r} — the reader loses {why}"
            for fmt, phrase, why in missing
        )
        + "\n\nEvery format that cites a round must say which round it is and "
        "that it is closed. Render the note from "
        "_round_provenance.round_provenance_paragraphs(); do not retype it."
    )


def test_the_paragraph_view_is_exactly_the_note():
    """The two shapes of the note cannot drift, because one builds the other.

    Excel needs the note one paragraph per row (a merged cell has a 409-pt
    ceiling and the note is longer). The wrong way to get that is a second
    copy in excel_builder -- the exact shape that produced the 1.2.0 defect
    where a sentence was corrected in one file and left live in another. So
    the paragraphs are the definition and the note is their join, and this
    asserts it stays that way.
    """
    joined = " ".join(rp.round_provenance_paragraphs())
    assert joined == rp.round_provenance_note(), (
        "round_provenance_note() is no longer the join of "
        "round_provenance_paragraphs(). Excel renders the paragraphs and the "
        "other three formats render the note, so they have just diverged: one "
        "artifact now says something the others do not."
    )


# ---------------------------------------------------------------------------
# The offline half
# ---------------------------------------------------------------------------

def test_the_round_claim_has_not_expired():
    """FAILS ON TIME. The only offline failure available here, and it is real.

    This does not check whether CY 2026 published. It checks whether anybody
    has looked recently, which is a different and weaker claim -- and it is the
    one that failed. See this module's header.
    """
    today = _dt.date.today()
    recheck = _iso(rp.RECHECK_AFTER)
    assert today <= recheck, (
        f"the {rp.UPCOMING_ROUND} round claim expired on {rp.RECHECK_AFTER} "
        f"(today is {today.isoformat()}). Nobody has verified it since "
        f"{rp.LAST_VERIFIED}.\n\n"
        "THIS IS NOT A FAILING BUILD, IT IS A SCHEDULED RE-CHECK. Do it now:\n\n"
        f"  1. {rp.CY2026_ANNOUNCEMENT_URL}\n"
        f"     -- is there a newer release than the 12 Aug 2026 one?\n"
        f"  2. {rp.PROGRAM_PAGE_URL}\n"
        f"     -- does the timeline still show the CY 2024-2025 dates "
        "(Opening 19 Nov 2024 / Deadline 29 Jan 2025 / Announcement "
        "23 Dec 2025)? Is there a CY 2026 NOAA or Allocation Application "
        "link yet?\n\n"
        "Then EITHER set UPCOMING_MATERIALS_PUBLISHED = True and open the "
        "re-verification work in _round_provenance.RECHECK_ITEMS, OR bump "
        "LAST_VERIFIED and RECHECK_AFTER.\n\n"
        "Bumping the dates WITHOUT opening those two pages turns this into a "
        "gate that cannot fail, which is the exact thing it was built to "
        "replace. `pytest -m network tests/test_round_provenance.py` does the "
        "check for you."
    )


def test_the_recheck_horizon_is_after_the_verification():
    """A horizon before its own verification date is already expired."""
    assert _iso(rp.RECHECK_AFTER) > _iso(rp.LAST_VERIFIED), (
        f"RECHECK_AFTER ({rp.RECHECK_AFTER}) is not after LAST_VERIFIED "
        f"({rp.LAST_VERIFIED})."
    )


def test_the_horizon_is_not_pushed_out_of_reach():
    """A far-enough horizon is an abstention wearing a gate's clothes.

    MAX_HORIZON_DAYS is the difference between "re-check quarterly" and "this
    will fire after I have stopped working on it". Without a ceiling, the
    cheapest response to a red expiry is +5 years, and the gate is gone with
    nobody having decided to remove it.
    """
    span = (_iso(rp.RECHECK_AFTER) - _iso(rp.LAST_VERIFIED)).days
    assert span <= 180, (
        f"the re-check horizon is {span} days. An NMTC round opens, runs and "
        "closes inside that window, so a horizon this long cannot catch the "
        "transition it exists for. Keep it to a quarter or so; if the round "
        "genuinely has not moved, re-verifying costs two page loads."
    )


def test_the_application_hash_is_pinned_in_exactly_one_place():
    """ONE COPY of 64 hex characters.

    The hash was typed into two module docstrings, split across two source
    lines each. Nobody proofreads 64 hex characters, so a divergence would
    read as provenance while pointing at nothing -- and this package has
    already shipped three hand-typed copies of one sentence that agreed only
    by luck (``Q25_QEI_BASIS_CLAUSE``).
    """
    sha = rp.APPLICATION_SHA256
    assert len(sha) == 64 and re.fullmatch(r"[0-9a-f]{64}", sha), (
        f"APPLICATION_SHA256 is not 64 lowercase hex characters: {sha!r}"
    )

    # Search the package for the hash with any whitespace/newlines removed,
    # which is how the two docstring copies were written.
    offenders = []
    for dirpath, _dirs, names in os.walk(os.path.join(_REPO_ROOT, "nmtcapp")):
        for name in sorted(names):
            if not name.endswith(".py") or name == "_round_provenance.py":
                continue
            path = os.path.join(dirpath, name)
            with open(path, encoding="utf-8") as handle:
                squashed = re.sub(r"\s+", "", handle.read())
            if sha in squashed:
                offenders.append(os.path.relpath(path, _REPO_ROOT))

    assert not offenders, (
        f"the CY 2024-2025 Application hash is typed again in {offenders}. It "
        "lives in nmtcapp/renderers/_round_provenance.APPLICATION_SHA256 and "
        "must be READ from there, not retyped -- import it, or reference the "
        "constant by name in prose."
    )


def test_the_disclosure_states_both_directions():
    """The note must not read as a warning label.

    A disclosure scoped only around understatement produces a correct fact
    leading somewhere wrong: a CDE told the guidance is stale, and not told it
    is nonetheless the right thing to prepare against, prepares against
    nothing. That outcome is WORSE than the stale citation this replaces, so
    both halves are asserted rather than only the cautionary one.
    """
    note = rp.round_provenance_note()

    for phrase, why in (
        ("NOT YET PUBLISHED", "must say the CY 2026 materials do not exist"),
        ("most recent PUBLISHED Application",
         "must say what the cited round IS, not only what it is not"),
        ("re-verified", "must tell the CDE what to do, not just what is wrong"),
        ("nothing here is unreliable",
         "must say the cited instrument is still the right basis -- this is "
         "the overstating-uncertainty half, and it is the one a "
         "caution-shaped rewrite drops first"),
        ("August 31, 2026",
         "must carry the certification deadline, which is the only CY 2026 "
         "date a CDE can miss TODAY"),
    ):
        assert phrase in note, (
            f"round_provenance_note() no longer contains {phrase!r}: it {why}."
        )

    assert "CY 2026" in note and rp.CITED_ROUND in note


def test_the_note_names_every_recheck_item():
    """The re-check list is what makes this actionable rather than ominous."""
    note = rp.round_provenance_note()
    missing = [item for item in rp.RECHECK_ITEMS if item not in note]
    assert not missing, (
        f"round_provenance_note() drops {missing} from the re-check list. A "
        "CDE cannot act on an item the note does not name."
    )


def test_upcoming_materials_are_still_unpublished():
    """Pins the premise the rest of the package is written against.

    DELIBERATELY WEAK, AND LABELLED AS SUCH. This asserts the module agrees
    with itself. Flipping UPCOMING_MATERIALS_PUBLISHED to True should be a
    deliberate act that breaks this and sends the author to RECHECK_ITEMS --
    that is its whole value. It is NOT evidence about the world; see the
    module header on why no offline test can be.
    """
    assert rp.UPCOMING_MATERIALS_PUBLISHED is False, (
        "UPCOMING_MATERIALS_PUBLISHED is True, so the CY 2026 materials have "
        "been published and this package still encodes CY 2024-2025.\n\n"
        "Work the list in _round_provenance.RECHECK_ITEMS against the new "
        "documents, then update the citations and this test together."
    )


# ---------------------------------------------------------------------------
# The network half -- a TOOL, not a gate
# ---------------------------------------------------------------------------

@pytest.mark.network
def test_live_cdfi_fund_check():
    """Ask cdfifund.gov the question no offline test can answer.

    Run it when the expiry fires::

        pytest -m network tests/test_round_provenance.py

    SKIPPED BY DEFAULT, and that is a real limitation rather than a
    configuration detail: a skipped test cannot fail in CI, so this is not a
    gate. Making CI depend on a federal website being reachable would produce
    red builds that say nothing about this repository, and a flaky gate is one
    people learn to ignore -- which is a worse outcome than an opt-in tool
    people run deliberately.
    """
    import urllib.request

    request = urllib.request.Request(
        rp.PROGRAM_PAGE_URL, headers={"User-Agent": "Mozilla/5.0"}
    )
    try:
        with urllib.request.urlopen(request, timeout=45) as response:
            body = response.read().decode("utf-8", "replace")
    except Exception as exc:                      # pragma: no cover - network
        pytest.skip(f"cdfifund.gov unreachable ({exc}); no conclusion drawn")

    assert len(body) > 5_000, "program page returned too little to read"

    # The CY 2024-2025 timeline is what the page showed on LAST_VERIFIED. If it
    # is gone, the round has turned over.
    still_prior_round = "January 29, 2025" in body and "November 19, 2024" in body
    names_new_round = bool(
        re.search(r"CY\s*2026\s+.*Allocation Application", body)
        or re.search(r"2026[-_]\s*NMTC[_-]", body)
    )

    assert still_prior_round and not names_new_round, (
        "cdfifund.gov's NMTC program page has changed.\n\n"
        f"  CY 2024-2025 timeline still shown: {still_prior_round}\n"
        f"  page now names CY 2026 materials:  {names_new_round}\n\n"
        "The CY 2026 round may have opened. Work "
        "_round_provenance.RECHECK_ITEMS against the published materials, "
        "then update the citations, UPCOMING_MATERIALS_PUBLISHED, and the "
        "offline tests in this module."
    )
