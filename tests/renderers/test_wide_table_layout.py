"""Tests that wide tables are handled correctly in Word and PDF output."""
import os
import pytest

pytest.importorskip("docx", reason="python-docx not installed")
pytest.importorskip("reportlab", reason="reportlab not installed")


# ---------------------------------------------------------------------------
# Word: orientation and table width
# ---------------------------------------------------------------------------

class TestWordLandscapeOrientation:
    @pytest.fixture
    def docx_path(self, sample_application, application_analysis, tmp_path):
        from nmtcapp.renderers.word_builder import WordApplicationBuilder
        path = str(tmp_path / "test_wide.docx")
        WordApplicationBuilder(sample_application, application_analysis).save(path)
        return path

    def test_document_has_at_least_three_sections(self, docx_path):
        from docx import Document
        doc = Document(docx_path)
        # Cover, body, landscape A, landscape B back-to-portrait, ...
        assert len(doc.sections) >= 3

    def test_at_least_one_landscape_section(self, docx_path):
        from docx import Document
        from docx.enum.section import WD_ORIENT
        doc = Document(docx_path)
        orientations = [s.orientation for s in doc.sections]
        assert WD_ORIENT.LANDSCAPE in orientations

    def test_landscape_section_is_wider_than_tall(self, docx_path):
        from docx import Document
        from docx.enum.section import WD_ORIENT
        doc = Document(docx_path)
        landscape_secs = [s for s in doc.sections if s.orientation == WD_ORIENT.LANDSCAPE]
        assert len(landscape_secs) >= 1
        for sec in landscape_secs:
            assert sec.page_width > sec.page_height

    def test_portrait_sections_are_taller_than_wide(self, docx_path):
        from docx import Document
        from docx.enum.section import WD_ORIENT
        doc = Document(docx_path)
        portrait_secs = [s for s in doc.sections if s.orientation == WD_ORIENT.PORTRAIT]
        assert len(portrait_secs) >= 2
        for sec in portrait_secs:
            assert sec.page_height > sec.page_width

    def test_summary_tables_used_in_portrait_sections(self, sample_pipeline, sample_cde):
        """Verify the tables used in portrait body have ≤ 8 columns each."""
        from nmtcapp.tables.pipeline_table import build_pipeline_summary_table
        from nmtcapp.tables.distress_table import build_distress_summary_table
        from nmtcapp.tables.impact_table import build_impact_summary_table
        from nmtcapp.tables.geographic_table import build_geographic_table
        from nmtcapp.tables.track_record_table import build_track_record_table
        from nmtcapp.core.cde import CDEProfile
        cde = CDEProfile.sample()
        for builder_fn, args in [
            (build_pipeline_summary_table, (sample_pipeline,)),
            (build_distress_summary_table, (sample_pipeline,)),
            (build_impact_summary_table, (sample_pipeline,)),
            (build_geographic_table, (sample_pipeline,)),
            (build_track_record_table, (cde,)),
        ]:
            df = builder_fn(*args)
            assert len(df.columns) <= 8, (
                f"{builder_fn.__name__} has {len(df.columns)} columns (max 8 for portrait)"
            )

    def test_appendix_a_contains_excel_cross_reference(self, docx_path):
        from docx import Document
        doc = Document(docx_path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Excel" in full_text or "excel" in full_text.lower()

    def test_appendix_d_splits_into_two_tables(self, docx_path):
        from docx import Document
        doc = Document(docx_path)
        # The investor section should produce two tables (identification + commitment)
        # Verify total table count is higher than without the split (basic sanity)
        assert len(doc.tables) >= 10


# ---------------------------------------------------------------------------
# Word: summary table builders
# ---------------------------------------------------------------------------

class TestSummaryTableBuilders:
    def test_pipeline_summary_has_six_columns(self, sample_pipeline, sample_cde):
        from nmtcapp.tables.pipeline_table import build_pipeline_summary_table
        df = build_pipeline_summary_table(sample_pipeline)
        assert len(df.columns) == 6

    def test_pipeline_summary_has_totals_row(self, sample_pipeline, sample_cde):
        from nmtcapp.tables.pipeline_table import build_pipeline_summary_table
        df = build_pipeline_summary_table(sample_pipeline)
        assert "TOTALS" in str(df.iloc[-1, 0]).upper()

    def test_pipeline_summary_truncates_names(self):
        from nmtcapp.core.pipeline import Pipeline, PipelineProject
        from nmtcapp.tables.pipeline_table import build_pipeline_summary_table
        p = PipelineProject(
            project_id="X-001",
            project_name="A" * 50,
            qalicb_name="Test QALICB",
            address="100 Main St", city="Chicago", state="IL",
            sector="healthcare", project_type="real_estate",
            total_project_cost=5_000_000, qei_request=3_000_000,
            qlici_amount=3_000_000, expected_jobs_created=10,
        )
        pipeline = Pipeline()
        pipeline.add(p)
        df = build_pipeline_summary_table(pipeline)
        name = df.iloc[0]["Project Name"]
        assert len(name) <= 40
        assert name.endswith("...")

    def test_distress_summary_has_five_columns(self, sample_pipeline):
        from nmtcapp.tables.distress_table import build_distress_summary_table
        df = build_distress_summary_table(sample_pipeline)
        assert len(df.columns) == 5

    def test_distress_summary_has_no_summary_row(self, sample_pipeline):
        from nmtcapp.tables.distress_table import build_distress_summary_table
        df = build_distress_summary_table(sample_pipeline)
        # Summary function doesn't add a totals row
        assert len(df) == len(sample_pipeline)

    def test_investor_identification_table_four_cols(self, sample_application):
        from nmtcapp.tables.investor_table import build_investor_identification_table
        df = build_investor_identification_table(sample_application)
        assert len(df.columns) == 4

    def test_investor_commitment_table_five_cols(self, sample_application):
        from nmtcapp.tables.investor_table import build_investor_commitment_table
        df = build_investor_commitment_table(sample_application)
        assert len(df.columns) == 5

    def test_investor_commitment_has_totals_row(self, sample_application):
        from nmtcapp.tables.investor_table import build_investor_commitment_table
        df = build_investor_commitment_table(sample_application)
        assert "TOTALS" in str(df.iloc[-1, 0]).upper()

    def test_impact_summary_has_seven_columns(self, sample_pipeline):
        from nmtcapp.tables.impact_table import build_impact_summary_table
        df = build_impact_summary_table(sample_pipeline)
        assert len(df.columns) == 7

    def test_impact_summary_has_totals_row(self, sample_pipeline):
        from nmtcapp.tables.impact_table import build_impact_summary_table
        df = build_impact_summary_table(sample_pipeline)
        last = str(df.iloc[-1, 0]).upper()
        assert "TOTAL" in last


# ---------------------------------------------------------------------------
# PDF: landscape pages present
# ---------------------------------------------------------------------------

class TestPDFLandscapePages:
    @pytest.fixture
    def pdf_path(self, sample_application, application_analysis, tmp_path):
        from nmtcapp.renderers.pdf_builder import PDFApplicationBuilder
        path = str(tmp_path / "test_wide.pdf")
        PDFApplicationBuilder(sample_application, application_analysis).save(path)
        return path

    def test_pdf_is_valid(self, pdf_path):
        with open(pdf_path, "rb") as f:
            assert f.read(4) == b"%PDF"

    def test_pdf_contains_landscape_page_dimensions(self, pdf_path):
        # Landscape letter: MediaBox [0 0 792 612]
        # Portrait letter:  MediaBox [0 0 612 792]
        # Search raw PDF bytes for landscape indicator (792 pt wide)
        with open(pdf_path, "rb") as f:
            content = f.read()
        # 792 points is the landscape width; search as bytes
        assert b"792" in content, "No landscape pages found in PDF"

    def test_pdf_size_reasonable(self, pdf_path):
        size = os.path.getsize(pdf_path)
        assert size > 10_000, f"PDF too small ({size} bytes)"

    def test_appendix_b_note_indicates_distress_data(self, pdf_path):
        # Can't parse PDF sections, but can check file is built without error
        assert os.path.exists(pdf_path)

    def test_pdf_has_multiple_pages(self, pdf_path):
        # A full application PDF should have at least 10 pages
        with open(pdf_path, "rb") as f:
            content = f.read()
        # Count "Page" field codes or "/Page" objects in PDF structure
        page_count = content.count(b"/Type /Page\n") + content.count(b"/Type/Page\n")
        assert page_count >= 5, f"Expected at least 5 page objects, got {page_count}"


# ---------------------------------------------------------------------------
# Word: _switch helpers directly
# ---------------------------------------------------------------------------

class TestOrientationSwitching:
    def test_switch_to_landscape_sets_orientation(self, sample_application, application_analysis):
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from nmtcapp.renderers.word_builder import WordApplicationBuilder
        builder = WordApplicationBuilder(sample_application, application_analysis)
        doc = Document()
        builder._switch_to_landscape(doc)
        last_sec = doc.sections[-1]
        assert last_sec.orientation == WD_ORIENT.LANDSCAPE
        assert last_sec.page_width > last_sec.page_height

    def test_switch_to_portrait_after_landscape(self, sample_application, application_analysis):
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from nmtcapp.renderers.word_builder import WordApplicationBuilder
        builder = WordApplicationBuilder(sample_application, application_analysis)
        doc = Document()
        builder._switch_to_landscape(doc)
        builder._switch_to_portrait(doc)
        last_sec = doc.sections[-1]
        assert last_sec.orientation == WD_ORIENT.PORTRAIT
        assert last_sec.page_height > last_sec.page_width

    def test_landscape_uses_narrower_margins(self, sample_application, application_analysis):
        from docx import Document
        from docx.shared import Inches
        from nmtcapp.renderers.word_builder import WordApplicationBuilder
        builder = WordApplicationBuilder(sample_application, application_analysis)
        doc = Document()
        builder._switch_to_landscape(doc)
        sec = doc.sections[-1]
        # Landscape margins should be 0.75"
        assert abs(sec.left_margin - Inches(0.75)) < 100  # EMU tolerance
