"""Auditor scenario B: partial-unverified pipelines disclose in every export.

Scenario B = 6 projects, 2 geocode failures. The pipeline status is "ok"
(the CDFI Fund dataset loaded), but two projects could not be
location-verified. Every export format (Word, PDF, Excel, Markdown) must:

1. Render a banner naming the unverified project IDs.
2. Attach an INLINE qualifier to every eligibility-dependent metric —
   "67% (2 of 6 unverified)" in the same cell/line as the number, never a
   separate paragraph (a separate paragraph can be stripped in editing;
   an inline qualifier cannot).
3. Never assert the bare percentage as an unqualified fact in narratives.
"""
import re
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pytest

from nmtcmapper import EligibilityDownloadError

from nmtcapp.core.application import Application
from nmtcapp.core.cde import CDEProfile
from nmtcapp.core.pipeline import Pipeline, PipelineProject

BANNER_MARK = "could not be location-verified"
QUALIFIER = "(2 of 6 unverified)"
UNVERIFIED_IDS = ("PRJ-B05", "PRJ-B06")


def _project(pid: str, city: str, state: str, qei: float) -> PipelineProject:
    return PipelineProject(
        project_id=pid, project_name=f"Scenario B {pid}",
        qalicb_name=f"{pid} QALICB LLC", address="100 Main St",
        city=city, state=state, sector="healthcare",
        project_type="real_estate", total_project_cost=qei * 1.4,
        qei_request=qei, qlici_amount=qei, expected_jobs_created=40,
    )


def _scenario_b_pipeline() -> Pipeline:
    return Pipeline([
        _project("PRJ-B01", "Chicago", "IL", 8_000_000),
        _project("PRJ-B02", "Houston", "TX", 7_000_000),
        _project("PRJ-B03", "Cleveland", "OH", 6_000_000),
        _project("PRJ-B04", "Memphis", "TN", 5_000_000),
        _project("PRJ-B05", "Nowhere", "MT", 4_000_000),
        _project("PRJ-B06", "Lost City", "WY", 3_000_000),
    ])


def _ok_result(address: str, tract: str, distress: str) -> SimpleNamespace:
    return SimpleNamespace(
        address=address, tract_id=tract, geocode_success=True,
        nmtc_eligible=True, distress_level=distress,
        is_nmtc_native_area=False, is_high_migration_rural=False,
        is_opportunity_zone=False,
    )


def _geocode_failed(address: str) -> SimpleNamespace:
    return SimpleNamespace(
        address=address, tract_id=None, geocode_success=False,
        nmtc_eligible=False, distress_level="ineligible",
        is_nmtc_native_area=False, is_high_migration_rural=False,
        is_opportunity_zone=False,
    )


def _mock_mapper_scenario_b() -> MagicMock:
    mapper = MagicMock()
    mapper.data_source = "cdfi_fund"
    mapper.check_address.side_effect = [
        _ok_result("PRJ-B01", "17031000100", "deep"),
        _ok_result("PRJ-B02", "48201000100", "deep"),
        _ok_result("PRJ-B03", "39035000100", "severe"),
        _ok_result("PRJ-B04", "47157000100", "lic"),
        _geocode_failed("PRJ-B05"),
        _geocode_failed("PRJ-B06"),
    ]
    return mapper


@pytest.fixture(scope="module")
def scenario_b(tmp_path_factory):
    """Analyze scenario B and generate all four export formats once."""
    app = Application(
        cde=CDEProfile.sample(), requested_allocation=55_000_000,
        application_round="CY2025",
    )
    app.add_pipeline(_scenario_b_pipeline())
    with patch("nmtcmapper.NMTCMapper", return_value=_mock_mapper_scenario_b()):
        analysis = app.analyze()
        out_dir = tmp_path_factory.mktemp("scenario_b_exports")
        paths = app.generate(str(out_dir), formats=["markdown", "word", "excel", "pdf"])
    return app, analysis, paths


# ---------------------------------------------------------------------------
# Corpus extractors
# ---------------------------------------------------------------------------

def _word_texts(path: str) -> list:
    from docx import Document
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return texts


def _excel_texts(path: str) -> list:
    from openpyxl import load_workbook
    wb = load_workbook(path)
    texts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for v in row:
                if isinstance(v, str):
                    texts.append(v)
    return texts


def _pdf_text(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    raw = "\n".join(page.extract_text() or "" for page in reader.pages)
    return re.sub(r"\s+", " ", raw)


def _assert_no_bare_pct(texts: list, pct: str = "67%") -> None:
    """Every text unit carrying the percentage must carry the qualifier inline."""
    offenders = [t for t in texts if pct in t and "unverified" not in t.lower()]
    assert not offenders, f"bare {pct} without inline qualifier: {offenders!r}"


# ---------------------------------------------------------------------------
# 1. Analysis-level ground truth
# ---------------------------------------------------------------------------

def test_scenario_b_analysis_ground_truth(scenario_b):
    _, analysis, _ = scenario_b
    pr = analysis.pipeline_result
    assert pr.eligibility_data_status == "ok"
    assert sorted(pr.unverified_project_ids) == list(UNVERIFIED_IDS)
    assert pr.eligibility_pct == pytest.approx(4 / 6)


# ---------------------------------------------------------------------------
# 2. Every export format: banner + IDs + inline qualifier, no bare 67%
# ---------------------------------------------------------------------------

def test_markdown_partial_unverified_disclosure(scenario_b):
    _, _, paths = scenario_b
    with open(paths["markdown"], encoding="utf-8") as f:
        md = f.read()
    assert BANNER_MARK in md
    for pid in UNVERIFIED_IDS:
        assert pid in md
    assert QUALIFIER in md
    _assert_no_bare_pct(md.splitlines())


def test_word_partial_unverified_disclosure(scenario_b):
    _, _, paths = scenario_b
    texts = _word_texts(paths["word"])
    joined = "\n".join(texts)
    assert BANNER_MARK in joined
    for pid in UNVERIFIED_IDS:
        assert pid in joined
    assert QUALIFIER in joined
    _assert_no_bare_pct(texts)


def test_excel_partial_unverified_disclosure(scenario_b):
    _, _, paths = scenario_b
    texts = _excel_texts(paths["excel"])
    joined = "\n".join(texts)
    assert BANNER_MARK in joined
    for pid in UNVERIFIED_IDS:
        assert pid in joined
    assert QUALIFIER in joined
    _assert_no_bare_pct(texts)


def test_pdf_partial_unverified_disclosure(scenario_b):
    _, _, paths = scenario_b
    text = _pdf_text(paths["pdf"])
    assert BANNER_MARK in text
    for pid in UNVERIFIED_IDS:
        assert pid in text
    assert QUALIFIER in text
    # Normalized text: every 67% must be followed by the qualifier nearby
    for m in re.finditer(r"67%", text):
        window = text[m.start():m.start() + 60]
        assert "unverified" in window.lower(), f"bare 67% in PDF near: {window!r}"


# ---------------------------------------------------------------------------
# 3. Section narratives never assert unverified figures as bare fact
# ---------------------------------------------------------------------------

def _section_lines(analysis, app) -> list:
    from nmtcapp.sections.section_a_business import SectionABusinessStrategy
    from nmtcapp.sections.section_b_outcomes import SectionBCommunityOutcomes
    lines = []
    for gen in (SectionABusinessStrategy(), SectionBCommunityOutcomes()):
        content = gen.generate_content(app, analysis)
        for sub in content["subsections"]:
            body = sub["body"]
            if isinstance(body, dict):
                lines.extend(f"{k}: {v}" for k, v in body.items())
            elif isinstance(body, list):
                lines.extend(str(x) for x in body)
            else:
                lines.extend(str(body).splitlines())
    return lines


def test_section_narratives_qualified_in_partial_mode(scenario_b):
    app, analysis, _ = scenario_b
    lines = _section_lines(analysis, app)
    pr = analysis.pipeline_result
    deep_pct = f"{pr.distress_breakdown.get('pct_deep_or_severe', 0):.0%}"
    for pct in ("67%", deep_pct):
        offenders = [
            ln for ln in lines
            if pct in ln and "of QEI" in ln and "unverified" not in ln.lower()
        ]
        assert not offenders, f"bare {pct} fact-claim in narrative: {offenders!r}"


def test_section_narratives_degraded_mode_no_zero_pct_claims():
    app = Application(
        cde=CDEProfile.sample(), requested_allocation=55_000_000,
        application_round="CY2025",
    )
    app.add_pipeline(_scenario_b_pipeline())
    with patch("nmtcmapper.NMTCMapper",
               side_effect=EligibilityDownloadError("CDFI Fund download failed")):
        analysis = app.analyze()
    lines = _section_lines(analysis, app)
    joined = "\n".join(lines)
    assert "0% of QEI" not in joined
    assert "targeting 0%" not in joined.lower()
    assert "unverified" in joined.lower() or "unavailable" in joined.lower()


# ---------------------------------------------------------------------------
# 4. Markdown full-unavailable branch (ported from Word) + methodology note
# ---------------------------------------------------------------------------

def test_markdown_full_unavailable_disclosure(tmp_path):
    from nmtcapp.renderers.markdown_builder import MarkdownApplicationBuilder
    app = Application(
        cde=CDEProfile.sample(), requested_allocation=55_000_000,
        application_round="CY2025",
    )
    app.add_pipeline(_scenario_b_pipeline())
    with patch("nmtcmapper.NMTCMapper",
               side_effect=EligibilityDownloadError("CDFI Fund download failed")):
        analysis = app.analyze()
        md = MarkdownApplicationBuilder(app, analysis).build()

    assert "ELIGIBILITY DATA UNAVAILABLE" in md
    assert "CDFI Fund download failed" in md
    assert "0% of QEI" not in md
    # The methodology note must never cite a data source that did not load
    method_note = md.split("## Methodology Note")[1]
    assert "CDFI Fund NMTC Eligibility Table" not in method_note
    assert "unavailable" in method_note.lower()


def test_markdown_partial_methodology_keeps_source_but_discloses(scenario_b):
    _, _, paths = scenario_b
    with open(paths["markdown"], encoding="utf-8") as f:
        md = f.read()
    method_note = md.split("## Methodology Note")[1]
    # Dataset DID load in scenario B — the citation stays, plus disclosure
    assert "CDFI Fund NMTC Eligibility Table" in method_note
    assert "2 of 6" in method_note and "unverified" in method_note.lower()


# ---------------------------------------------------------------------------
# 5. Readiness score: partial=True with a distinct unverified note
# ---------------------------------------------------------------------------

def test_readiness_partial_on_unverified_projects(scenario_b):
    _, analysis, _ = scenario_b
    rs = analysis.readiness_score
    assert rs.partial is True
    assert "2 projects unverified" in rs.partial_note
    assert "unverified" in rs.summary().lower()


def test_readiness_partial_when_all_projects_unverified():
    """Auditor amendment (3): ALL projects unverified → partial regardless
    of pipeline status."""
    from nmtcapp.validation.readiness_score import compute_readiness_score

    app = Application(
        cde=CDEProfile.sample(), requested_allocation=55_000_000,
        application_round="CY2025",
    )
    app.add_pipeline(_scenario_b_pipeline())
    mapper = MagicMock()
    mapper.data_source = "cdfi_fund"
    mapper.check_address.side_effect = [
        _geocode_failed(f"PRJ-B{i:02d}") for i in range(1, 7)
    ]
    with patch("nmtcmapper.NMTCMapper", return_value=mapper):
        analysis = app.analyze()

    pr = analysis.pipeline_result
    assert pr.eligibility_data_status == "ok"          # dataset loaded fine
    assert len(pr.unverified_project_ids) == 6          # ...but nothing verified
    rs = compute_readiness_score(pr, analysis.validation_results)
    assert rs.partial is True
    assert "unverified" in rs.partial_note
