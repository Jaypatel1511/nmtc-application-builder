"""Tests for WordApplicationBuilder."""
import pytest

pytest.importorskip("docx", reason="python-docx not installed")

from docx import Document
from docx.document import Document as DocumentClass
from nmtcapp.renderers.word_builder import WordApplicationBuilder, _write_df_to_doc, _set_cell_text


class TestWordBuilder:
    @pytest.fixture
    def builder(self, sample_application, application_analysis):
        return WordApplicationBuilder(sample_application, application_analysis)

    @pytest.fixture
    def doc(self, builder):
        return builder.build()

    # ---- Build returns Document ----

    def test_build_returns_document(self, doc):
        assert isinstance(doc, DocumentClass)

    def test_document_has_content(self, doc):
        texts = [p.text for p in doc.paragraphs]
        combined = " ".join(texts)
        assert len(combined) > 100

    # ---- Cover page ----

    def test_cover_has_cde_name(self, doc, sample_application):
        texts = " ".join(p.text for p in doc.paragraphs)
        assert sample_application.cde.name in texts

    def test_cover_has_application_round(self, doc, sample_application):
        # See tests/renderers/test_markdown_builder.test_cover_contains_round:
        # the fixture takes the default, which no longer invents a round
        # (1.5.5 T1). The cover must state the FIELD, which now means the
        # unsupplied-round disclosure.
        from nmtcapp.core.application_round import round_label
        texts = " ".join(p.text for p in doc.paragraphs)
        assert sample_application.application_round is None
        assert round_label(sample_application.application_round) in texts

    # ---- Executive summary ----

    def test_doc_has_executive_summary_heading(self, doc):
        texts = [p.text for p in doc.paragraphs]
        assert any("Executive Summary" in t for t in texts)

    # ---- Sections ----

    def test_doc_has_all_section_headings(self, doc):
        texts = " ".join(p.text for p in doc.paragraphs)
        for letter in ("A", "B", "C", "D", "E"):
            assert f"Section {letter}" in texts

    # ---- Appendices ----

    def test_doc_has_appendix_headings(self, doc):
        texts = " ".join(p.text for p in doc.paragraphs)
        for label in ("Appendix A", "Appendix B", "Appendix C", "Appendix D", "Appendix E"):
            assert label in texts

    def test_doc_has_methodology_appendix(self, doc):
        texts = " ".join(p.text for p in doc.paragraphs)
        assert "Appendix F" in texts or "Methodology" in texts

    # ---- Tables ----

    def test_doc_has_tables(self, doc):
        assert len(doc.tables) > 0

    def test_cover_table_has_cells(self, doc):
        # First table is the banner
        assert doc.tables[0].rows[0].cells[0] is not None

    # ---- Page numbers in footer ----

    def test_sections_have_footers(self, doc):
        for section in doc.sections:
            assert section.footer is not None

    # ---- Save ----

    def test_save_creates_docx(self, builder, tmp_path):
        import os
        path = str(tmp_path / "application.docx")
        builder.save(path)
        assert os.path.exists(path)
        assert os.path.getsize(path) > 1000

    def test_save_creates_parent_dir(self, builder, tmp_path):
        import os
        path = str(tmp_path / "new_dir" / "application.docx")
        builder.save(path)
        assert os.path.exists(path)

    # ---- Configure document ----

    def test_document_uses_calibri_font(self, builder):
        doc = Document()
        builder._configure_document(doc)
        assert doc.styles["Normal"].font.name == "Calibri"


class TestWriteDfToDoc:
    def test_empty_df_adds_no_data_paragraph(self):
        import pandas as pd
        doc = Document()
        _write_df_to_doc(doc, pd.DataFrame())
        texts = [p.text for p in doc.paragraphs]
        assert any("No data" in t for t in texts)

    def test_none_df_adds_no_data_paragraph(self):
        doc = Document()
        _write_df_to_doc(doc, None)
        texts = [p.text for p in doc.paragraphs]
        assert any("No data" in t for t in texts)

    def test_df_adds_table(self):
        import pandas as pd
        doc = Document()
        df = pd.DataFrame({"A": [1, 2, 3], "B": ["x", "y", "z"]})
        _write_df_to_doc(doc, df)
        assert len(doc.tables) == 1

    def test_max_rows_respected(self):
        import pandas as pd
        doc = Document()
        df = pd.DataFrame({"A": range(100)})
        _write_df_to_doc(doc, df, max_rows=10)
        texts = " ".join(p.text for p in doc.paragraphs)
        assert "10 of 100" in texts


class TestSetCellText:
    def test_sets_text(self):
        doc = Document()
        tbl = doc.add_table(rows=1, cols=1)
        cell = tbl.cell(0, 0)
        _set_cell_text(cell, "Hello")
        assert cell.paragraphs[0].text == "Hello"

    def test_bold_flag(self):
        doc = Document()
        tbl = doc.add_table(rows=1, cols=1)
        cell = tbl.cell(0, 0)
        _set_cell_text(cell, "Bold Text", bold=True)
        assert cell.paragraphs[0].runs[0].font.bold is True
