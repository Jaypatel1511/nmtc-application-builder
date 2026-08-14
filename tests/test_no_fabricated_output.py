"""Release gate: no fabricated claim may reach a rendered application.

WHAT THIS ASSERTS, AND WHY IT IS SHAPED THIS WAY

1. It reads the RENDERED ARTIFACT, not the source. Text is extracted back out
   of the .docx and .pdf and read off the .md and .xlsx. A gate that greps
   nmtcapp/*.py cannot see a fabrication reintroduced through a different
   string, an f-string, or a constant — and every literal below reached a
   submitted document through exactly such a path.

2. The denylist is PARAMETRIZED, so `empty_parameter_set_mark = fail_at_collect`
   (pyproject.toml) turns an empty list into a collection ERROR rather than a
   silent skip. A denylist that quietly empties is the failure mode this repo's
   release gates exist to prevent, so it must not be possible here either.

3. It runs under TWO pipelines — one fully enriched, one with every eligibility
   field None. Most of these fabrications were reachable only on the degraded
   path, which is precisely the path a CDE hits when the CDFI Fund download
   fails, and the path where a confident assertion is most damaging.

Each entry names the release-note finding it belongs to. Do not delete an entry
to make a build pass: the string is in the denylist because it was published to
a federal filing and must never return.
"""
from __future__ import annotations

import os

import pytest

from nmtcapp.core.application import Application
from nmtcapp.core.cde import CDEProfile
from nmtcapp.core.pipeline import Pipeline, PipelineProject


# ---------------------------------------------------------------------------
# The denylist — every literal 1.1.6 removed from a rendered document.
#
# (finding, needle, why it must never reappear)
# Matching is case-insensitive on the extracted text.
# ---------------------------------------------------------------------------
FABRICATIONS = [
    # B1 — invented HMDA statistics and a citation for data never fetched.
    ("B1", "2.3×", "invented racial-disparity ratio in a federal filing"),
    ("B1", "2.3x", "invented racial-disparity ratio (ASCII form)"),
    ("B1", "denial rates", "denial-rate claim the tool cannot compute"),
    ("B1", "denial rate", "denial-rate claim the tool cannot compute"),
    ("B1", "HMDA 5-year data", "citation for data the package never fetched"),
    ("B1", "2018–2022", "vintage of an HMDA dataset never loaded"),
    ("B1", "racial disparities", "disparity claim the tool cannot substantiate"),
    ("B1", "minority applicants", "disparity framing with no underlying data"),
    ("B1", "lending desert", "count produced by integer division"),
    # B2 — thresholds invented and attributed to the CDFI Fund.
    ("B2", "CDFI Fund Competitive Minimum", "fabricated CDFI Fund attribution"),
    ("B2", "CDFI Fund Target (Deep/Severe)", "fabricated CDFI Fund attribution"),
    # B3 — unconditional clean-compliance and deployment assertions.
    ("B3", "zero compliance violations", "asserted for CDEs that declared otherwise"),
    ("B3", "no compliance events or repayments have occurred",
     "asserted for CDEs that declared otherwise"),
    ("B3", "All compliance obligations have been met",
     "asserted for CDEs that declared otherwise"),
    ("B3", "within 18 months of award", "heuristic rendered as measured fact"),
    # B4 — pipeline total mislabelled as the amount requested.
    ("B4", "Total QEI Requested", "row label that contradicted the cover page"),
    ("B4", "has established relationships with",
     "investor-relationship claim the tool cannot verify"),
    ("B4", "from one or two lead investors",
     "investor-count claim the tool cannot verify"),
]


@pytest.fixture(scope="module")
def enriched_pipeline() -> Pipeline:
    """A pipeline with every eligibility field populated."""
    return Pipeline([
        _project("EN-001", distress_level="deep", is_nmtc_eligible=True,
                 census_tract="17031838200", is_native_area=True,
                 is_high_migration_rural=False, is_opportunity_zone=True,
                 geocode_success=True),
        _project("EN-002", distress_level="severe", is_nmtc_eligible=True,
                 census_tract="17031320100", is_native_area=False,
                 is_high_migration_rural=True, is_opportunity_zone=False,
                 geocode_success=True),
        _project("EN-003", distress_level="lic", is_nmtc_eligible=True,
                 census_tract="36061011300", is_native_area=False,
                 is_high_migration_rural=False, is_opportunity_zone=False,
                 geocode_success=True),
    ])


@pytest.fixture(scope="module")
def unenriched_pipeline() -> Pipeline:
    """A pipeline where every eligibility field is None.

    This is the degraded path — what a CDE gets when the CDFI Fund
    eligibility download fails. Nothing here may render as a confident
    negative or as a verified statistic.
    """
    return Pipeline([
        _project("UN-001"), _project("UN-002"), _project("UN-003"),
    ])


def _project(pid: str, **kw) -> PipelineProject:
    base = dict(
        project_id=pid, project_name=f"Project {pid}", qalicb_name=f"{pid} QALICB LLC",
        address="100 Main St", city="Springfield", state="IL",
        sector="healthcare", project_type="real_estate",
        total_project_cost=12_000_000, qei_request=8_000_000,
        qlici_amount=8_000_000, expected_jobs_created=45,
        expected_jobs_retained=12, expected_units_built=20,
        expected_sq_ft=15_000,
    )
    base.update(kw)
    return PipelineProject(**base)


# ---------------------------------------------------------------------------
# Rendering + text extraction
# ---------------------------------------------------------------------------

def _rendered_text(pipeline: Pipeline, tmp_path) -> dict:
    """Generate all four formats and extract the TEXT BACK OUT of each.

    Returns {format: extracted_text}. Reading the artifact rather than the
    source is the whole point of this gate.
    """
    app = Application(cde=CDEProfile.sample(), requested_allocation=55_000_000)
    app.add_pipeline(pipeline)
    out = str(tmp_path)
    paths = app.generate(out, formats=["markdown", "word", "excel", "pdf"])

    texts = {}
    for fmt, path in paths.items():
        assert os.path.exists(path), f"{fmt} artifact was not written: {path}"
        texts[fmt] = _extract(fmt, path)
        assert texts[fmt].strip(), f"{fmt} artifact extracted as empty text"
    assert set(texts) == {"markdown", "word", "excel", "pdf"}, (
        f"expected all four formats, got {sorted(texts)}"
    )
    return texts


def _extract(fmt: str, path: str) -> str:
    if fmt == "markdown":
        with open(path, encoding="utf-8") as fh:
            return fh.read()
    if fmt == "word":
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    if fmt == "excel":
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                parts.extend(str(v) for v in row if v is not None)
        return "\n".join(parts)
    if fmt == "pdf":
        from pypdf import PdfReader
        return "\n".join((page.extract_text() or "") for page in PdfReader(path).pages)
    raise AssertionError(f"unknown format {fmt}")


@pytest.fixture(scope="module")
def enriched_text(enriched_pipeline, tmp_path_factory) -> dict:
    return _rendered_text(enriched_pipeline, tmp_path_factory.mktemp("enriched"))


@pytest.fixture(scope="module")
def unenriched_text(unenriched_pipeline, tmp_path_factory) -> dict:
    return _rendered_text(unenriched_pipeline, tmp_path_factory.mktemp("unenriched"))


# ---------------------------------------------------------------------------
# The gate
# ---------------------------------------------------------------------------

def test_denylist_is_populated():
    """An empty denylist certifies nothing. Fail loudly rather than pass."""
    assert FABRICATIONS, "FABRICATIONS is empty — this gate would assert nothing"
    assert len(FABRICATIONS) >= 15, (
        f"FABRICATIONS shrank to {len(FABRICATIONS)} entries; entries are removed "
        "only when the claim is provably unreachable, never to make a build pass"
    )


@pytest.mark.parametrize("fmt", ["markdown", "word", "excel", "pdf"])
@pytest.mark.parametrize(
    "finding,needle,why",
    FABRICATIONS,
    ids=[f"{f}-{n}" for f, n, _ in FABRICATIONS],
)
def test_no_fabrication_in_enriched_output(enriched_text, fmt, finding, needle, why):
    haystack = enriched_text[fmt].lower()
    assert needle.lower() not in haystack, (
        f"{finding}: fabricated text {needle!r} reached the rendered {fmt} "
        f"application ({why})"
    )


@pytest.mark.parametrize("fmt", ["markdown", "word", "excel", "pdf"])
@pytest.mark.parametrize(
    "finding,needle,why",
    FABRICATIONS,
    ids=[f"{f}-{n}" for f, n, _ in FABRICATIONS],
)
def test_no_fabrication_in_unenriched_output(unenriched_text, fmt, finding, needle, why):
    haystack = unenriched_text[fmt].lower()
    assert needle.lower() not in haystack, (
        f"{finding}: fabricated text {needle!r} reached the rendered {fmt} "
        f"application on the DEGRADED path ({why})"
    )


# ---------------------------------------------------------------------------
# Positive assertions — the placeholders that replaced the fabrications must
# actually be present, or "no fabrication" could be satisfied by rendering
# nothing at all.
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("fmt", ["markdown", "word", "pdf"])
def test_cde_todo_placeholders_are_present(unenriched_text, fmt):
    assert "CDE TO COMPLETE".lower() in unenriched_text[fmt].lower(), (
        f"the {fmt} application carries no [CDE TO COMPLETE] marker — an "
        "unfinished application must look unfinished"
    )


def test_requested_allocation_matches_the_cover_page(enriched_text):
    """Section D's requested figure must be the CDE's request, not the pipeline."""
    md = enriched_text["markdown"]
    assert "$55,000,000" in md, "the requested allocation is missing from the document"
    assert "Allocation Requested" in md, (
        "Section D no longer labels the CDE's actual request"
    )


def test_degraded_path_asserts_no_eligibility_source(unenriched_pipeline, tmp_path):
    """A run that loaded no eligibility data must not cite the CDFI Fund table.

    markdown_builder branched on this correctly; word_builder and pdf_builder
    asserted the citation unconditionally until 1.1.6.
    """
    from unittest.mock import patch

    from nmtcmapper import EligibilityDownloadError

    app = Application(cde=CDEProfile.sample(), requested_allocation=55_000_000)
    app.add_pipeline(unenriched_pipeline)
    with patch("nmtcmapper.NMTCMapper",
               side_effect=EligibilityDownloadError("CDFI Fund download failed")):
        paths = app.generate(str(tmp_path), formats=["markdown", "word", "pdf"])
        texts = {f: _extract(f, p) for f, p in paths.items()}

    for fmt, text in texts.items():
        assert "ELIGIBILITY: UNAVAILABLE" in text or "Eligibility data: UNAVAILABLE" in text, (
            f"the {fmt} methodology note does not disclose that eligibility "
            "data was unavailable"
        )
        assert "2016–2020 ACS 5-Year Estimates" not in text, (
            f"the {fmt} methodology note cites the CDFI Fund eligibility table "
            "on a run where the download failed"
        )
