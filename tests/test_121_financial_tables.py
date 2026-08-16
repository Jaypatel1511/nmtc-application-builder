"""The 1.2.1 fixes, asserted against rendered output rather than against intent.

One module per release is a bad habit, so this is scoped deliberately: it holds
the checks that do not belong to an existing module because they are about a
DEFECT CLASS rather than a unit — invented columns, a type-based format rule
standing in for a column-based one, a stale format config, a shipped input that
does not resolve, a refusal message naming a path the reader does not have.

Each test names the item it closes.
"""
from __future__ import annotations

import os
import subprocess
import sys

import pytest

from nmtcapp.core.cde import CDEProfile
from nmtcapp.core.pipeline import Pipeline, PipelineProject
from nmtcapp.tables.pipeline_table import build_pipeline_table


# ---------------------------------------------------------------------------
# B-2: Appendix A prints no column the CDFI Fund never asked for
# ---------------------------------------------------------------------------

# Retrieved from the CY 2024-2025 NMTC Allocation Application, Exhibit A,
# TABLE A5: PROPOSED TRANSACTIONS (pp. 82-84). A full-text search of all 142
# pages returns ZERO occurrences of every string below.
_COLUMNS_THE_FUND_NEVER_ASKS_FOR = (
    "QLICI A Loan ($)",
    "QLICI B Loan ($)",
    "Senior Debt ($)",
    "Subordinate Debt ($)",
    "Annual Operating Budget ($)",
)


@pytest.fixture()
def pipeline() -> Pipeline:
    projects = []
    for i in range(3):
        p = PipelineProject(
            project_id=f"B2-{i}", project_name=f"B2 Project {i}",
            qalicb_name=f"B2 {i} QALICB LLC", address=f"{i} Table A5 Road",
            city="Toledo", state="OH", sector="healthcare",
            project_type="real_estate",
            total_project_cost=float(8_000_000 + i),
            qei_request=float(5_000_000 + i),
            qlici_amount=float(4_900_000 + i),
            expected_jobs_created=20 + i, expected_jobs_retained=4,
            expected_sq_ft=float(12_000 + i),
        )
        p.census_tract = "39095006900"
        p.is_nmtc_eligible = True
        p.distress_level = "severe"
        p.geocode_success = True
        projects.append(p)
    pl = Pipeline(projects)
    pl.eligibility_data_status = "ok"
    return pl


def test_appendix_a_drops_the_five_invented_columns(pipeline):
    df = build_pipeline_table(pipeline)
    present = [c for c in _COLUMNS_THE_FUND_NEVER_ASKS_FOR if c in df.columns]
    assert not present, (
        f"Appendix A still prints {present} as data. None of these appears "
        "anywhere in the CDFI Fund's CY 2024-2025 Allocation Application; the "
        "Fund's per-project attachment is Table A5, which collects one QLICI "
        "total (row h) and defines no tranches, no debt layers and no QALICB "
        "operating budget."
    )


def test_appendix_a_prints_the_cdes_own_qlici_total(pipeline):
    """The figure Table A5 row (h) actually asks for, as the CDE supplied it."""
    df = build_pipeline_table(pipeline)
    assert "Total QLICI ($)" in df.columns
    rows = df[df["Project ID"] != "TOTALS"]
    for _, row in rows.iterrows():
        source = next(p for p in pipeline if p.project_id == row["Project ID"])
        assert row["Total QLICI ($)"] == source.qlici_amount, (
            "the QLICI column must be the CDE's own number, unmodified"
        )


def test_appendix_a_leverage_plus_equity_equals_qei(pipeline):
    """B-3: the identity that replaced the flat 80%, asserted per project."""
    df = build_pipeline_table(pipeline)
    rows = df[df["Project ID"] != "TOTALS"]
    for _, row in rows.iterrows():
        qei = row["QEI Request ($)"]
        implied = row["Leverage Loan ($)"] + row["Estimated Investor Equity ($)"]
        assert abs(implied - qei) < 0.01, (
            f"leverage ${row['Leverage Loan ($)']:,.0f} + equity "
            f"${row['Estimated Investor Equity ($)']:,.0f} != QEI ${qei:,.0f}. "
            "A leverage figure that does not close against the equity printed "
            "beside it describes no structure."
        )


def test_appendix_a_docstring_no_longer_claims_template_parity():
    """The claim went with the columns.

    The module said it "mirrors CDFI Fund CY2025 Excel template format" while
    printing five fields the form does not contain. A false statement about
    provenance is the same defect class as a false figure.
    """
    import nmtcapp.tables.pipeline_table as pt
    doc = pt.__doc__ or ""
    assert 'used to say it "mirrors CDFI Fund' in doc, (
        "the claim must be recorded as withdrawn, not silently deleted — a "
        "future reader needs to know it was checked and found false"
    )
    assert doc.lstrip().startswith("Section A pipeline detail table — the "
                                   "per-project attachment."), (
        "the module still opens by asserting template parity"
    )
    assert "Table A5" in doc, "name the form it was checked against"


# ---------------------------------------------------------------------------
# B-7a: the cell formatter reads the column, not the magnitude
# ---------------------------------------------------------------------------

def test_currency_columns_get_a_dollar_sign():
    from nmtcapp.renderers._cell_format import format_cell
    assert format_cell("QEI ($)", 8_500_000.0) == "$8,500,000"
    assert format_cell("Cost per Job ($)", 125_000) == "$125,000"


def test_non_currency_floats_do_not_get_a_dollar_sign():
    """The instance B-7 asks for, stated as a rule rather than hunted for.

    Whether a live column currently trips the old rule is beside the point: the
    rule was wrong, and any column added tomorrow would trip it. Square footage
    and a Herfindahl index are both non-currency values that clear 1000.
    """
    from nmtcapp.renderers._cell_format import format_cell
    assert format_cell("Square Feet", 24_000.0) == "24,000"
    assert format_cell("Commercial Sq Ft", 58_000.0) == "58,000"
    assert format_cell("HHI", 2_450.0) == "2,450"
    assert format_cell("Jobs/$MM QEI", 7.126) == "7.13"
    for header in ("Square Feet", "HHI", "Total Jobs"):
        assert "$" not in format_cell(header, 5_000.0)


def test_fraction_shares_render_as_percentages():
    """B-1's defect, fixed at the formatter instead of worked around upstream."""
    from nmtcapp.renderers._cell_format import format_cell
    assert format_cell("QEI (% of Total)", 0.3284) == "32.8%"
    assert format_cell("QEI (% of Total)", 1.0) == "100.0%"
    assert format_cell("Poverty Rate (%)", 0.207) == "20.7%"


def test_appendix_c_share_is_a_float_again(pipeline):
    """It sorts and sums in the workbook again; 1.2.0 traded that away."""
    from nmtcapp.tables.geographic_table import build_geographic_table
    df = build_geographic_table(pipeline)
    values = df["QEI (% of Total)"].tolist()
    assert all(isinstance(v, float) for v in values), values
    assert abs(values[-1] - 1.0) < 1e-9, "the TOTAL row must be 1.0, not '100.0%'"


# ---------------------------------------------------------------------------
# B-7b: no sheet may name a column it does not have
# ---------------------------------------------------------------------------

def test_every_excel_sheet_config_names_real_columns(pipeline):
    """Five of six sheets named columns that do not exist. Now it raises.

    The old configs were silent: an unmatched name formats nothing and the
    column falls through to magnitude-based auto-detection, which is how
    Appendix C's share column read 0 0 0 0 0 0 1 in the attachment the Word and
    PDF documents cross-reference as authoritative.
    """
    pytest.importorskip("openpyxl")
    from nmtcapp.core.application import Application
    from nmtcapp.renderers.excel_builder import ExcelApplicationBuilder

    cde = CDEProfile(
        name="Sheet Config CDE, LLC", cde_id="CDE-2022-0001",
        certification_date="2022-01-01", mission="Fixture.",
        target_markets=["Ohio"], prior_awards=[],
        contact={"name": "S", "email": "s@example.org"},
        governance={"board_members": 5, "community_representatives": 2},
    )
    app = Application(cde=cde, requested_allocation=15_000_000.0)
    app.add_pipeline(pipeline)
    # build() raises ValueError from _write_df_to_sheet on a stale name.
    ExcelApplicationBuilder(app, app.analyze()).build()


def test_a_stale_sheet_config_raises(pipeline):
    """Prove the guard fails: name a column that is not there."""
    pytest.importorskip("openpyxl")
    from nmtcapp.core.application import Application
    from nmtcapp.renderers.excel_builder import ExcelApplicationBuilder
    from openpyxl import Workbook

    cde = CDEProfile(
        name="Stale Config CDE, LLC", cde_id="CDE-2022-0002",
        certification_date="2022-01-01", mission="Fixture.",
        target_markets=["Ohio"], prior_awards=[],
        contact={"name": "S", "email": "s@example.org"},
        governance={"board_members": 5, "community_representatives": 2},
    )
    app = Application(cde=cde, requested_allocation=15_000_000.0)
    app.add_pipeline(pipeline)
    builder = ExcelApplicationBuilder(app, app.analyze())

    with pytest.raises(ValueError, match="do not exist"):
        builder._write_df_to_sheet(
            Workbook(), "Stale", build_pipeline_table(pipeline),
            title="Stale", currency_cols=["Leverage Debt ($)"],
        )


def test_excel_weight_column_matches_the_published_weights(pipeline):
    """The Excel dashboard printed 0.0% for two components weighted 25% and 10%.

    weight_map's keys were "eligibility" and "validation"; ReadinessScore's
    component keys are "eligibility_quality" and "validation_pass_rate", so
    .get(comp, 0) returned 0 for both. The Weight column summed to 65% and
    declared two components weightless, against a methodology appendix in the
    same package stating 25% and 10%.
    """
    pytest.importorskip("openpyxl")
    import tempfile

    import openpyxl

    from nmtcapp.core.application import Application
    from nmtcapp.data.schema import READINESS_SCORING_WEIGHTS

    cde = CDEProfile(
        name="Weights CDE, LLC", cde_id="CDE-2022-0003",
        certification_date="2022-01-01", mission="Fixture.",
        target_markets=["Ohio"], prior_awards=[],
        contact={"name": "W", "email": "w@example.org"},
        governance={"board_members": 5, "community_representatives": 2},
    )
    app = Application(cde=cde, requested_allocation=15_000_000.0)
    app.add_pipeline(pipeline)
    out = tempfile.mkdtemp()
    path = app.generate(out, formats=["excel"])["excel"]

    ws = openpyxl.load_workbook(path, data_only=True)["Summary Dashboard"]
    weights = {}
    for row in range(18, 30):
        label = ws.cell(row=row, column=1).value
        if not label or str(label).startswith("OVERALL"):
            break
        weights[str(label)] = ws.cell(row=row, column=4).value

    assert weights, "no readiness weights rendered at all"
    assert all(w for w in weights.values()), (
        f"a component rendered with weight 0: {weights}"
    )
    assert abs(sum(weights.values()) - 1.0) < 1e-9, (
        f"the Weight column sums to {sum(weights.values())}, not 1.0: {weights}"
    )
    assert sorted(weights.values()) == sorted(READINESS_SCORING_WEIGHTS.values())


# ---------------------------------------------------------------------------
# B-5: the Fund's Deep-only bar gets a Deep-only figure
# ---------------------------------------------------------------------------

def test_section_b_reports_deep_distress_alone():
    """The 20% bar is Deep Distress only; deep-only was reported nowhere."""
    from nmtcapp.core.application import Application
    from nmtcapp.sections.section_b_outcomes import SectionBCommunityOutcomes

    projects = []
    for i, level in enumerate(("deep", "severe", "severe", "lic")):
        p = PipelineProject(
            project_id=f"B5-{i}", project_name=f"B5 {i}",
            qalicb_name=f"B5 {i} QALICB LLC", address=f"{i} Bar Street",
            city="Akron", state="OH", sector="healthcare",
            project_type="real_estate", total_project_cost=8_000_000.0,
            qei_request=5_000_000.0, qlici_amount=5_000_000.0,
            expected_jobs_created=20,
        )
        p.census_tract = "39153531500"
        p.is_nmtc_eligible = True
        p.distress_level = level
        p.geocode_success = True
        projects.append(p)
    pl = Pipeline(projects)
    pl.eligibility_data_status = "ok"

    cde = CDEProfile(
        name="Deep Bar CDE, LLC", cde_id="CDE-2023-0005",
        certification_date="2023-03-03", mission="Fixture.",
        target_markets=["Ohio"], prior_awards=[],
        contact={"name": "D", "email": "d@example.org"},
        governance={"board_members": 5, "community_representatives": 2},
    )
    app = Application(cde=cde, requested_allocation=15_000_000.0)
    app.add_pipeline(pl)
    content = SectionBCommunityOutcomes().generate_content(app, app.analyze())
    rows = next(s for s in content["subsections"]
                if s["heading"] == "Distress Level Commitments")["body"]

    deep_only = [v for k, v in rows.items() if "Deep Distress Tracts (" in k]
    assert deep_only, f"deep-only is still reported nowhere: {list(rows)}"
    # 1 of 4 equal-QEI projects is deep.
    assert deep_only[0].startswith("25.0%"), deep_only

    severe = [v for k, v in rows.items() if "Severely Distressed Tracts" in k]
    assert severe, f"the severe share is reported nowhere: {list(rows)}"
    # 3 of 4: one deep (deep IS severe) plus two severe.
    assert severe[0].startswith("75.0%"), severe
    assert deep_only[0] != severe[0], (
        "deep-only and severe render the same string; one number is again "
        "being left to answer both of the Fund's bars"
    )


# ---------------------------------------------------------------------------
# C-1: Native Area is labelled as the CDE's declaration on every surface
# ---------------------------------------------------------------------------

def test_native_area_is_labelled_cde_declared_in_every_table(pipeline):
    from nmtcapp.tables.distress_table import (
        build_distress_summary_table, build_distress_table,
    )
    from nmtcapp.tables.geographic_table import build_geographic_table
    from nmtcapp.tables.impact_table import build_impact_table

    tables = {
        "pipeline": build_pipeline_table(pipeline),
        "distress": build_distress_table(pipeline),
        "distress_summary": build_distress_summary_table(pipeline),
        "impact": build_impact_table(pipeline),
        "geographic": build_geographic_table(pipeline),
    }
    for name, df in tables.items():
        native = [c for c in df.columns if "Native Area" in c]
        assert native, f"{name} has no Native Area column at all"
        for col in native:
            assert "CDE-declared" in col, (
                f"{name}.{col!r} presents a CDE declaration as a tool-derived "
                "finding. The CDFI Fund publishes no tract-keyed NMTC Native "
                "Areas resource, so this tool cannot verify it, and Special "
                "Targeting scores it."
            )


def test_cli_summary_labels_native_area_and_the_house_band(pipeline):
    """The block `nmtcapp analyze` prints, read the same way as the documents."""
    from nmtcapp.intelligence.pipeline_analyzer import PipelineAnalyzer

    text = PipelineAnalyzer().analyze(pipeline).summary()
    assert "Native Area" in text
    native_line = next(l for l in text.splitlines() if "Native Area" in l)
    assert "CDE-declared" in native_line, native_line

    deep_line = next(l for l in text.splitlines() if "Deep/Severe:" in l)
    assert "not a CDFI Fund threshold" in deep_line, (
        f"the target tick cites an undisclosed house band: {deep_line}"
    )


def test_the_mapper_still_does_not_supply_native_area():
    """C-1's chain, asserted at the one link that could silently change.

    nmtc-mapper 0.5.0 dropped is_nmtc_native_area. If a future adapter change
    reads it again the value becomes a dependency artifact rather than the
    CDE's declaration, and the "CDE-declared" labels above become false.
    tests/integrations/test_mapper_contract.py owns the regex; this asserts the
    field's default so the two cannot drift.
    """
    p = PipelineProject(
        project_id="NA-1", project_name="Native Default",
        qalicb_name="NA QALICB LLC", address="1 Default Way", city="Tulsa",
        state="OK", sector="healthcare", project_type="real_estate",
        total_project_cost=5_000_000.0, qei_request=3_000_000.0,
        qlici_amount=3_000_000.0, expected_jobs_created=10,
    )
    assert p.is_native_area is None, (
        "is_native_area defaults to something other than None — an undeclared "
        "project would then carry an assertion the CDE never made"
    )


# ---------------------------------------------------------------------------
# C-2: the exit-code ruling, implemented
# ---------------------------------------------------------------------------

def _templates_dir() -> str:
    import nmtcapp
    return os.path.join(os.path.dirname(nmtcapp.__file__), "templates")


def _run_cli(*args) -> subprocess.CompletedProcess:
    return subprocess.run(
        [sys.executable, "-m", "nmtcapp.cli", *args],
        capture_output=True, text=True,
    )


def test_analyze_help_states_what_the_exit_code_means():
    """'Do not leave it implicit' — so the help says it."""
    result = _run_cli("analyze", "--help")
    assert result.returncode == 0
    text = result.stdout
    assert "--strict" in text
    assert "exit code" in text.lower()
    assert "not what it found" in text or "not what it finds" in text


def test_strict_flag_exists_and_is_documented():
    result = _run_cli("analyze", "--help")
    assert "Exit non-zero when any validation check FAILS" in result.stdout


# ---------------------------------------------------------------------------
# C-3: the shipped sample resolves, and its rows are internally consistent
# ---------------------------------------------------------------------------

def test_shipped_sample_rows_have_matching_state_and_tract_fips():
    """PRJ-S017 said Kansas for a Missouri street. FIPS is the cheap check.

    The Census geocoder resolves "3500 Troost Ave, Kansas City" to MO 64109 and
    refuses it when the state is given as KS — which is why a wrong-state row
    surfaced as "could not be verified" rather than as wrong data. This test
    needs no network: a state and an 11-digit GEOID in the same row must agree
    on the leading two digits.
    """
    import csv

    from nmtcapp.validation.consistency_check import _STATE_FIPS

    path = os.path.join(_templates_dir(), "pipeline_sample_strong.csv")
    with open(path, encoding="utf-8") as fh:
        rows = list(csv.DictReader(l for l in fh if not l.startswith("#")))

    assert rows, "the shipped sample parsed as empty"
    mismatched = [
        (r["project_id"], r["state"], r["census_tract"])
        for r in rows
        if r.get("census_tract") and _STATE_FIPS.get(r["state"].upper())
        and not r["census_tract"].startswith(_STATE_FIPS[r["state"].upper()])
    ]
    assert not mismatched, (
        "shipped sample rows whose declared state and census tract are in "
        f"different states: {mismatched}"
    )


def test_pipeline_sample_rows_have_matching_state_and_tract_fips():
    """The same two addresses live in Pipeline.sample(); same rule."""
    from nmtcapp.validation.consistency_check import _STATE_FIPS

    mismatched = [
        (p.project_id, p.state, p.census_tract)
        for p in Pipeline.sample(n=20)
        if p.census_tract and _STATE_FIPS.get(p.state.upper())
        and not p.census_tract.startswith(_STATE_FIPS[p.state.upper()])
    ]
    assert not mismatched, mismatched


def test_troost_avenue_is_not_in_kansas():
    """Named explicitly so the fix cannot be reverted by a tidy-up."""
    kc = [p for p in Pipeline.sample(n=20) if "Troost" in p.address]
    assert kc, "the Troost Avenue row disappeared; this test needs updating"
    assert kc[0].state == "MO", (
        "Troost Avenue is a Kansas City, MISSOURI street. Kansas City KS is a "
        "separate municipality across the state line."
    )


# ---------------------------------------------------------------------------
# C-5: refusal messages point somewhere the reader can actually go
# ---------------------------------------------------------------------------

def test_refusal_message_does_not_assume_a_git_checkout():
    from nmtcapp.core.sample_identity import SampleDataError, assert_not_sample_identity

    with pytest.raises(SampleDataError) as exc:
        assert_not_sample_identity(name="Riverbend Community Capital CDE, LLC")
    message = str(exc.value)
    assert "nmtcapp/templates/pipeline_template.xlsx" not in message, (
        "the message names a repo-relative path a pip user does not have"
    )
    assert "nmtcapp init" in message
    assert os.path.isdir(_templates_dir())
    assert _templates_dir() in message, (
        "the packaged template directory should be resolved at runtime"
    )


def test_the_refusal_guard_is_still_multi_signal():
    """C-5 is cosmetic and must not weaken the strongest artifact of 1.2.0."""
    from nmtcapp.core.sample_identity import matched_sample_field

    assert matched_sample_field(name="Riverbend Community Capital CDE, LLC") == "CDE name"
    assert matched_sample_field(cde_id="CDE-2018-0117") == "CDE ID"
    assert matched_sample_field(ein="82-1234567") == "EIN"
    assert matched_sample_field(ein="82-4491073") == "EIN"
    # And still not fuzzy.
    assert matched_sample_field(name="Riverbend Housing Partners") is None
    assert matched_sample_field(cde_id="CDE-2018-0118") is None
    assert matched_sample_field(ein="82-1234568") is None


# ---------------------------------------------------------------------------
# C-4: the docs hook cannot report success on a partial build
# ---------------------------------------------------------------------------

_DOCS_HOOK = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "docs", "hooks", "generate_sample_output.py",
)

# THIS TEST IS ABOUT THE DOCS BUILD, AND MANIFEST.in PRUNES docs/.
#
# The hook is loaded by absolute path because mkdocs loads it that way too
# (mkdocs.yml `hooks:`), and it is not importable as a module. In a git
# checkout that path always resolves. Inside an unpacked sdist it never can:
# MANIFEST.in:88 is `prune docs`, so the tarball carries no docs/ at all,
# while release.yml's sdist job runs this suite out of the tarball with no
# checkout to fall back on. The test then raised FileNotFoundError and the
# release job exited 1 — CI stayed green only because CI runs from a checkout.
# A tag on that artifact would have failed before publishing anything.
#
# So the question this test asks — "does the docs hook fail loudly on a
# partial build?" — is not merely unanswerable in the sdist; it is not the
# sdist's question. The sdist does not build the docs and does not ship the
# hook. Skip explicitly, the same way and for the same reason as
# tests/test_no_committed_generated_artifacts.py.
#
# The skip is safe because its condition is unambiguous and cannot occur in
# the repository: docs/hooks/generate_sample_output.py is tracked, and
# tests/test_no_committed_generated_artifacts.py fails if the repo stops being
# a git checkout. It cannot be reached by a regression in the hook — deleting
# or renaming the hook in a CHECKOUT still fails this test, because the
# checkout has a docs/ directory and the file is simply absent from it. What
# it cannot distinguish is a checkout with no docs/ directory at all, which
# would fail the mkdocs build long before this suite ran.
#
# NOT a try/except around the load. Swallowing the error would make the test
# pass while asserting nothing, on every platform, forever. A skip is counted:
# release.yml's FLOOR subtracts skips from the executed total, so a skip that
# started firing in CI shows up as a smaller number rather than as green.
_DOCS_HOOK_PRESENT = os.path.isfile(_DOCS_HOOK)


@pytest.mark.skipif(
    not _DOCS_HOOK_PRESENT,
    reason=(
        "docs/hooks/generate_sample_output.py is absent (this is an unpacked "
        "sdist or an installed tree, not a checkout). MANIFEST.in prunes docs/, "
        "the sdist does not build the documentation, and this gate is about the "
        "docs build's own failure mode."
    ),
)
def test_docs_hook_raises_when_a_configured_format_is_missing(monkeypatch, tmp_path):
    """It logged 'sample output generated' over two of four files under --strict."""
    import importlib.util

    spec = importlib.util.spec_from_file_location("_gen_sample", _DOCS_HOOK)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)

    from nmtcapp.core.application import Application

    def _only_two(self, output_dir="./drafts", formats=None):
        os.makedirs(output_dir, exist_ok=True)
        for name in ("sample.md", "sample.xlsx"):
            open(os.path.join(output_dir, name), "w").close()
        return {"markdown": os.path.join(output_dir, "sample.md"),
                "excel": os.path.join(output_dir, "sample.xlsx")}

    monkeypatch.setattr(Application, "generate", _only_two)
    with pytest.raises(RuntimeError, match="missing 2 of 4 formats"):
        module._build_sample(str(tmp_path))


def _docs_extra_deps():
    """The [docs] extra's dependency strings, parsed WITHOUT tomllib.

    tomllib is 3.11+. requires-python is >=3.9 and CI runs 3.9-3.12, so an
    `import tomllib` here fails on the two oldest interpreters — which is
    exactly where a floor claim needs testing.

    NOT wrapped in try/except ImportError with a skip. A skip exits 0, so the
    test would be silently absent on precisely the two interpreters where it
    currently fails: a gate that cannot fail, which is the class this release
    exists to close. hmda-analyzer hit this identical defect and recorded the
    rule — a test that can only run on 3.11 asserts nothing about a 3.9 floor.

    NOT a tomli dependency either: adding a runtime-version-shim package to
    [dev] to read one list is a heavier answer than the question. The portfolio
    precedent is regex.

    FAILS CLOSED. The two asserts below are the reason this is a helper and not
    an inline findall: a parser that silently returned [] on a table it could
    not find would make the "unpinned operator" check below pass vacuously
    (all() over an empty list is True). It raises instead.
    """
    import re

    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    with open(os.path.join(root, "pyproject.toml"), encoding="utf-8") as fh:
        text = fh.read()
    block = re.search(
        r"^\[project\.optional-dependencies\](.*?)(?=^\[)", text,
        re.MULTILINE | re.DOTALL)
    assert block, "no [project.optional-dependencies] table"
    docs = re.search(r"^docs\s*=\s*\[(.*?)\]", block.group(1),
                     re.MULTILINE | re.DOTALL)
    assert docs, "there is still no [docs] extra"
    return re.findall(r'"([^"]+)"', docs.group(1))


def test_docs_extra_exists_and_covers_the_toolchain():
    """`pip install -e ".[docs]"` used to warn and install only the core."""
    deps = _docs_extra_deps()
    assert deps, (
        "the [docs] extra parsed as empty. Every assertion below would then "
        "pass or fail for the wrong reason."
    )
    names = " ".join(deps).lower()
    for required in ("mkdocs", "mkdocs-material", "python-docx",
                     "openpyxl", "reportlab"):
        assert required in names, (
            f"[docs] does not pin {required}; the docs build renders the "
            "sample application and needs the output libraries"
        )
    assert all(
        any(op in dep for op in (">=", "==", "~="))
        for dep in deps
    ), f"an unpinned docs dependency: {deps}"


# ---------------------------------------------------------------------------
# L-2: no sentence may credit the CDE with a value an adapter can overwrite
# ---------------------------------------------------------------------------

_PACKAGE_SOURCE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "nmtcapp"
)


@pytest.mark.skipif(
    not os.path.isdir(_PACKAGE_SOURCE),
    reason=(
        "nmtcapp/ is not present next to tests/ (this is an unpacked sdist or "
        "an installed tree, not a checkout). This test derives its subject — "
        "the fields an adapter can assign and the columns a CDE fills in — "
        "from the package SOURCE, and release.yml's sdist job deliberately "
        "runs with no source tree in the working directory."
    ),
)
def test_no_rendered_claim_credits_the_cde_with_an_overwritable_flag():
    """The sweep behind the High Migration Rural fix, kept runnable.

    Section A printed "Per the flags supplied in this CDE's own pipeline
    submission, the pipeline places 13% in tracts declared High Migration
    Rural" over a figure the mapper had CORRECTED — the CDE declared 36.2%,
    _prefer_determinate(False, True) returned False, and the document reported
    12.6% as the CDE's own. The mapper was right; the sentence was wrong about
    where the number came from.

    That defect is a JOIN, not a string: a field that is both a CDE-supplied
    column and an adapter-assignable attribute, rendered under a sentence
    crediting the CDE. This test recomputes the join rather than re-checking
    the one sentence, so a second instance fails here the day it is written.

    The current answer is two fields, and neither is falsely attributed:
      is_high_migration_rural   attributed to BOTH authors since 1.2.1 L-2
      is_opportunity_zone       rendered under bare headings ("Opportunity
                                Zone (Y/N)", "OZ") that credit nobody
    """
    import ast
    import re

    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    # Fields any integration assigns onto a project.
    assigned = set()
    integrations = os.path.join(root, "nmtcapp", "integrations")
    for name in sorted(os.listdir(integrations)):
        if not name.endswith(".py"):
            continue
        src = open(os.path.join(integrations, name), encoding="utf-8").read()
        assigned |= set(re.findall(r"project\.(\w+)\s*=", src))
    assert assigned, "no adapter assigns anything — the join would be empty"

    # Boolean columns the CDE fills in on the shipped template.
    pipeline_src = open(
        os.path.join(root, "nmtcapp", "core", "pipeline.py"), encoding="utf-8"
    ).read()
    cde_columns = set(re.findall(r'_optional_bool\(row\.get\("(\w+)"', pipeline_src))
    assert cde_columns, "no CDE-declared boolean columns found — regex is stale"

    overwritable = sorted({f"is_{c}" for c in cde_columns} & assigned)
    assert overwritable == ["is_high_migration_rural", "is_opportunity_zone"], (
        f"the set of CDE-declared, adapter-overwritable fields changed to "
        f"{overwritable}. Each one needs its rendered attribution checked: a "
        "sentence saying 'supplied by this CDE' over a value an adapter can "
        "correct is wrong about its own author, and it reaches a scored "
        "Special Targeting criterion. Adjudicate the new field, then update "
        "this list."
    )

    # And Section A must not have gone back to crediting the CDE for HMR.
    section_a = open(
        os.path.join(root, "nmtcapp", "sections", "section_a_business.py"),
        encoding="utf-8",
    ).read()
    rendered = "\n".join(
        line for line in section_a.splitlines() if not line.strip().startswith("#")
    )
    tree = ast.parse(rendered)
    strings = [
        n.value for n in ast.walk(tree)
        if isinstance(n, ast.Constant) and isinstance(n.value, str)
    ]
    hmr_claims = [
        s for s in strings
        if "High Migration Rural" in s and "CDE's own" in s
    ]
    assert not hmr_claims, (
        "Section A credits the CDE with the High Migration Rural share again: "
        f"{hmr_claims}. That share is overwritten by nmtc_mapper_adapter "
        "wherever the mapper returned a determination."
    )


# ---------------------------------------------------------------------------
# L-1: Word's Appendix A must not drop a column without saying so
# ---------------------------------------------------------------------------

def test_word_appendix_a_asks_for_columns_the_table_actually_has(pipeline):
    """Eleven columns rendered where the comment claimed twelve.

    word_builder asked for "NMTC Native Area (Y/N)". The table renders "NMTC
    Native Area (CDE-declared, Y/N)" — the heading 1.2.1 widened so every
    surface would say whose declaration the flag is. `[c for c in
    landscape_cols if c in full_df.columns]` dropped it without a word, and the
    CHANGELOG's "every surface now says so" became false for Word's Appendix A,
    which had stopped saying anything at all.
    """
    import inspect
    import re

    from nmtcapp.renderers import word_builder

    src = inspect.getsource(word_builder.WordApplicationBuilder._build_appendix_pipeline)
    requested = re.findall(r'"([^"]+)"', src.split("landscape_cols = [", 1)[1].split("]", 1)[0])
    assert len(requested) == 12, f"expected the 12-column subset, found {requested}"

    produced = set(build_pipeline_table(pipeline).columns)
    missing = [c for c in requested if c not in produced]
    assert not missing, (
        f"Word's Appendix A asks for {missing}, which tables/pipeline_table "
        "does not produce. A silent drop is how a disclosure disappears."
    )
    assert "NMTC Native Area (CDE-declared, Y/N)" in requested, (
        "the Native Area column must carry its (CDE-declared) disclosure into "
        "Word's Appendix A, not just into the Excel attachment"
    )


def test_word_appendix_a_raises_on_an_unmatched_column(pipeline, monkeypatch):
    """The filter is gone; a shape change must fail loudly."""
    pytest.importorskip("docx")

    from nmtcapp.core.application import Application
    from nmtcapp.core.cde import CDEProfile
    from nmtcapp.renderers import word_builder

    real = word_builder.build_pipeline_table

    def _renamed(pl, cde=None):
        df = real(pl, cde)
        return df.rename(columns={"NMTC Native Area (CDE-declared, Y/N)": "Native Area"})

    monkeypatch.setattr(word_builder, "build_pipeline_table", _renamed)

    cde = CDEProfile(
        name="L1 Fixture CDE, LLC", cde_id="CDE-2019-0001",
        certification_date="2019-01-01", mission="L1 fixture.",
        target_markets=["Ohio"], prior_awards=[],
        contact={"name": "L1", "email": "l1@example.org"},
        governance={"board_members": 7, "community_representatives": 3},
    )
    app = Application(cde=cde, requested_allocation=15_000_000.0)
    app.add_pipeline(pipeline)

    with pytest.raises(KeyError, match="Native Area"):
        word_builder.WordApplicationBuilder(app, app.analyze()).build()


# ---------------------------------------------------------------------------
# L-5: three smaller ones, each asserted against rendered output
# ---------------------------------------------------------------------------

def test_an_unsupplied_affordable_unit_count_is_not_a_zero(pipeline):
    """`expected_units_built=None` used to render "0 affordable units".

    A supplied zero and an absent value are different statements to a federal
    reviewer, and the field feeds a Community Outcomes measure. Same class as
    the "Quarterly" governance default 1.2.0 removed.
    """
    from nmtcapp.renderers._cell_format import NOT_SUPPLIED, format_cell
    from nmtcapp.tables.impact_table import build_impact_table

    for p in pipeline:
        assert p.expected_units_built is None, "fixture must not supply units"

    for df, column in (
        (build_pipeline_table(pipeline), "Affordable Units Built"),
        (build_impact_table(pipeline), "Affordable Units"),
    ):
        rows = df[df[df.columns[0]].astype(str).str.contains("TOTAL") == False]  # noqa: E712
        for value in rows[column]:
            assert format_cell(column, value) == NOT_SUPPLIED, (
                f"{column} rendered {format_cell(column, value)!r} for a project "
                "that supplied no unit count. A zero is a claim."
            )

    # A real zero still reads as a zero: a CDE writing 0 has said something.
    assert format_cell("Affordable Units Built", 0) == "0"


def test_an_unrecognised_sector_is_marked_as_unrecognised(pipeline):
    """`retail` used to render as "Retail", exactly like a recognised sector.

    The list is advisory on the way in — the project loads, and every derived
    figure is unaffected — and declarative on the way out.
    """
    from nmtcapp.data.schema import VALID_SECTORS
    from nmtcapp.tables.impact_table import build_impact_summary_table

    assert "retail" not in VALID_SECTORS, "fixture assumes retail is unrecognised"
    # A recognised sector is NOT marked. Asserted before the mutation, because
    # `pipeline` hands out the same PipelineProject objects.
    assert build_pipeline_table(pipeline).iloc[0]["Sector (as supplied)"] == "Healthcare"

    projects = list(pipeline)
    projects[0].sector = "retail"
    pl = Pipeline(projects)
    pl.eligibility_data_status = "ok"

    cell = build_pipeline_table(pl).iloc[0]["Sector (as supplied)"]
    assert cell == "Retail (not one of the sectors this tool recognises)", cell

    summary = build_impact_summary_table(pl).iloc[0]["Sector"]
    assert "not one of the sectors this tool recognises" in summary, summary


def test_section_e_does_not_promise_detail_it_does_not_have():
    """It printed "including states served, sectors financed, and outcomes
    achieved" directly above rows reading "States: N/A. Sectors: N/A."."""
    from nmtcapp.core.application import Application
    from nmtcapp.core.cde import CDEProfile
    from nmtcapp.sections.section_e_prior_awards import SectionEPriorAwards

    cde = CDEProfile(
        name="Section E Fixture CDE, LLC", cde_id="CDE-2015-0002",
        certification_date="2015-02-02", mission="Section E fixture.",
        target_markets=["Ohio"],
        prior_awards=[{"year": 2019, "amount": 30_000_000,
                       "deployment_status": "fully_deployed"}],
        contact={"name": "E", "email": "e@example.org"},
        governance={"board_members": 7, "community_representatives": 3},
    )
    app = Application(cde=cde, requested_allocation=20_000_000.0)
    content = SectionEPriorAwards().generate_content(app, None)
    body = "\n".join(
        s["body"] if isinstance(s["body"], str) else "\n".join(s["body"])
        for s in content["subsections"]
    )

    assert "including states served, sectors financed, and outcomes achieved" not in body, (
        "Section E promises detail the table does not contain"
    )
    assert "States: N/A" not in body and "Sectors: N/A" not in body, (
        "'N/A' reads as a value the CDE supplied; nothing was collected"
    )
    assert "not collected by this tool" in body, (
        "the absence must be stated, not left as a blank or an N/A"
    )

    # And when the CDE DOES supply them, the sentence names them.
    cde.prior_awards[0]["states"] = ["OH", "KY"]
    cde.prior_awards[0]["sectors"] = ["healthcare"]
    body2 = "\n".join(
        s["body"] if isinstance(s["body"], str) else "\n".join(s["body"])
        for s in SectionEPriorAwards().generate_content(app, None)["subsections"]
    )
    assert "including states served and sectors financed" in body2, body2[:400]


# ---------------------------------------------------------------------------
# L-4: the attribution gate does not scan docs/, so this does
# ---------------------------------------------------------------------------

_DOCS_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs"
)

# Claims withdrawn from the RENDERED document that the DOCUMENTATION SITE went
# on making. Each is (phrase, why it was withdrawn).
#
# WHY THIS EXISTS SEPARATELY. tests/test_attributed_claims.py and
# tests/test_no_fabricated_output.py both read artifacts they render
# themselves into a temp directory; docs/ is outside their world, exactly as
# examples/sample_output/ was before it was deleted. So "Net subsidy to
# QALICB" was removed from the document in 1.2.1 and stayed on
# docs/workflow/output-formats.md as a row the document contains — a
# description of output that no longer exists, published on the site a CDE
# reads before running anything.
#
# NOT a general fabrication scan. It is a list of claims this repository has
# already established to be wrong, checked against the pages that describe the
# output, so a withdrawal has to happen in both places or fail here.
_WITHDRAWN_FROM_DOCS = (
    ("net subsidy to the QALICB",
     "1.2.1 B-4: the row is QEI minus the CDE fee, ~97.5% of QEI, and includes "
     "the whole leverage loan. Renamed to 'QEI Less CDE Fees ($)'"),
    ("Net subsidy to QALICB",
     "1.2.1 B-4: same row, as listed in the Excel sheet contents"),
    ("Sector (NAICS)",
     "1.2.0: the column mapped a coarse sector label to a NAICS code this tool "
     "invented. Deleted; the column is now 'Sector (as supplied)'"),
    ("QLICI A Loan",
     "1.2.1 B-2: an invented tranche. Table A5 collects one QLICI total"),
    ("top quartile",
     "1.2.0: a percentile claim against a distribution nothing publishes"),
)


@pytest.mark.skipif(
    not os.path.isdir(_DOCS_DIR),
    reason=(
        "docs/ is absent (this is an unpacked sdist or an installed tree). "
        "MANIFEST.in prunes docs/, and the sdist does not build or publish the "
        "documentation site."
    ),
)
def test_the_docs_site_does_not_still_describe_withdrawn_output():
    """A claim removed from the document must be removed from the site too."""
    pages = []
    for dirpath, _dirs, files in os.walk(_DOCS_DIR):
        for name in sorted(files):
            if name.endswith(".md"):
                pages.append(os.path.join(dirpath, name))
    assert pages, "no documentation pages found — this scan would check nothing"

    failures = []
    for path in pages:
        text = open(path, encoding="utf-8").read()
        for lineno, line in enumerate(text.splitlines(), 1):
            for phrase, why in _WITHDRAWN_FROM_DOCS:
                if phrase.lower() not in line.lower():
                    continue
                # A page may DISCUSS a withdrawn claim, and 1.2.1 adds notes
                # that do exactly that. What it may not do is present one as
                # current output. The marker is the withdrawal note itself.
                context = "\n".join(text.splitlines()[max(0, lineno - 12):lineno + 12])
                if any(m in context for m in
                       ("not a net subsidy", "Renamed in 1.2.1", "renamed in 1.2.1",
                        "WAS DELETED", "was deleted", "withdrawn", "no longer")):
                    continue
                rel = os.path.relpath(path, os.path.dirname(_DOCS_DIR))
                failures.append(f"  {rel}:{lineno}: {line.strip()[:100]}\n      {why}")

    assert not failures, (
        f"{len(failures)} documentation line(s) describe output this package "
        "has withdrawn:\n\n" + "\n".join(failures) + "\n\n"
        "The attribution and fabrication gates render artifacts into a temp "
        "directory and never look at docs/, so a claim can be removed from the "
        "document and left on the site indefinitely. Fix the page, or add the "
        "withdrawal note that explains what the phrase used to mean.\n\n"
        "NOTE: docs/ changes do not reach the published site until "
        "`mkdocs gh-deploy` is run by hand."
    )


# ---------------------------------------------------------------------------
# FIX-2 B-1: one filing may not state two different severe-distress shares
# ---------------------------------------------------------------------------

def _b1_pipeline(levels) -> Pipeline:
    """Equal-QEI projects at the given distress levels, verified-shaped."""
    projects = []
    for i, level in enumerate(levels):
        p = PipelineProject(
            project_id=f"B1-{i}", project_name=f"B1 Project {i}",
            qalicb_name=f"B1 {i} QALICB LLC", address=f"{100 + i} Subset Way",
            city="Youngstown", state="OH", sector="healthcare",
            project_type="real_estate", total_project_cost=9_000_000.0,
            qei_request=6_000_000.0, qlici_amount=6_000_000.0,
            expected_jobs_created=30,
        )
        p.census_tract = "39099810900"
        p.is_nmtc_eligible = True
        p.distress_level = level
        p.geocode_success = True
        projects.append(p)
    pl = Pipeline(projects)
    pl.eligibility_data_status = "ok"
    return pl


def _b1_application(levels):
    from nmtcapp.core.application import Application

    cde = CDEProfile(
        name="Subset Relation CDE, LLC", cde_id="CDE-2022-0011",
        certification_date="2022-02-02", mission="Fixture.",
        target_markets=["Ohio"], prior_awards=[],
        contact={"name": "S", "email": "s@example.org"},
        governance={"board_members": 6, "community_representatives": 3},
    )
    app = Application(cde=cde, requested_allocation=24_000_000.0)
    app.add_pipeline(_b1_pipeline(levels))
    return app


@pytest.mark.parametrize("levels", [
    # THE SHAPE THAT SHIPPED THE CONTRADICTION: every distressed tract is deep,
    # so the exclusive "severe" bucket is empty. Section B printed 0.0% under
    # "QEI in Severely Distressed Tracts" while Appendix B flagged the same
    # projects "Yes".
    ("deep", "deep", "lic", "lic", "lic"),
    ("deep", "severe", "severe", "lic"),
    ("severe", "severe", "lic", "lic"),
    ("deep",),
])
def test_rendered_severe_share_is_never_below_the_rendered_deep_share(levels):
    """Deep distress is a SUBSET of severe, so severe >= deep. Always.

    Verified against the CDFI Fund's own NMTC LIC Eligibility workbook
    (2016-2020 ACS), columns O and P, over all 85,395 tracts: 8,061 rows are
    flagged deep and every one of them is also flagged severe. The count of
    deep-but-not-severe tracts is ZERO. See
    intelligence/distress_analysis.DEEP_IS_SUBSET_OF_SEVERE.

    THE PIN IS ON THE RENDERED STRINGS, not on the dict. 1.2.1's own arithmetic
    was correct in every bucket; what was wrong was which bucket a label was
    attached to, and no test that reads the dict can see that.
    """
    from nmtcapp.sections.section_b_outcomes import SectionBCommunityOutcomes

    app = _b1_application(levels)
    content = SectionBCommunityOutcomes().generate_content(app, app.analyze())
    rows = next(s for s in content["subsections"]
                if s["heading"] == "Distress Level Commitments")["body"]

    def _pct(label_fragment: str) -> float:
        hits = [v for k, v in rows.items()
                if label_fragment in k and v.rstrip("%").replace(".", "").isdigit()]
        assert hits, f"no row matching {label_fragment!r} in {list(rows)}"
        return float(hits[0].rstrip("%"))

    deep = _pct("Deep Distress Tracts (")
    severe = _pct("Severely Distressed Tracts")

    assert severe >= deep, (
        f"the document states {severe}% of QEI in severely distressed tracts "
        f"and {deep}% in deep distress tracts. Deep distress is a strict "
        "subset of severe distress in the Fund's own workbook, so a severe "
        "share below the deep share is arithmetically impossible and one of "
        "the two rows is measuring a bucket its label does not describe."
    )


def test_section_b_severe_row_agrees_with_appendix_b_flag():
    """Two surfaces, one fact. They disagreed in 1.2.1.

    Appendix B carries a per-project "Severely Distressed Flag" that counts a
    deep-distress tract as severely distressed — correctly. Section B carried a
    percentage under the same words computed from a bucket that excluded them.
    A CDE filing both was filing two answers to one question.
    """
    from nmtcapp.sections.section_b_outcomes import SectionBCommunityOutcomes
    from nmtcapp.tables.distress_table import build_distress_table

    levels = ("deep", "deep", "severe", "lic")
    app = _b1_application(levels)
    analysis = app.analyze()

    content = SectionBCommunityOutcomes().generate_content(app, analysis)
    rows = next(s for s in content["subsections"]
                if s["heading"] == "Distress Level Commitments")["body"]
    severe_row = next(v for k, v in rows.items()
                      if "Severely Distressed Tracts" in k)

    df = build_distress_table(app.pipeline)
    # The last row is Appendix B's own SUMMARY line ("3/4 (75%)"), not a project.
    project_rows = df[df["Project ID"] != "SUMMARY"]
    flags = [str(v) for v in project_rows["Severely Distressed Flag"]]
    flagged = sum(1 for f in flags if f == "Yes")

    # Equal QEI per project, so the count share and the dollar share coincide.
    appendix_b_pct = 100.0 * flagged / len(flags)
    section_b_pct = float(severe_row.rstrip("%"))

    assert section_b_pct == pytest.approx(appendix_b_pct, abs=0.1), (
        f"Section B says {section_b_pct}% of QEI is in severely distressed "
        f"tracts; Appendix B flags {flagged} of {len(flags)} projects "
        f"({appendix_b_pct}%) as severely distressed. Same document, same "
        "word, two numbers."
    )


def test_the_ambiguous_pct_severe_key_is_gone():
    """The rename is the fix. An alias would restore the ambiguity.

    'severe' meant the exclusive bucket in distress_analysis and the inclusive
    flag in distress_table, and the package rendered a label using the third
    meaning. Only one of the three may keep the bare word.
    """
    from nmtcapp.intelligence.distress_analysis import (
        DISTRESS_SHARE_SEMANTICS, analyze_distress_concentration,
    )

    result = analyze_distress_concentration(_b1_pipeline(("deep", "severe")))
    assert "pct_severe" not in result, (
        "pct_severe is back. It named the severe-EXCLUDING-deep bucket while "
        "reading as the severe share, which is how a filing came to state "
        "both 0.0% and 40%."
    )
    for key in DISTRESS_SHARE_SEMANTICS:
        assert key in result, f"{key} is documented but not emitted"
    # The arithmetic the three names promise.
    assert result["pct_deep"] + result["pct_severe_excluding_deep"] == pytest.approx(
        result["pct_deep_or_severe"]
    )


# ---------------------------------------------------------------------------
# FIX-2 B-2: the aggregate surface, which the L-5 gate did not reach
# ---------------------------------------------------------------------------

def test_section_b_does_not_claim_zero_units_over_an_unfilled_column():
    """"0 affordable or mixed-income housing units developed" over no input.

    The 1.2.1 gate above (test_an_unsupplied_affordable_unit_count_is_not_a_zero)
    checks the TABLE cells and passed while this line was still wrong, because
    the aggregate takes a different path: impact_aggregator summed
    `p.expected_units_built or 0` and Section B formatted the result with
    `{units:,}`. The table said "—" and the paragraph above it said "0", in one
    document.
    """
    from nmtcapp.intelligence.impact_aggregator import aggregate_impact
    from nmtcapp.sections.section_b_outcomes import SectionBCommunityOutcomes

    app = _b1_application(("deep", "severe", "lic"))
    for p in app.pipeline:
        assert p.expected_units_built is None, "fixture must not supply units"

    impact = aggregate_impact(app.pipeline)
    assert impact["total_units_built"] is None, (
        "an unfilled column aggregated to a number again"
    )
    assert impact["projects_reporting_units"] == 0

    content = SectionBCommunityOutcomes().generate_content(app, app.analyze())
    body = next(s for s in content["subsections"]
                if s["heading"] == "Aggregate Community Impact Projections")["body"]
    assert "0 affordable or mixed-income housing units developed" not in body, body
    assert "not reported" in body, body


def test_a_supplied_zero_still_renders_as_a_zero():
    """The other half. A CDE that answers 0 has answered."""
    from nmtcapp.intelligence.impact_aggregator import aggregate_impact
    from nmtcapp.sections.section_b_outcomes import SectionBCommunityOutcomes

    app = _b1_application(("deep", "severe", "lic"))
    for p in app.pipeline:
        p.expected_units_built = 0

    impact = aggregate_impact(app.pipeline)
    assert impact["total_units_built"] == 0
    assert impact["projects_reporting_units"] == 3
    assert impact["projects_with_units"] == 0

    content = SectionBCommunityOutcomes().generate_content(app, app.analyze())
    body = next(s for s in content["subsections"]
                if s["heading"] == "Aggregate Community Impact Projections")["body"]
    assert "0 affordable or mixed-income housing units developed" in body, body


def test_projects_with_units_is_read_by_something():
    """It was computed every run and read by nothing in the package.

    A derived value nobody consumes cannot be wrong in a way anybody notices,
    which is how the distinction it encodes stayed unused while the figure
    beside it was collapsing a supplied zero into an absent one.
    """
    from nmtcapp.intelligence.impact_aggregator import aggregate_impact

    app = _b1_application(("deep", "severe", "lic"))
    projects = list(app.pipeline)
    projects[0].expected_units_built = 40
    projects[1].expected_units_built = 0
    projects[2].expected_units_built = None

    impact = aggregate_impact(app.pipeline)
    assert impact["total_units_built"] == 40
    assert impact["projects_reporting_units"] == 2, "0 is an answer; None is not"
    assert impact["projects_with_units"] == 1, "only one project builds housing"


# ---------------------------------------------------------------------------
# FIX-2 B-3: a year is an identifier, not a quantity
# ---------------------------------------------------------------------------

_IDENTIFIER_COLUMNS = (
    ("Award Year", 2019, "2019"),
    ("Census Tract (11-digit)", 17031838200, "17031838200"),
    ("Census Tract (GEOID)", 17031838200, "17031838200"),
    ("ZIP Code", 60653, "60653"),
    ("Project ID", 1042, "1042"),
    ("State FIPS Code", 17, "17"),
)


@pytest.mark.parametrize("header,value,expected", _IDENTIFIER_COLUMNS)
def test_an_identifier_column_never_takes_a_thousands_separator(header, value, expected):
    """`f"{2019:,}"` is "2,019", and nothing is 2,019 of anything.

    1.2.1's _cell_format unification correctly gave the currency columns their
    dollar signs and gave a prior award year a thousands separator on the way
    past. A GEOID would have been worse: 17,031,838,200 in a column citing the
    CDFI Fund's eligibility table is not obviously a formatting fault.
    """
    from nmtcapp.renderers._cell_format import format_cell, is_identifier_column

    assert is_identifier_column(header), f"{header!r} is not classed as an identifier"
    assert format_cell(header, value) == expected


def test_a_quantity_column_still_takes_its_separator():
    """The rule must discriminate, or it is just a disabled formatter."""
    from nmtcapp.renderers._cell_format import format_cell, is_identifier_column

    for header, value, expected in (
        ("Jobs Created", 1200, "1,200"),
        ("Square Feet", 662900, "662,900"),
        ("Affordable Units Built", 1500, "1,500"),
        ("QEI ($)", 88500000, "$88,500,000"),
        ("QEI (% of Total)", 0.307, "30.7%"),
    ):
        assert not is_identifier_column(header), f"{header!r} is not an identifier"
        assert format_cell(header, value) == expected


def test_the_track_record_year_renders_unseparated_on_every_surface():
    """Markdown, Word, PDF and Excel each printed 2,019. All four are checked.

    Excel is the one that predates 1.2.1: "Award Year" is named in no format
    list on that sheet, so it fell through to the magnitude-based auto-detect
    and took #,##0, which Excel displays with the separator. The three prose
    renderers acquired it when 1.2.1 unified their cell formatting.
    """
    import openpyxl
    from docx import Document
    from pypdf import PdfReader

    from nmtcapp.core.application import Application
    from nmtcapp.renderers._cell_format import format_cell
    from nmtcapp.renderers.excel_builder import FMT_IDENTIFIER
    from nmtcapp.tables.track_record_table import build_track_record_table

    cde = CDEProfile(
        name="Comma Free CDE, LLC", cde_id="CDE-2014-0099",
        certification_date="2014-04-04", mission="Fixture.",
        target_markets=["Ohio"],
        prior_awards=[{"year": 2019, "amount": 45_000_000,
                       "deployment_status": "fully_deployed"}],
        contact={"name": "C", "email": "c@example.org"},
        governance={"board_members": 5, "community_representatives": 2},
    )

    df = build_track_record_table(cde)
    rendered = [format_cell("Award Year", v) for v in df["Award Year"]]
    assert "2,019" not in rendered, rendered
    assert "2019" in rendered, rendered

    app = Application(cde=cde, requested_allocation=20_000_000.0)
    app.add_pipeline(_b1_pipeline(("deep", "severe", "lic")))

    import tempfile
    with tempfile.TemporaryDirectory() as out:
        paths = app.generate(out, formats=["markdown", "word", "excel", "pdf"])
        assert set(paths) == {"markdown", "word", "excel", "pdf"}, sorted(paths)

        with open(paths["markdown"], encoding="utf-8") as fh:
            assert "2,019" not in fh.read()

        doc = Document(paths["word"])
        word_text = "\n".join(
            [p.text for p in doc.paragraphs]
            + [c.text for t in doc.tables for r in t.rows for c in r.cells]
        )
        assert "2,019" not in word_text

        pdf_text = "\n".join(
            (page.extract_text() or "") for page in PdfReader(paths["pdf"]).pages
        )
        assert "2,019" not in pdf_text

        # Excel holds a NUMBER, so the separator lives in the number format,
        # not in the cell value. Reading the value alone would pass while the
        # workbook still displayed 2,019 — which is exactly how it survived.
        ws = openpyxl.load_workbook(paths["excel"])["Track Record"]
        year_col = next(
            c for c in range(1, ws.max_column + 1)
            if ws.cell(row=3, column=c).value == "Award Year"
        )
        cell = ws.cell(row=4, column=year_col)
        assert cell.value == 2019
        assert cell.number_format == FMT_IDENTIFIER, (
            f"Award Year carries number format {cell.number_format!r}; Excel "
            "displays #,##0 as '2,019'."
        )
        assert "#" not in cell.number_format


# ---------------------------------------------------------------------------
# FIX-2 G-5 sweep: the THIRD instance of the duplicated-list class
# ---------------------------------------------------------------------------

def test_the_required_project_field_list_is_not_duplicated():
    """Three independent statements of the same twelve names.

    data/schema.REQUIRED_PROJECT_FIELDS (which validation/completeness_check
    imports), the ``required_cols`` set inside Pipeline.from_csv, and
    PipelineProject's own non-default dataclass fields. They agreed by hand.

    The mutation profile is identical to the required-CDE-fields triplication:
    drop a name from the schema list and completeness_check stops flagging it
    while from_csv still demands it and the suite stays green; add one and a
    CSV loads that later fails at render time. Only the dataclass copy fails
    loudly.

    from_csv now derives. The dataclass cannot derive — it IS the type — so it
    is asserted against the same list here.
    """
    import dataclasses

    from nmtcapp.data.schema import REQUIRED_PROJECT_FIELDS

    assert REQUIRED_PROJECT_FIELDS, "the authoritative list is empty"

    required_dataclass_fields = {
        f.name for f in dataclasses.fields(PipelineProject)
        if f.default is dataclasses.MISSING
        and f.default_factory is dataclasses.MISSING
    }
    assert required_dataclass_fields == set(REQUIRED_PROJECT_FIELDS), (
        "PipelineProject's mandatory fields and REQUIRED_PROJECT_FIELDS have "
        "diverged.\n"
        f"  dataclass only: {sorted(required_dataclass_fields - set(REQUIRED_PROJECT_FIELDS))}\n"
        f"  schema only:    {sorted(set(REQUIRED_PROJECT_FIELDS) - required_dataclass_fields)}"
    )

    # from_csv must DERIVE, not merely agree. An equal copy is what let the
    # CDE-field list drop 'governance' with 955 tests green.
    import inspect
    source = inspect.getsource(Pipeline.from_csv)
    assert "REQUIRED_PROJECT_FIELDS" in source, (
        "Pipeline.from_csv no longer reads REQUIRED_PROJECT_FIELDS. It held "
        "its own hand-typed copy of the same twelve names."
    )
    assert '"project_id", "project_name", "qalicb_name",' not in source, (
        "the hand-typed required-column set is back in from_csv"
    )


def test_from_csv_demands_exactly_the_schema_required_columns(tmp_path):
    """The derivation has to bite: drop each column and require a refusal."""
    from nmtcapp.data.schema import REQUIRED_PROJECT_FIELDS

    values = {
        "project_id": "P-1", "project_name": "P One",
        "qalicb_name": "P One QALICB LLC", "address": "1 Derived Way",
        "city": "Akron", "state": "OH", "sector": "healthcare",
        "project_type": "real_estate", "total_project_cost": "8000000",
        "qei_request": "5000000", "qlici_amount": "5000000",
        "expected_jobs_created": "20",
    }
    assert set(values) == set(REQUIRED_PROJECT_FIELDS), (
        "this fixture no longer matches the authoritative list"
    )

    for dropped in REQUIRED_PROJECT_FIELDS:
        cols = [c for c in values if c != dropped]
        path = tmp_path / f"drop_{dropped}.csv"
        path.write_text(
            ",".join(cols) + "\n" + ",".join(values[c] for c in cols) + "\n",
            encoding="utf-8",
        )
        with pytest.raises(ValueError) as exc:
            Pipeline.from_csv(str(path))
        assert dropped in str(exc.value), (
            f"dropping column {dropped!r} was not reported by name: {exc.value}"
        )


def test_the_sector_priority_tiers_are_derived_not_retyped():
    """Fourth statement of one classification, and one had already drifted.

    schema.TARGET_SECTORS carries a "priority" per sector. Three other places
    restated the tiers by hand:

        intelligence/sector_analysis._HIGH_PRIORITY_SECTORS   agreed
        visualization/maps._HIGH_PRIORITY / _MED_PRIORITY     DID NOT
        streamlit_app/utils.VALID_SECTORS                     retyped under
            the comment "(from schema)", which it was not

    maps' _MED_PRIORITY held {small_business, mixed_use} and omitted
    community_facility and clean_energy, which schema classes medium. The
    sector-mix chart gave those two bars the low-priority grey under a legend
    reading "Medium Priority (Small Business/Mixed Use)", while the Streamlit
    page rendering the same pipeline printed "Priority: Medium" for them in the
    table beside the chart. One screen, two classifications.
    """
    import inspect

    from nmtcapp.data.schema import (
        SECTORS_BY_PRIORITY, TARGET_SECTORS, VALID_SECTORS,
    )
    from nmtcapp.intelligence import sector_analysis
    from nmtcapp.visualization import maps

    assert set(SECTORS_BY_PRIORITY) == {"high", "medium", "low"}
    covered = set().union(*SECTORS_BY_PRIORITY.values())
    assert covered == set(TARGET_SECTORS), (
        f"a sector belongs to no tier: {sorted(set(TARGET_SECTORS) - covered)}"
    )

    assert sector_analysis._HIGH_PRIORITY_SECTORS is SECTORS_BY_PRIORITY["high"], (
        "sector_analysis holds its own copy of the high-priority tier again"
    )

    source = inspect.getsource(maps.plot_sector_distribution)
    assert "SECTORS_BY_PRIORITY" in source, (
        "the sector-mix chart no longer reads the tiers from schema"
    )
    assert '{"small_business", "mixed_use"}' not in source, (
        "the drifted hand-typed medium tier is back in visualization/maps"
    )
    # The legend must be built from the tiers too, or it goes on naming the
    # old pair after the sets are corrected. Matched on the Patch call rather
    # than the phrase: the comment recording the defect quotes it on purpose.
    assert 'label="Medium Priority' not in source, (
        "the legend still hard-names its sectors instead of deriving them"
    )
    assert "_tier_label" in source, "the legend is not built from the tiers"

    from streamlit_app.utils import VALID_SECTORS as STREAMLIT_SECTORS
    assert STREAMLIT_SECTORS is VALID_SECTORS, (
        "streamlit_app/utils retyped the sector vocabulary again; a sector "
        "added to schema would be rejected by the uploader"
    )
