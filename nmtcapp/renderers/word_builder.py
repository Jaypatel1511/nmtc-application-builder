"""Professional Word document builder for NMTC applications."""
from __future__ import annotations

import logging
import os
from datetime import date
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm

from nmtcapp.renderers._disclosure import (
    is_partial_unverified, qualified_pct, unverified_banner,
    unverified_ids, unverified_qualifier,
)
from nmtcapp.renderers._word_helpers import (
    add_styled_paragraph, add_styled_table, shade_cell, _hex_to_rgb,
)
from nmtcapp.renderers.styles import COLORS, TYPOGRAPHY, PAGE_LAYOUT
from nmtcapp.sections import ALL_SECTIONS
from nmtcapp.tables.distress_table import build_distress_table, build_distress_summary_table
from nmtcapp.tables.geographic_table import build_geographic_table
from nmtcapp.tables.impact_table import build_impact_table, build_impact_summary_table
from nmtcapp.tables.investor_table import (
    build_investor_identification_table, build_investor_commitment_table,
)
from nmtcapp.tables.pipeline_table import build_pipeline_table, build_pipeline_summary_table
from nmtcapp.tables.track_record_table import build_track_record_table

if TYPE_CHECKING:
    from nmtcapp.core.application import Application, ApplicationAnalysis

logger = logging.getLogger(__name__)

_PRIMARY = COLORS["primary"].lstrip("#")
_ACCENT = COLORS["accent"].lstrip("#")


class WordApplicationBuilder:
    """Builds a professional Word document for the full NMTC application.

    The output document includes a cover page, executive summary, all five
    sections (A–E), and all supporting attachments.

    Example::

        builder = WordApplicationBuilder(application, analysis)
        doc = builder.build()
        builder.save("./drafts/application.docx")
    """

    def __init__(self, application: "Application", analysis: "ApplicationAnalysis") -> None:
        self.application = application
        self.analysis = analysis

    def build(self) -> Document:
        """Build and return a python-docx Document.

        Example::

            doc = builder.build()
        """
        doc = Document()
        self._configure_document(doc)
        self._build_cover_page(doc)
        self._build_executive_summary(doc)
        self._build_readiness_callout(doc)

        for section_gen in ALL_SECTIONS:
            doc.add_page_break()
            section_gen.generate_word(doc, self.application, self.analysis)

        doc.add_page_break()
        self._build_appendix_pipeline(doc)
        doc.add_page_break()
        self._build_appendix_distress(doc)
        doc.add_page_break()
        self._build_appendix_geographic(doc)
        doc.add_page_break()
        self._build_appendix_impact(doc)
        doc.add_page_break()
        self._build_investor_section(doc)
        doc.add_page_break()
        self._build_appendix_track_record(doc)
        doc.add_page_break()
        self._build_methodology(doc)

        self._add_page_numbers(doc)
        return doc

    def save(self, path: str) -> None:
        """Save the document to a .docx file.

        Example::

            builder.save("./output/application.docx")
        """
        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
        doc = self.build()
        doc.save(path)
        size_kb = os.path.getsize(path) // 1024
        logger.info("Word document saved: %s (%d KB)", path, size_kb)

    # ------------------------------------------------------------------
    # Document configuration
    # ------------------------------------------------------------------

    def _configure_document(self, doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin    = Inches(PAGE_LAYOUT["margin_top"])
        section.bottom_margin = Inches(PAGE_LAYOUT["margin_bottom"])
        section.left_margin   = Inches(PAGE_LAYOUT["margin_left"])
        section.right_margin  = Inches(PAGE_LAYOUT["margin_right"])
        # Configure core style
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(TYPOGRAPHY["size_body"])

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _build_cover_page(self, doc: Document) -> None:
        app = self.application
        score = self.analysis.readiness_score

        # ---- Top blue banner ----
        banner_tbl = doc.add_table(rows=1, cols=1)
        banner_tbl.style = "Table Grid"
        banner_cell = banner_tbl.cell(0, 0)
        banner_cell.width = Inches(6.0)
        banner_tbl.rows[0].height = Inches(1.2)
        shade_cell(banner_cell, _PRIMARY)

        p_banner = banner_cell.paragraphs[0]
        p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p_banner.add_run("NEW MARKETS TAX CREDIT")
        r1.font.bold = True
        r1.font.size = Pt(22)
        r1.font.color.rgb = RGBColor(255, 255, 255)
        p_banner.add_run("\n")
        r2 = p_banner.add_run("ALLOCATION APPLICATION")
        r2.font.bold = True
        r2.font.size = Pt(18)
        r2.font.color.rgb = RGBColor(255, 255, 255)

        # ---- Gold accent line ----
        doc.add_paragraph()
        accent_tbl = doc.add_table(rows=1, cols=1)
        accent_tbl.style = "Table Grid"
        shade_cell(accent_tbl.cell(0, 0), _ACCENT)
        accent_tbl.rows[0].height = Cm(0.3)

        # ---- CDE name ----
        doc.add_paragraph()
        p_cde = doc.add_paragraph()
        p_cde.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_cde.add_run(app.cde.name)
        r.font.bold = True
        r.font.size = Pt(TYPOGRAPHY["size_h1"])
        r.font.color.rgb = RGBColor(*_hex_to_rgb(COLORS["primary"]))

        # ---- Application details ----
        details_tbl = doc.add_table(rows=4, cols=2)
        details_tbl.style = "Table Grid"
        details = [
            ("Application Round:", app.application_round),
            ("Requested NMTC Allocation:", f"${app.requested_allocation:,.0f}"),
            ("Preparation Date:", date.today().strftime("%B %d, %Y")),
            ("Readiness Assessment:",
             f"Grade {score.grade} — {score.overall_score:.1f}/100"
             + (" (PARTIAL)" if getattr(score, "partial", False) else "")),
        ]
        for i, (label, value) in enumerate(details):
            label_cell = details_tbl.cell(i, 0)
            value_cell = details_tbl.cell(i, 1)
            shade_cell(label_cell, COLORS["bg_light"].lstrip("#"))
            _set_cell_text(label_cell, label, bold=True)
            _set_cell_text(value_cell, value)

        # ---- Contact info ----
        contact = app.cde.contact
        doc.add_paragraph()
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_contact.add_run(
            f"Contact: {contact.get('name', '')} | {contact.get('email', '')} | "
            f"{contact.get('phone', '')}"
        )
        r.font.size = Pt(TYPOGRAPHY["size_caption"])
        r.font.color.rgb = RGBColor(*_hex_to_rgb(COLORS["text_muted"]))

        # ---- Confidentiality notice ----
        doc.add_paragraph()
        p_conf = doc.add_paragraph()
        p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_conf.add_run(
            "CONFIDENTIAL — Prepared for CDFI Fund Review Only | "
            "Do Not Distribute Without CDE Authorization"
        )
        r.font.size = Pt(TYPOGRAPHY["size_footnote"])
        r.font.italic = True
        r.font.color.rgb = RGBColor(*_hex_to_rgb(COLORS["text_muted"]))

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Executive summary
    # ------------------------------------------------------------------

    def _build_executive_summary(self, doc: Document) -> None:
        app = self.application
        pr = self.analysis.pipeline_result
        distress = pr.distress_breakdown

        add_styled_paragraph(doc, "Executive Summary", level=1, color_key="primary")

        degraded = getattr(pr, "eligibility_data_status", "ok") != "ok"
        partial_unverified = is_partial_unverified(pr)
        if degraded:
            notice = doc.add_paragraph()
            r = notice.add_run(
                "ELIGIBILITY DATA UNAVAILABLE — "
                f"{getattr(pr, 'eligibility_data_error', None) or 'reason unknown'}. "
                "Census tract eligibility and distress figures could not be "
                "verified and are excluded from this document. Do not submit "
                "until eligibility data has been restored and re-verified."
            )
            r.font.bold = True
            r.font.size = Pt(TYPOGRAPHY["size_body"])
            r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
            summary_text = (
                f"{app.cde.name} requests ${app.requested_allocation/1e6:.1f} million in "
                f"New Markets Tax Credit allocation for {app.application_round}. "
                f"Our {pr.total_projects}-project pipeline spans "
                f"{pr.geographic_diversity.get('states_count', 0)} states. Census "
                "tract eligibility and distress concentration are unverified — "
                "eligibility data was unavailable when this draft was generated."
            )
        elif partial_unverified:
            notice = doc.add_paragraph()
            r = notice.add_run("UNVERIFIED PROJECT LOCATIONS — " + unverified_banner(pr))
            r.font.bold = True
            r.font.size = Pt(TYPOGRAPHY["size_body"])
            r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
            summary_text = (
                f"{app.cde.name} requests ${app.requested_allocation/1e6:.1f} million in "
                f"New Markets Tax Credit allocation for {app.application_round}. "
                f"Our {pr.total_projects}-project pipeline spans "
                f"{pr.geographic_diversity.get('states_count', 0)} states, with "
                f"{distress.get('pct_deep_or_severe', 0):.0%} of QEI "
                f"{unverified_qualifier(pr)} committed to deep and severely "
                "distressed census tracts — figures reflect location-verified "
                "projects only."
            )
        else:
            summary_text = (
                f"{app.cde.name} requests ${app.requested_allocation/1e6:.1f} million in "
                f"New Markets Tax Credit allocation for {app.application_round}. "
                f"Our {pr.total_projects}-project pipeline spans "
                f"{pr.geographic_diversity.get('states_count', 0)} states, with "
                f"{distress.get('pct_deep_or_severe', 0):.0%} of QEI committed to deep and "
                f"severely distressed census tracts."
            )
        p = doc.add_paragraph(summary_text)
        p.runs[0].font.size = Pt(TYPOGRAPHY["size_body"])

        doc.add_paragraph()

        # Key metrics table
        impact = pr.aggregate_impact
        add_styled_paragraph(doc, "Key Metrics", level=2, color_key="secondary")
        unverified = "Unverified — eligibility data unavailable"

        def _elig_metric(value: float) -> str:
            if degraded:
                return unverified
            if partial_unverified:
                return qualified_pct(value, pr)
            return f"{value:.0%}"

        metrics_data = [
            ["Total Pipeline QEI", f"${pr.total_qei_request:,.0f}"],
            ["Total Project Cost", f"${pr.total_project_cost:,.0f}"],
            ["States Represented", str(pr.geographic_diversity.get("states_count", 0))],
            ["Deep/Severe Distress Concentration",
             _elig_metric(distress.get("pct_deep_or_severe", 0))],
            ["NMTC Eligibility Rate", _elig_metric(pr.eligibility_pct)],
            ["Jobs to Be Created", f"{impact.get('total_jobs_created', 0):,}"],
            ["Affordable Units to Be Built", f"{impact.get('total_units_built', 0):,}"],
            ["Jobs per $1MM QEI", f"{impact.get('jobs_per_million_qei', 0):.1f}"],
        ]
        add_styled_table(doc, ["Metric", "Value"], metrics_data)

    def _build_readiness_callout(self, doc: Document) -> None:
        """Add a readiness score callout box after the executive summary."""
        score = self.analysis.readiness_score
        doc.add_paragraph()
        add_styled_paragraph(doc, "Application Readiness Assessment", level=2,
                              color_key="secondary")

        score_tbl = doc.add_table(rows=1, cols=2)
        score_tbl.style = "Table Grid"

        # Left: grade
        left = score_tbl.cell(0, 0)
        shade_cell(left, _PRIMARY)
        left.width = Inches(1.2)
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Grade\n{score.grade}")
        r.font.bold = True
        r.font.size = Pt(28)
        r.font.color.rgb = RGBColor(*_hex_to_rgb(COLORS["accent"]))

        # Right: component scores
        right = score_tbl.cell(0, 1)
        shade_cell(right, COLORS["bg_light"].lstrip("#"))
        p2 = right.paragraphs[0]
        partial_line = ""
        if getattr(score, "partial", False):
            partial_line = f"PARTIAL — {score.partial_note}\n"
        r2 = p2.add_run(
            partial_line +
            f"Overall Score: {score.overall_score:.1f}/100"
            + ("  (PARTIAL)" if getattr(score, "partial", False) else "") + "\n" +
            "\n".join(
                f"  {k.replace('_', ' ').title()}: {v:.0f}/100"
                for k, v in score.component_scores.items()
            )
        )
        r2.font.size = Pt(TYPOGRAPHY["size_body"])

    # ------------------------------------------------------------------
    # Appendices
    # ------------------------------------------------------------------

    def _build_appendix_pipeline(self, doc: Document) -> None:
        """Appendix A: 6-col summary in portrait + full detail in landscape."""
        add_styled_paragraph(doc, "Appendix A: Pipeline Detail", level=1, color_key="primary")

        # Summary table in portrait — fits on page
        summary_df = build_pipeline_summary_table(self.application.pipeline)
        _write_df_to_doc(doc, summary_df)

        # Note pointing to Excel for full detail
        p = doc.add_paragraph(
            "Full 33-column pipeline detail (deal economics, QLICI structure, timeline) "
            "is provided in the accompanying Excel workbook, Pipeline Detail tab. "
            "The landscape pages below contain key pipeline fields for reference."
        )
        p.runs[0].font.italic = True
        p.runs[0].font.size = Pt(TYPOGRAPHY["size_caption"])

        # Landscape section with the full table (key 12-col subset)
        self._switch_to_landscape(doc)
        add_styled_paragraph(doc, "Appendix A (continued): Pipeline Detail — Full View",
                              level=2, color_key="secondary")
        full_df = build_pipeline_table(self.application.pipeline, self.application.cde)
        # Key 12 columns that fit in landscape with 8pt
        landscape_cols = [
            "Project ID", "Project Name", "City", "State",
            "Distress Level", "NMTC Eligible (Y/N)", "NMTC Native Area (Y/N)",
            "QEI Request ($)", "Total Project Cost ($)",
            "Jobs Created", "Jobs Retained", "Sector (NAICS)",
        ]
        existing = [c for c in landscape_cols if c in full_df.columns]
        _write_df_to_doc(doc, full_df[existing], max_rows=50, font_size=8)
        self._switch_to_portrait(doc)

    def _build_appendix_distress(self, doc: Document) -> None:
        """Appendix B: distress table in landscape (15 cols fits at 8pt)."""
        self._switch_to_landscape(doc)
        add_styled_paragraph(doc, "Appendix B: Distress Documentation", level=1,
                              color_key="primary")
        df = build_distress_table(self.application.pipeline)
        _write_df_to_doc(doc, df, font_size=8)
        self._switch_to_portrait(doc)

    def _build_appendix_geographic(self, doc: Document) -> None:
        add_styled_paragraph(doc, "Appendix C: Geographic Targeting", level=1,
                              color_key="primary")
        df = build_geographic_table(self.application.pipeline)
        _write_df_to_doc(doc, df)

    def _build_appendix_impact(self, doc: Document) -> None:
        """Appendix D: 7-col summary in portrait."""
        add_styled_paragraph(doc, "Appendix D: Impact Projections", level=1,
                              color_key="primary")
        df = build_impact_summary_table(self.application.pipeline)
        _write_df_to_doc(doc, df, max_rows=25)
        p = doc.add_paragraph(
            "Full impact detail (units, square footage, OZ/HMR flags) is provided "
            "in the accompanying Excel workbook, Impact Projections tab."
        )
        p.runs[0].font.italic = True
        p.runs[0].font.size = Pt(TYPOGRAPHY["size_caption"])

    def _build_appendix_track_record(self, doc: Document) -> None:
        add_styled_paragraph(doc, "Appendix E: CDE Track Record", level=1, color_key="primary")
        df = build_track_record_table(self.application.cde)
        _write_df_to_doc(doc, df)

    # ------------------------------------------------------------------
    # Investor section (Section D body) — split into two narrow tables
    # ------------------------------------------------------------------

    def _build_investor_section(self, doc: Document) -> None:
        """Add the investor table split into identification + commitment detail."""
        from nmtcapp.tables.investor_table import (
            INVESTOR_TABLE_NOTE, INVESTOR_TABLE_TITLE,
        )
        add_styled_paragraph(doc, INVESTOR_TABLE_TITLE, level=2,
                              color_key="secondary")
        doc.add_paragraph(INVESTOR_TABLE_NOTE)
        id_df = build_investor_identification_table(self.application)
        commit_df = build_investor_commitment_table(self.application)
        if not id_df.empty:
            _write_df_to_doc(doc, id_df)
        doc.add_paragraph()
        if not commit_df.empty:
            _write_df_to_doc(doc, commit_df)

    # ------------------------------------------------------------------
    # Orientation helpers
    # ------------------------------------------------------------------

    def _switch_to_landscape(self, doc: Document) -> None:
        """Add a page break into a new landscape section."""
        new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
        new_sec.orientation = WD_ORIENT.LANDSCAPE
        new_sec.page_width = Inches(11)
        new_sec.page_height = Inches(8.5)
        new_sec.left_margin = Inches(0.75)
        new_sec.right_margin = Inches(0.75)
        new_sec.top_margin = Inches(0.75)
        new_sec.bottom_margin = Inches(0.75)

    def _switch_to_portrait(self, doc: Document) -> None:
        """Add a page break back to portrait layout."""
        new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
        new_sec.orientation = WD_ORIENT.PORTRAIT
        new_sec.page_width = Inches(8.5)
        new_sec.page_height = Inches(11)
        new_sec.left_margin = Inches(PAGE_LAYOUT["margin_left"])
        new_sec.right_margin = Inches(PAGE_LAYOUT["margin_right"])
        new_sec.top_margin = Inches(PAGE_LAYOUT["margin_top"])
        new_sec.bottom_margin = Inches(PAGE_LAYOUT["margin_bottom"])

    def _eligibility_source_line(self) -> str:
        """Cite the eligibility source only when it actually loaded.

        markdown_builder._methodology_note branched on this correctly; the
        Word and PDF methodology appendices asserted the citation
        unconditionally, so a run whose eligibility download failed still
        told a federal reviewer the CDFI Fund table had been used.
        """
        pr = self.analysis.pipeline_result
        if getattr(pr, "eligibility_data_status", "ok") != "ok":
            reason = getattr(pr, "eligibility_data_error", None) or "reason unknown"
            return (
                f"ELIGIBILITY: UNAVAILABLE — {reason}. No CDFI Fund eligibility "
                "data was loaded or used anywhere in this draft."
            )
        base = (
            "ELIGIBILITY: CDFI Fund NMTC Eligibility Table using 2016–2020 ACS "
            "5-Year Estimates. Tract eligibility determined by Low Income "
            "Community (LIC) criteria per IRC §45D."
        )
        if is_partial_unverified(pr):
            ids = unverified_ids(pr)
            return (
                f"{base} {len(ids)} of {pr.total_projects} projects could not be "
                f"location-verified (no census tract assigned): {', '.join(ids)}."
            )
        return base

    def _build_methodology(self, doc: Document) -> None:
        from nmtcapp import __version__
        add_styled_paragraph(doc, "Appendix F: Methodology", level=1, color_key="primary")
        p = doc.add_paragraph(
            f"Generated by nmtc-application-builder v{__version__} on {date.today().isoformat()}.\n\n"
            f"{self._eligibility_source_line()}\n\n"
            "DISTRESS LEVELS: Deep distress = poverty rate >30% or unemployment >1.5× national "
            "average. Severe distress = LIC plus additional qualifying factors. "
            "(Source: CDFI Fund NMTC Program guidance and the NMTC Allocation "
            "Application Review Process. NMTC rounds are announced by a NOAA — "
            "Notice of Allocation Availability — and the CY 2026 NOAA is not "
            "yet published.)\n\n"
            "DEAL ECONOMICS: Computed using nmtc-calc library. Standard leveraged NMTC structure "
            "with credit price $0.83/credit, CDE fee 2.5% of QEI, 7-year compliance period.\n\n"
            "IMPACT BENCHMARKS: CDFI Fund Annual Reports, FY2018–FY2023. "
            "Average jobs per $1MM QEI: 12.0 FTE. Top quartile: ≥20.0 FTE per $1MM.\n\n"
            "READINESS SCORE: An UNSOURCED HOUSE HEURISTIC of this tool "
            "(eligibility 25%, distress concentration 25%, geographic diversity "
            "15%, impact metrics 20%, validation 10%, completeness 5%). The "
            "weights are this tool's own judgement; they are not calibrated "
            "against award data, the CDFI Fund publishes no such weighting, and "
            "the score does not predict an award outcome."
        )
        p.runs[0].font.size = Pt(TYPOGRAPHY["size_body"])

    # ------------------------------------------------------------------
    # Utility: page numbers in footer
    # ------------------------------------------------------------------

    def _add_page_numbers(self, doc: Document) -> None:
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.clear()
            run = p.add_run()
            run.font.size = Pt(TYPOGRAPHY["size_footnote"])
            run.font.color.rgb = RGBColor(*_hex_to_rgb(COLORS["text_muted"]))
            # Add field code for page number
            fldChar1 = OxmlElement("w:fldChar")
            fldChar1.set(qn("w:fldCharType"), "begin")
            instrText = OxmlElement("w:instrText")
            instrText.text = "PAGE"
            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "end")
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            conf_run = p.add_run(
                f"  |  {self.application.cde.name} — NMTC {self.application.application_round} "
                f"Application  |  CONFIDENTIAL"
            )
            conf_run.font.size = Pt(TYPOGRAPHY["size_footnote"])
            conf_run.font.color.rgb = RGBColor(*_hex_to_rgb(COLORS["text_muted"]))


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def _write_df_to_doc(doc: Document, df, max_rows: int = 50, font_size: int = None) -> None:
    """Write a DataFrame as a styled table to the document."""
    if df is None or df.empty:
        doc.add_paragraph("No data available.")
        return
    display = df.head(max_rows)
    headers = list(display.columns)
    data = [list(row) for _, row in display.iterrows()]
    add_styled_table(doc, headers, data, totals_last=True, font_size=font_size)
    if len(df) > max_rows:
        p = doc.add_paragraph(f"Showing {max_rows} of {len(df)} rows. See Excel attachment for complete data.")
        p.runs[0].font.italic = True
        p.runs[0].font.size = Pt(TYPOGRAPHY["size_caption"])


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    """Set plain text in a table cell, clearing existing content."""
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.bold = bold
    run.font.size = Pt(TYPOGRAPHY["size_body"])
