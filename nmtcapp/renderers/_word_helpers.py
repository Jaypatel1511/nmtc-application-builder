"""Internal helpers shared by word_builder and section generators."""
from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from nmtcapp.renderers.styles import COLORS, TYPOGRAPHY


def shade_cell(cell, hex_color: str) -> None:
    """Set the background fill of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shd element
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_borders(cell, **kwargs) -> None:
    """Set borders on a table cell. kwargs: top, bottom, left, right with color/size."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if side in kwargs or True:  # always set clean borders
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), kwargs.get(side, "single"))
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), COLORS["border"].lstrip("#"))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_styled_paragraph(doc, text: str, level: int = 0, center: bool = False,
                          color_key: str = "text_body", bold: bool = False,
                          font_size: int = None) -> object:
    """Add a styled paragraph to the document."""
    if level == 1:
        p = doc.add_heading("", level=1)
    elif level == 2:
        p = doc.add_heading("", level=2)
    elif level == 3:
        p = doc.add_heading("", level=3)
    else:
        p = doc.add_paragraph()

    run = p.add_run(text)
    run.font.bold = bold or level in (1, 2, 3)
    run.font.color.rgb = RGBColor(*_hex_to_rgb(COLORS.get(color_key, COLORS["text_body"])))
    if font_size:
        run.font.size = Pt(font_size)
    elif level == 1:
        run.font.size = Pt(TYPOGRAPHY["size_h1"])
    elif level == 2:
        run.font.size = Pt(TYPOGRAPHY["size_h2"])
    elif level == 3:
        run.font.size = Pt(TYPOGRAPHY["size_h3"])
    else:
        run.font.size = Pt(TYPOGRAPHY["size_body"])

    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_styled_table(doc, headers: list, data: list,
                     col_widths: list = None, totals_last: bool = False,
                     font_size: int = None) -> object:
    """Add a professionally styled table with dark header and alternating rows."""
    if not headers:
        return None

    body_pt = font_size or TYPOGRAPHY["size_table_body"]
    header_pt = max(body_pt - 1, 7) if font_size else TYPOGRAPHY["size_table_header"]

    n_rows = 1 + len(data)
    n_cols = len(headers)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.autofit = True

    # Header row
    hrow = table.rows[0]
    hrow.height = Inches(0.30)
    for j, hdr in enumerate(headers):
        cell = hrow.cells[j]
        _clear_and_set(cell, str(hdr))
        shade_cell(cell, COLORS["row_header"])
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(header_pt)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        bg = COLORS["row_total"] if (totals_last and i == len(data) - 1) else (
            COLORS["row_even"] if i % 2 == 0 else COLORS["row_odd"]
        )
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            # NOT truncated. An 80-char cut with an ellipsis silently removed
            # the tail of any long cell — including Section B's severe-distress
            # row, whose value ended "…Review Process; the CY 2026 NOAA is not
            # yet published)". Word rendered "…NMTC Allocation Applicatio..."
            # and dropped the one sentence warning that the current round's
            # rules are unknown, in the format most likely to be submitted.
            # Markdown and PDF carried it in full.
            #
            # Word wraps cell text on its own; there is no rendering reason to
            # cut. If a width cap is ever reintroduced it must exempt any cell
            # carrying a disclosure — see test_word_cells_are_not_truncated.
            text = _fmt_cell(val)
            _clear_and_set(cell, text)
            shade_cell(cell, bg)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(body_pt)
            if totals_last and i == len(data) - 1:
                run.font.bold = True

    # Column widths
    if col_widths:
        for j, w in enumerate(col_widths):
            if j < n_cols:
                for cell in table.columns[j].cells:
                    cell.width = w

    return table


def write_section_to_doc(doc, content: dict) -> None:
    """Write a structured section content dict to a python-docx Document."""
    section_id = content.get("section_id", "")
    title = content.get("title", "")
    add_styled_paragraph(doc, f"Section {section_id}: {title}", level=1,
                         color_key="primary")
    doc.add_paragraph()

    for sub in content.get("subsections", []):
        heading = sub.get("heading", "")
        body = sub.get("body", "")
        sub_type = sub.get("type", "narrative")

        add_styled_paragraph(doc, heading, level=2, color_key="secondary")

        if sub_type == "table_ref":
            if isinstance(body, dict):
                headers = ["Item", "Value"]
                data = [[k, v] for k, v in body.items()]
                add_styled_table(doc, headers, data)
            elif isinstance(body, str):
                p = doc.add_paragraph(body)
                p.runs[0].font.italic = True
                p.runs[0].font.size = Pt(TYPOGRAPHY["size_body"])
        elif sub_type == "list":
            if isinstance(body, list):
                for item in body:
                    p = doc.add_paragraph(str(item), style="List Bullet")
                    if p.runs:
                        p.runs[0].font.size = Pt(TYPOGRAPHY["size_body"])
            else:
                _add_body_text(doc, str(body))
        else:
            _add_body_text(doc, str(body))

        doc.add_paragraph()


def _add_body_text(doc, text: str) -> None:
    """Add body text, splitting on double newlines into separate paragraphs."""
    for para in text.split("\n\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph(para)
            if p.runs:
                p.runs[0].font.size = Pt(TYPOGRAPHY["size_body"])


def _truncate(text: str, max_len: int) -> str:
    """Truncate cell text with ellipsis if it exceeds max_len.

    NO LONGER USED FOR TABLE CELLS. Kept only because it is imported
    elsewhere; see the note at the former call site. Do not reintroduce it
    for any cell that can carry a disclosure or a qualifier.
    """
    if len(text) <= max_len:
        return text
    return text[:max_len - 3] + "..."


def _clear_and_set(cell, text: str) -> None:
    """Clear a cell's paragraphs and set plain text."""
    for para in cell.paragraphs[:-1]:
        p = para._element
        p.getparent().remove(p)
    cell.paragraphs[0].clear()
    cell.paragraphs[0].add_run(text)


def _fmt_cell(val) -> str:
    """Format a cell value for display."""
    if val is None:
        return ""
    if isinstance(val, float):
        if val > 1000:
            return f"${val:,.0f}"
        return f"{val:.2f}"
    return str(val)


def _hex_to_rgb(hex_color: str) -> tuple:
    """Convert #RRGGBB to (R, G, B) tuple of ints."""
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
